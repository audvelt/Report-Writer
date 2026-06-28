"""
Author: Daniel Fraser
Created: 12/2025
Description:
    Samsung Quality Lab report writing tool.
"""

import sys
import os
import json
import shutil
import threading
import platform
import ctypes
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                                 QHBoxLayout, QLabel, QPushButton, QCheckBox, 
                                 QTextEdit, QFileDialog, QScrollArea, QGroupBox,
                                 QLineEdit, QMessageBox, QTabWidget, QGridLayout,
                                 QToolButton, QProgressDialog, QDialog)
    from PyQt5.QtCore import Qt, QTimer, QThread, pyqtSignal
    from PyQt5.QtGui import QPixmap, QFont, QIcon
except ImportError as e:
    print(f"Error importing PyQt5: {e}")
    print("Please install PyQt5: pip install PyQt5")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError as e:
    print(f"Error importing python-docx: {e}")
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

BEARING_WEAR_PATTERNS = {
    "No issues detected": {
        "observations": [
            "No significant wear patterns observed."
        ]
    },
    "Electrical Fluting (Bearing Flare)": {
        "observations": [
            "Closely spaced corrugation patterns visible on raceway surfaces.",
            "Flute-like marks observed across bearing races.",
            "Gray frosted appearance present on raceway surfaces.",
            "Washboard-like texture evident on rolling surfaces.",
            "Material removal visible in regular, repeating patterns.",
            "Symmetric damage pattern across bearing circumference.",
            "Burnt grease observed in bearing."
        ]
    },
    "Pitting/Spalling": {
        "observations": [
            "Small craters visible on raceway surfaces.",
            "Material removal observed on rolling surfaces.",
            "Progressive surface deterioration noted.",
            "Irregular surface texture present.",
            "Flaking of surface material evident.",
            "Subsurface initiated damage visible."
        ]
    },
    "Excessive Heat Discoloration": {
        "observations": [
            "Blue discoloration visible on bearing surfaces.",
            "Brown tinting present on bearing components.",
            "Straw-colored oxidation observed.",
            "Temper colors evident across races.",
            "Heat patterns visible on rolling elements.",
            "Surface oxidation layers present."
        ]
    },
    "Cage Wear/Damage": {
        "observations": [
            "Material loss observed on cage pockets.",
            "Cracks visible in cage structure.",
            "Distortion of cage components noted.",
            "Bending of cage bars evident.",
            "Broken cage segments present.",
            "Wear marks visible on cage surfaces."
        ]
    },
    "Corrosion/Rust": {
        "observations": [
            "Red-brown oxidation visible on races.",
            "Rust deposits present on rolling elements.",
            "Surface pitting from corrosion observed.",
            "Moisture staining evident on components.",
            "Oxidation layers covering bearing surfaces.",
            "Corrosion products accumulated in bearing."
        ]
    },
    "Brinelling/False Brinelling": {
        "observations": [
            "Evenly spaced depressions visible in raceways.",
            "Indentations corresponding to rolling element positions.",
            "Material displacement evident at contact points.",
            "Regular impression pattern matching element spacing.",
            "Permanent deformation observed in raceway.",
            "Circular indentation marks present."
        ]
    },
    "Lubrication Failure": {
        "observations": [
            "Dried lubricant present in bearing.",
            "Hardened grease observed.",
            "Discolored lubricant visible.",
            "Absence of lubricant in critical areas noted.",
            "Contaminated lubricant with foreign particles.",
            "Moisture present in lubricant."
        ]
    },
    "Misalignment Wear": {
        "observations": [
            "Asymmetric wear patterns observed.",
            "Uneven contact marks visible on raceways.",
            "One-sided loading evident on rolling elements.",
            "Diagonal wear pattern across bearing width.",
            "Heavier wear on one side of raceway.",
            "Non-uniform contact areas present."
        ]
    }
}

BEARING_REFERENCE_IMAGES = {
    "Electrical Fluting (Bearing Flare)": "fluting.png",
    "Pitting/Spalling": "spalling.png",
    "Excessive Heat Discoloration": "heat.png",
    "Cage Wear/Damage": "cage.png",
    "Corrosion/Rust": "corrosion.png",
    "Brinelling/False Brinelling": "brinelling.png",
    "Lubrication Failure": "lubrication.png",
    "Misalignment Wear": "misalignment.png"
}

ELECTRICAL_CONNECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Electrical connections inspected and found to be acceptable with no significant issues observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Loose wire connections observed.",
                "Corroded terminals present.",
                "Damaged insulation noted on wiring.",
                "Burnt or discolored connections evident.",
                "Missing or damaged connection hardware.",
                "Improper wire gauge observed.",
                "Evidence of arcing at connection points.",
                "Moisture ingress at electrical connections."
            ]
        }
    }
}

MOTOR_SHAFT_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Motor shaft inspected and found to be acceptable with no significant wear observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Shaft surface shows scoring marks.",
                "Rust or corrosion present on shaft.",
                "Shaft exhibits wear at bearing contact areas.",
                "Shaft runout exceeds acceptable limits.",
                "Shaft shows signs of bending or deformation.",
                "Keyway damage observed.",
                "Surface finish degradation noted.",
                "Shaft diameter reduction evident."
            ]
        }
    }
}

HOUSING_WEAR_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Motor housing inspected and found to be acceptable with no significant wear observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Scratches visible on bore surfaces.",
                "Grooves present on mounting surfaces.",
                "Material removal evident in bearing seat.",
                "Out-of-round condition observed.",
                "Scoring marks visible on bore.",
                "Loss of surface finish noted.",
                "Wear marks at mounting interfaces.",
                "Corrosion and rust deposits evident on housing surfaces.",
                "Moisture staining evident in housing cavity.",
                "Seal groove damage observed.",
                "Distortion of groove geometry noted.",
                "Thread damage present on fastener holes.",
                "Stripped or cross-threaded holes observed.",
                "Cracks visible in housing structure.",
                "Material separation noted at critical locations.",
                "Mounting surface wear present.",
                "Elongated bolt holes visible.",
                "Deformation at mounting flange noted.",
                "Foreign particles and debris visible in housing.",
                "Contaminated lubricant present."
            ]
        }
    }
}

ELECTRICAL_TESTING_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Electrical testing completed with no issues detected."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Grounded windings detected.",
                "Open windings detected.",
                "Resistance values out of acceptable range."
            ]
        }
    }
}

OIL_EVALUATION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Oil evaluation shows normal condition with acceptable volume and color."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Oil volume below normal levels.",
                "Oil volume above normal levels.",
                "Oil color indicates contamination.",
                "Oil color darker than acceptable.",
                "Metal particles present in oil.",
                "Moisture contamination detected in oil.",
                "Oil exhibits burnt odor.",
                "Oil exhibits acidic odor.",
                "Oil viscosity degraded."
            ]
        }
    }
}

SLEEVE_BEARING_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Sleeve bearing inspected and found to be acceptable with no significant wear observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Sleeve bearing is seized on the shaft and further inspection cannot proceed.",
                "Sleeve bearing shows signs of damage.",
                "Excessive wear visible on sleeve bearing surface.",
                "Material loss evident on bearing surfaces.",
                "Surface pitting visible on bearing.",
                "Discoloration indicating heat damage on bearing.",
                "Scoring marks observed on bearing surface.",
                "Foreign material deposits on bearing surfaces.",
                "Corrosion present on bearing surfaces."
            ]
        }
    }
}

SCROLL_PLATE_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Scroll plate inspected and found to be acceptable with no significant wear observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Corrosion present on scroll plate surfaces.",
                "Scoring marks observed on scroll plate.",
                "Surface pitting visible on scroll components.",
                "Discoloration indicating heat damage.",
                "Cracks visible in scroll plate structure.",
                "Foreign material deposits on scroll surfaces.",
                "Seal surface damage observed.",
                "Scroll plate is broken/shattered."
            ]
        }
    }
}

COIL_VISUAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Evaporator coil inspected and found to be in acceptable condition with no significant issues observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Corrosion visible on fins or tubes.",
                "Corrosion visible on sheet metal.",
                "Bent or damaged fins observed.",
                "Heavy dirt and debris accumulation present.",
                "Evidence of refrigerant leak detected.",
                "Ice formation observed on coil surfaces.",
                "Blocked airflow passages evident.",
                "Coil casing damage visible.",
                "Insulation deterioration noted.",
                "Fin spacing irregular or compressed.",
                "Tube damage or deformation observed.",
                "Biological growth (mold/mildew) present.",
                "Excessive oil residue on coil surfaces."
            ]
        }
    }
}

RECEIVER_VISUAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Receiver inspected and found to be in acceptable condition with no significant issues observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Corrosion visible on receiver shell.",
                "Dents or physical damage observed on housing.",
                "Evidence of refrigerant leak detected at joints or welds.",
                "Oil staining present on exterior surfaces.",
                "Insulation deterioration or damage noted.",
                "Mounting bracket corrosion or damage visible.",
                "Excessive frost or ice buildup on receiver exterior.",
                "Swelling or bulging of receiver shell observed.",
                "Paint blistering or peeling indicating moisture intrusion.",
                "Connection fitting corrosion or damage present.",
                "Vibration-induced wear marks visible.",
                "Weld seam degradation or cracking observed.",
                "Foreign material deposits on external surfaces.",
                "Broken or missing fusible plug."
            ]
        }
    }
}

PCB_VISUAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "PCB/Inverter board inspected and found to be in acceptable condition with no significant issues observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Burnt or scorched components visible on board.",
                "Capacitor bulging, leaking, or damage observed.",
                "Cracked or broken solder joints present.",
                "Corrosion or oxidation on board traces or components.",
                "Physical impact damage or cracked PCB substrate observed.",
                "Delamination or bubbling of board layers detected.",
                "Discoloration from overheating evident on board surface.",
                "Damaged or burnt MOSFET or IGBT transistors visible.",
                "Electrolytic capacitor venting or swelling detected.",
                "Trace damage or lifted pads observed.",
                "Missing or dislodged surface-mount components found.",
                "Insulation breakdown or carbon tracking visible.",
                "Connector pins corroded, bent, or damaged.",
                "Relay contact pitting or arcing damage present.",
                "Resistor failure or burn marks detected.",
                "Damaged or cracked voltage regulators observed.",
                "Heat sink separation or poor thermal contact evident.",
                "Moisture intrusion or condensation damage present.",
                "Contamination or foreign debris accumulation on board.",
                "Fuse blown or visibly damaged."
            ]
        }
    }
}

PCB_ELECTRICAL_TESTING_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Diode testing completed with all readings within specification. No failures detected."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Diode short circuit failure detected.",
                "Diode open circuit failure detected.",
                "Diode forward voltage out of specification.",
                "Diode reverse leakage current excessive.",
                "Freewheeling diode failure detected.",
                "Bridge rectifier diode failure detected.",
                "Zener diode breakdown voltage out of specification.",
                "Schottky diode forward voltage drop abnormal."
            ]
        }
    }
}

# TO ADD NEW EQUIPMENT: Create pattern dictionaries following this structure
VALVE_EXTERNAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Four-way valve external inspection completed with no significant issues observed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Corrosion visible on valve body or connections.",
                "Physical damage to valve housing observed.",
                "Oil leakage detected around valve seals.",
                "Refrigerant leak evidence at valve ports.",
                "Loose or damaged mounting hardware.",
                "Electrical connector damage or corrosion visible.",
                "Capillary tube damage or kinking observed.",
                "Pilot valve assembly showing wear or damage.",
                "Port caps missing or damaged.",
                "Foreign debris accumulation on valve body."
            ]
        }
    }
}

VALVE_ELECTRICAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Solenoid coil electrical inspection completed with no issues detected."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Solenoid coil resistance out of specification.",
                "Open circuit detected in solenoid coil.",
                "Short circuit detected in solenoid coil.",
                "Magnetic coil winding damage observed.",
                "Insulation breakdown detected on coil wires.",
                "Solenoid plunger sticking or seized.",
                "Electrical connector pins corroded or damaged.",
                "Coil overheating evidence visible."
            ]
        }
    }
}

VALVE_MECHANICAL_INSPECTION_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Valve mechanical inspection completed with smooth operation confirmed."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Valve body bore shows excessive wear.",
                "Piston movement restricted or binding.",
                "Slide valve difficult to actuate left and right.",
                "Internal spring damage or weakness detected.",
                "Piston rings worn or damaged.",
                "Valve seats showing excessive wear or pitting.",
                "Cylinder wall scoring or scratches present.",
                "O-ring grooves damaged or worn."
            ]
        }
    }
}

VALVE_INTERNAL_CYLINDER_A_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Internal inspection completed with valve chamber bore, slider assembly, support frame contact surfaces, pilot flow interfaces, and valve seat areas within specification."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Valve chamber bore scoring, abrasion, or localized wear observed.",
                "Valve chamber bore surface finish outside specification.",
                "Slider contact surface wear, polishing, or scoring observed.",
                "Slider deformation, warping, or loss of flatness detected.",
                "Slider edge lifting, curling, or heat-induced distortion observed, potentially restricting movement.",
                "Evidence of slider sticking, drag marks, or interrupted travel.",
                "Uneven wear pattern indicating slider misalignment or improper load distribution.",
                "Valve seat edge damage, deformation, or erosion observed at internal interfaces.",
                "Pilot flow passage wear, restriction, or contamination observed.",
                "Foreign debris or particulate contamination present within valve chamber.",
                "Metal particles embedded in valve chamber or slider surfaces.",
                "Thermal discoloration indicating localized overheating near slider or valve seat areas.",
                "Evidence of abnormal contact or rubbing between support frame and slider.",
                "Support frame contact surface wear or imprinting observed."
            ]
        }
    }
}


VALVE_INTERNAL_CYLINDER_B_PATTERNS = {
    "Status: OK": {
        "observations": [
            "Internal inspection completed with piston travel bore, piston assembly, polymer guide ring, internal surface finish, and dimensional characteristics within specification."
        ]
    },
    "Status: NOT OK": {
        "observations": [],
        "sub_observations": {
            "": [
                "Internal bore scoring, galling, or abrasion observed along piston travel path.",
                "Surface finish degradation beyond specification on internal bore.",
                "Deep scratches or longitudinal scoring present on bore surface.",
                "Pitting observed on internal bore surfaces.",
                "Piston surface wear, scoring, or material transfer observed.",
                "Polymer piston guide ring wear, deformation, cracking, or material loss observed.",
                "Guide ring extrusion or flattening indicating excessive thermal or mechanical load.",
                "Localized wear indicating piston misalignment or uneven loading.",
                "Metal debris or particulate contamination present within piston bore.",
                "Corrosion or chemical attack observed on bore or piston surfaces.",
                "Dimensional deviation or out-of-round condition detected in internal bore.",
                "Evidence of piston sticking, hesitation, or interrupted travel.",
                "Thermal discoloration or heat-affected areas observed on piston or bore surfaces."
            ]
        }
    }
}

# Add page number field
def add_field(run, field_code):
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = field_code

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


class ImageUploadWidget(QWidget):
    def __init__(self, section_name):
        super().__init__()
        self.section_name = section_name
        self.image_paths = []
        self.thumbnail_widgets = []
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        

        btn_layout = QHBoxLayout()
        self.upload_btn = QPushButton(f"Upload Images for {self.section_name}")
        self.upload_btn.clicked.connect(self.upload_images)
        btn_layout.addWidget(self.upload_btn)
        
        self.clear_all_btn = QPushButton("Clear All")
        self.clear_all_btn.clicked.connect(self.clear_images)
        btn_layout.addWidget(self.clear_all_btn)
        
        layout.addLayout(btn_layout)
        

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setMinimumHeight(120)
        self.scroll_area.setStyleSheet("border: 2px dashed #ccc;")
        
        self.scroll_area.setAcceptDrops(True)
        self.scroll_area.dragEnterEvent = self.dragEnterEvent
        self.scroll_area.dropEvent = self.dropEvent

        self.thumbnails_container = QWidget()
        self.thumbnails_layout = QHBoxLayout()
        self.thumbnails_layout.setAlignment(Qt.AlignLeft)
        self.thumbnails_container.setLayout(self.thumbnails_layout)
        
        self.scroll_area.setWidget(self.thumbnails_container)
        layout.addWidget(self.scroll_area)
        

        self.no_images_label = QLabel("No images uploaded\n(or drag and drop images here)")
        self.no_images_label.setAlignment(Qt.AlignCenter)
        self.no_images_label.setStyleSheet("color: #999; padding: 40px;")
        self.thumbnails_layout.addWidget(self.no_images_label)
        
        self.setLayout(layout)
    
    def upload_images(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, f"Select Images for {self.section_name}", "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.gif)"
        )
        if files:
            for file_path in files:
                if file_path not in self.image_paths:
                    self.image_paths.append(file_path)
                    self.add_thumbnail(file_path)
            self.update_display()
    
    def add_image(self, img_path):
        """Programmatically add an image (used for loading saved data)"""
        if img_path not in self.image_paths and os.path.exists(img_path):
            self.image_paths.append(img_path)
            self.add_thumbnail(img_path)
            self.update_display()
    
    def add_thumbnail(self, img_path):
        """Add a thumbnail widget for an image"""
        thumb_widget = QWidget()
        thumb_layout = QVBoxLayout()
        thumb_layout.setContentsMargins(5, 5, 5, 5)
        thumb_layout.setSpacing(5)
        

        img_label = QLabel()
        
        # Load image with PIL and apply EXIF orientation
        try:
            from PIL import Image as PILImage
            from PIL import ImageOps
            from io import BytesIO
            
            # Open image and apply EXIF transpose
            pil_img = PILImage.open(img_path)
            pil_img = ImageOps.exif_transpose(pil_img)
            
            # Convert PIL image to QPixmap
            # Save to BytesIO first, then load as QPixmap
            buffer = BytesIO()
            pil_img.save(buffer, format='PNG')
            buffer.seek(0)
            
            pixmap = QPixmap()
            pixmap.loadFromData(buffer.read())
            
            if not pixmap.isNull():
                scaled_pixmap = pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                img_label.setPixmap(scaled_pixmap)
            else:
                img_label.setText("Error")
        except:
            # Fallback to direct loading if PIL fails
            pixmap = QPixmap(img_path)
            if not pixmap.isNull():
                scaled_pixmap = pixmap.scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                img_label.setPixmap(scaled_pixmap)
            else:
                img_label.setText("Error")
        
        img_label.setAlignment(Qt.AlignCenter)
        img_label.setStyleSheet("border: 1px solid #ddd; background: white; padding: 5px;")
        img_label.setFixedSize(110, 110)
        thumb_layout.addWidget(img_label)
        

        filename = os.path.basename(img_path)
        if len(filename) > 15:
            filename = filename[:12] + "..."
        name_label = QLabel(filename)
        name_label.setAlignment(Qt.AlignCenter)
        name_label.setStyleSheet("font-size: 9px;")
        name_label.setWordWrap(True)
        thumb_layout.addWidget(name_label)
        

        remove_btn = QPushButton("✕")
        remove_btn.setFixedSize(30, 25)
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        remove_btn.clicked.connect(lambda: self.remove_image(img_path, thumb_widget))
        thumb_layout.addWidget(remove_btn, alignment=Qt.AlignCenter)
        
        thumb_widget.setLayout(thumb_layout)
        thumb_widget.setFixedWidth(120)
        
        self.thumbnail_widgets.append((img_path, thumb_widget))
        self.thumbnails_layout.addWidget(thumb_widget)
    
    def remove_image(self, img_path, thumb_widget):
        """Remove a specific image"""
        if img_path in self.image_paths:
            self.image_paths.remove(img_path)
        

        self.thumbnail_widgets = [(path, widget) for path, widget in self.thumbnail_widgets if path != img_path]
        thumb_widget.setParent(None)
        thumb_widget.deleteLater()
        
        self.update_display()
    
    def clear_images(self):
        """Clear all images"""
        self.image_paths = []
        

        for path, widget in self.thumbnail_widgets:
            widget.setParent(None)
            widget.deleteLater()
        
        self.thumbnail_widgets = []
        self.update_display()
    
    def update_display(self):
        """Update the display based on whether images exist"""
        if self.image_paths:
            self.no_images_label.hide()
        else:
            self.no_images_label.show()
    
    def dragEnterEvent(self, event):
        """Accept drag events with image files"""
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                file_path = url.toLocalFile()
                if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
                    event.acceptProposedAction()
                    return
        event.ignore()
    
    def dropEvent(self, event):
        """Handle dropped image files"""
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                file_path = url.toLocalFile()
                if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
                    if file_path not in self.image_paths and os.path.exists(file_path):
                        self.image_paths.append(file_path)
                        self.add_thumbnail(file_path)
            self.update_display()
            event.acceptProposedAction()
        else:
            event.ignore()

class ComponentSection(QWidget):
    def __init__(self, title, wear_patterns_dict):
        super().__init__()
        self.title = title
        self.wear_patterns = wear_patterns_dict
        self.checkboxes = {}
        self.pattern_image_widgets = {}
        self.observation_checkboxes = {}
        self.observation_containers = {}
        self.custom_patterns = []
        self.custom_pattern_widgets = {}
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        

        title_label = QLabel(self.title)
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        

        # To add new equipment: update this condition for custom labels
        if self.title == "Visual Inspection":
            wear_group = QGroupBox("Select Observed Issues:")
        else:
            wear_group = QGroupBox("Select Observed Wear Patterns:")
        self.wear_layout = QVBoxLayout()
        
        for pattern_name, pattern_data in self.wear_patterns.items():

            pattern_container = QWidget()
            pattern_layout = QVBoxLayout()
            pattern_layout.setContentsMargins(0, 5, 0, 5)
            

            cb = QCheckBox(pattern_name)
            cb.setStyleSheet("font-weight: bold;")
            
            if self.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                cb_font = cb.font()
                cb_font.setPointSize(9)
                cb.setFont(cb_font)
            # 10pt font for status checkboxes
            elif pattern_name in ["Status: OK", "Status: NOT OK"]:
                cb_font = cb.font()
                cb_font.setPointSize(10)
                cb.setFont(cb_font)
            

            if self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]:
                if pattern_name == "Status: OK":
                    cb.stateChanged.connect(lambda state, pn=pattern_name: self.toggle_status_checkbox(pn, state))
                elif pattern_name == "Status: NOT OK":
                    cb.stateChanged.connect(lambda state, pn=pattern_name: self.toggle_status_checkbox(pn, state))
                else:
                    cb.stateChanged.connect(lambda state, pn=pattern_name: self.toggle_pattern_section(pn, state))
            else:
                cb.stateChanged.connect(lambda state, pn=pattern_name: self.toggle_pattern_section(pn, state))
            
            self.checkboxes[pattern_name] = cb
            pattern_layout.addWidget(cb)
            

            sub_container = QWidget()
            sub_layout = QVBoxLayout()
            sub_layout.setContentsMargins(30, 5, 0, 5)
            

            has_direct_sub_obs = "sub_observations" in pattern_data and "" in pattern_data["sub_observations"]
            

            is_status_ok = (len(pattern_data.get("observations", [])) == 1 and 
                          "sub_observations" not in pattern_data)
            
            if is_status_ok:

                pass
            elif has_direct_sub_obs:

                sub_obs_group = QGroupBox("Specific Issues:")
                sub_obs_group_layout = QVBoxLayout()
                
                self.observation_checkboxes[pattern_name] = []
                

                is_special_section = self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]
                

                if is_special_section and pattern_name not in self.pattern_image_widgets:
                    self.pattern_image_widgets[pattern_name] = {}
                
                for sub_obs in pattern_data["sub_observations"][""]:
                    if is_special_section:

                        obs_container = QWidget()
                        obs_layout = QVBoxLayout()
                        obs_layout.setContentsMargins(0, 2, 0, 2)
                        
                        checkbox_row = QHBoxLayout()
                        
                        sub_cb = QCheckBox(sub_obs)
                        
                        if self.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                            sub_cb_font = sub_cb.font()
                            sub_cb_font.setPointSize(9)
                            sub_cb.setFont(sub_cb_font)
                        
                        checkbox_row.addWidget(sub_cb)
                        
                        toggle_btn = QPushButton("📷 Upload Images")
                        toggle_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px 8px; font-size: 10px;")
                        toggle_btn.setMaximumWidth(120)
                        toggle_btn.setVisible(False)
                        checkbox_row.addWidget(toggle_btn)
                        checkbox_row.addStretch()
                        
                        obs_layout.addLayout(checkbox_row)
                        

                        obs_img_widget = ImageUploadWidget(f"{self.title} - {sub_obs[:30]}...")
                        obs_img_widget.setVisible(False)
                        obs_layout.addWidget(obs_img_widget)
                        

                        self.pattern_image_widgets[pattern_name][sub_obs] = obs_img_widget
                        
                        sub_cb.stateChanged.connect(
                            lambda state, btn=toggle_btn: btn.setVisible(state == 2)
                        )
                        
                        toggle_btn.clicked.connect(
                            lambda checked=False, img_w=obs_img_widget, btn=toggle_btn: (
                                img_w.setVisible(not img_w.isVisible()),
                                btn.setText("📷 Hide Images" if img_w.isVisible() else "📷 Upload Images")
                            )
                        )
                        
                        obs_container.setLayout(obs_layout)
                        sub_obs_group_layout.addWidget(obs_container)
                    else:

                        sub_cb = QCheckBox(sub_obs)
                        
                        if self.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                            sub_cb_font = sub_cb.font()
                            sub_cb_font.setPointSize(9)
                            sub_cb.setFont(sub_cb_font)
                        
                        sub_obs_group_layout.addWidget(sub_cb)
                    
                    self.observation_checkboxes[pattern_name].append(sub_cb)
                
                sub_obs_group.setLayout(sub_obs_group_layout)
                sub_layout.addWidget(sub_obs_group)
                
                                # Special sections get inline custom issues with per-issue images
                if self.title in ["Motor Shaft", "Electrical Connection", "Motor Housing", "Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"] and pattern_name == "Status: NOT OK":

                    self.custom_section_container = QWidget()
                    self.custom_section_layout = QVBoxLayout()
                    self.custom_section_layout.setContentsMargins(0, 10, 0, 5)
                    self.custom_section_container.setLayout(self.custom_section_layout)
                    


                    self.first_custom_btn = QPushButton("+ Add Custom Issue")
                    self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
                    self.first_custom_btn.clicked.connect(self.add_custom_pattern)
                    self.custom_section_layout.addWidget(self.first_custom_btn)
                    

                    sub_layout.addWidget(self.custom_section_container)
                    

                    img_widget = ImageUploadWidget(f"{self.title} - {pattern_name}")
                    



                    is_motor_section = self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]
                    
                    if is_motor_section and isinstance(self.pattern_image_widgets.get(pattern_name), dict):

                        self.pattern_image_widgets[pattern_name]['__pattern_level__'] = img_widget
                    else:

                        self.pattern_image_widgets[pattern_name] = img_widget
                    
                    sub_layout.addWidget(img_widget)
                    

                    image_widget_added = True
                else:
                    image_widget_added = False
                

                if not image_widget_added and pattern_name not in self.pattern_image_widgets:
                    img_widget = ImageUploadWidget(f"{self.title} - {pattern_name}")
                    self.pattern_image_widgets[pattern_name] = img_widget
                    sub_layout.addWidget(img_widget)
            else:

                obs_group = QGroupBox("Select Specific Observations:")
                obs_layout = QVBoxLayout()
                
                self.observation_checkboxes[pattern_name] = []
                

                has_direct_sub_obs_check = "sub_observations" in pattern_data and "" in pattern_data["sub_observations"]
                
                if has_direct_sub_obs_check and len(pattern_data["observations"]) == 0:
                # Nested checkbox structure

                    pass
                else:

                    for observation in pattern_data["observations"]:

                        if not observation:
                            continue
                            
                        obs_cb = QCheckBox(observation)
                        
                        if self.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                            obs_cb_font = obs_cb.font()
                            obs_cb_font.setPointSize(9)
                            obs_cb.setFont(obs_cb_font)
                        
                        self.observation_checkboxes[pattern_name].append(obs_cb)
                        obs_layout.addWidget(obs_cb)
                        

                        if "sub_observations" in pattern_data and observation in pattern_data["sub_observations"]:

                            sub_obs_container = QWidget()
                            sub_obs_layout = QVBoxLayout()
                            sub_obs_layout.setContentsMargins(30, 5, 0, 5)
                            
                            sub_obs_group = QGroupBox("Specific Issues:")
                            sub_obs_group_layout = QVBoxLayout()
                            
                            sub_obs_checkboxes = []
                            for sub_obs in pattern_data["sub_observations"][observation]:
                                sub_cb = QCheckBox(sub_obs)
                                
                                if self.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                                    sub_cb_font = sub_cb.font()
                                    sub_cb_font.setPointSize(9)
                                    sub_cb.setFont(sub_cb_font)
                                
                                sub_obs_checkboxes.append(sub_cb)
                                sub_obs_group_layout.addWidget(sub_cb)
                            
                            sub_obs_group.setLayout(sub_obs_group_layout)
                            sub_obs_layout.addWidget(sub_obs_group)
                            sub_obs_container.setLayout(sub_obs_layout)
                            sub_obs_container.setVisible(False)
                            

                            obs_cb.stateChanged.connect(
                                lambda state, container=sub_obs_container: container.setVisible(state == 2)
                            )
                            

                            for sub_cb in sub_obs_checkboxes:
                                self.observation_checkboxes[pattern_name].append(sub_cb)
                            
                            obs_layout.addWidget(sub_obs_container)
                
                obs_group.setLayout(obs_layout)
                sub_layout.addWidget(obs_group)
            

            if not is_status_ok:

                if pattern_name not in self.pattern_image_widgets:

                    img_widget = ImageUploadWidget(f"{self.title} - {pattern_name}")
                    self.pattern_image_widgets[pattern_name] = img_widget
                    sub_layout.addWidget(img_widget)
                
                sub_container.setLayout(sub_layout)
                sub_container.setVisible(False)
                self.observation_containers[pattern_name] = sub_container
                
                pattern_layout.addWidget(sub_container)
            else:

                img_widget = ImageUploadWidget(f"{self.title} - {pattern_name}")
                self.pattern_image_widgets[pattern_name] = img_widget
                
                img_container = QWidget()
                img_layout = QVBoxLayout()
                img_layout.setContentsMargins(30, 5, 0, 5)
                img_layout.addWidget(img_widget)
                img_container.setLayout(img_layout)
                img_container.setVisible(False)
                self.observation_containers[pattern_name] = img_container
                
                pattern_layout.addWidget(img_container)
            pattern_container.setLayout(pattern_layout)
            self.wear_layout.addWidget(pattern_container)
        
        # Bearing sections get custom wear patterns; special sections use inline custom issues

        if self.title not in ["Motor Housing", "Motor Shaft", "Electrical Connection", "Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]:
            custom_section_container = QWidget()
            custom_section_layout = QVBoxLayout()
            custom_section_layout.setContentsMargins(0, 10, 0, 5)
            custom_section_container.setLayout(custom_section_layout)
            self.wear_layout.addWidget(custom_section_container)
            

            self.custom_section_container = custom_section_container
            self.custom_section_layout = custom_section_layout
            

            self.first_custom_btn = QPushButton("+ Add Custom Wear Pattern")
            self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.first_custom_btn.clicked.connect(self.add_custom_pattern)
            custom_section_layout.addWidget(self.first_custom_btn)
        
        wear_group.setLayout(self.wear_layout)
        layout.addWidget(wear_group)
        

        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Enter any additional observations or notes here...")
        self.notes_text.setMaximumHeight(100)
        layout.addWidget(self.notes_text)
        

        general_label = QLabel("General Section Images:")
        general_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(general_label)
        
        self.image_widget = ImageUploadWidget(self.title)
        layout.addWidget(self.image_widget)
        
        self.setLayout(layout)
    
    def add_custom_pattern(self):
        """Add a new custom wear pattern entry"""

        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)
        
        custom_id = len(self.custom_patterns)
        

        custom_container = QWidget()
        custom_layout = QVBoxLayout()
        custom_layout.setContentsMargins(0, 5, 0, 5)
        

        is_special = self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection"]
        
        if is_special:

            inline_layout = QHBoxLayout()
            
            text_input = QLineEdit()
            text_input.setPlaceholderText("Enter custom issue description...")
            inline_layout.addWidget(text_input)
            
            # Add per-observation image upload for special sections (motor/valve)
            is_motor_section = self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]
            if is_motor_section:
                toggle_img_btn = QPushButton("📷 Upload Images")
                toggle_img_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px 8px; font-size: 10px;")
                toggle_img_btn.setMaximumWidth(120)
                inline_layout.addWidget(toggle_img_btn)
            
            delete_btn = QPushButton("Delete")
            delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 5px;")
            delete_btn.clicked.connect(lambda: self.delete_custom_pattern(custom_id, custom_container))
            inline_layout.addWidget(delete_btn)
            
            custom_layout.addLayout(inline_layout)
            
            # Toggled by button
            if is_motor_section:
                img_widget = ImageUploadWidget(f"{self.title} - Custom Issue {custom_id + 1}")
                img_widget.setVisible(False)
                custom_layout.addWidget(img_widget)
                
                toggle_img_btn.clicked.connect(
                    lambda checked=False, img_w=img_widget, btn=toggle_img_btn: (
                        img_w.setVisible(not img_w.isVisible()),
                        btn.setText("📷 Hide Images" if img_w.isVisible() else "📷 Upload Images")
                    )
                )
            else:
                img_widget = None

            
            custom_container.setLayout(custom_layout)
            

            custom_data = {
                'container': custom_container,
                'text_input': text_input,
                'image_widget': img_widget,
                'id': custom_id
            }
            self.custom_patterns.append(custom_data)
            

            if not hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn = QPushButton("+ Add Another Custom Issue")
                self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
                self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)

                self.custom_section_layout.addWidget(self.shared_add_another_btn)
            else:

                self.shared_add_another_btn.setVisible(True)
            

            button_position = self.custom_section_layout.indexOf(self.shared_add_another_btn)
            if button_position >= 0:
                self.custom_section_layout.insertWidget(button_position, custom_container)
            else:

                self.custom_section_layout.insertWidget(self.custom_section_layout.count(), custom_container)
        else:


            name_layout = QHBoxLayout()
            name_label = QLabel("Wear Pattern Name:")
            name_input = QLineEdit()
            name_input.setPlaceholderText("Enter custom wear pattern name...")
            name_layout.addWidget(name_label)
            name_layout.addWidget(name_input)
            

            delete_btn = QPushButton("Delete")
            delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 5px;")
            delete_btn.clicked.connect(lambda: self.delete_custom_pattern(custom_id, custom_container))
            name_layout.addWidget(delete_btn)
            
            custom_layout.addLayout(name_layout)
            

            obs_container = QWidget()
            obs_layout = QVBoxLayout()
            obs_layout.setContentsMargins(30, 5, 0, 5)
            
            obs_group = QGroupBox("Custom Observations:")
            obs_group_layout = QVBoxLayout()
            

            obs_inputs_container = QWidget()
            obs_inputs_layout = QVBoxLayout()
            obs_inputs_container.setLayout(obs_inputs_layout)
            obs_group_layout.addWidget(obs_inputs_container)
            
            obs_group.setLayout(obs_group_layout)
            obs_container.setLayout(obs_layout)
            obs_layout.addWidget(obs_group)
            
            img_widget = ImageUploadWidget(f"{self.title} - Custom Pattern {custom_id + 1}")
            obs_layout.addWidget(img_widget)
            
            custom_layout.addWidget(obs_container)
            

            
            custom_container.setLayout(custom_layout)
            

            custom_data = {
                'container': custom_container,
                'name_input': name_input,
                'obs_inputs_layout': obs_inputs_layout,
                'obs_inputs': [],
                'image_widget': img_widget,
                'obs_count': 1,
                'id': custom_id
            }
            self.custom_patterns.append(custom_data)
            

            self.add_custom_observation_input(obs_inputs_layout, custom_id, 0)
            

            if not hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn = QPushButton("+ Add Another Custom Wear Pattern")
                self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
                self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)

                self.custom_section_layout.addWidget(self.shared_add_another_btn)
            else:

                self.shared_add_another_btn.setVisible(True)
            

            button_position = self.custom_section_layout.indexOf(self.shared_add_another_btn)
            if button_position >= 0:
                self.custom_section_layout.insertWidget(button_position, custom_container)
            else:

                self.custom_section_layout.insertWidget(self.custom_section_layout.count(), custom_container)
    
    def delete_custom_pattern(self, pattern_id, container):
        """Delete a custom wear pattern"""

        container.setParent(None)
        container.deleteLater()
        

        self.custom_patterns = [cp for cp in self.custom_patterns if cp.get('id') != pattern_id]
        

        if len(self.custom_patterns) == 0:
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)
            if hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn.setVisible(False)
    
    def add_custom_observation_input(self, parent_layout, pattern_id, obs_id):
        """Add a custom observation text input with an 'Add another' button"""
        obs_widget = QWidget()
        obs_layout = QHBoxLayout()
        obs_layout.setContentsMargins(0, 2, 0, 2)
        
        obs_input = QLineEdit()
        obs_input.setPlaceholderText(f"Enter observation {obs_id + 1}...")
        obs_layout.addWidget(obs_input)
        
        add_obs_btn = QPushButton("+ Add Observation")
        add_obs_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px;")
        add_obs_btn.clicked.connect(
            lambda: self.add_another_observation(pattern_id, parent_layout)
        )
        obs_layout.addWidget(add_obs_btn)
        
        obs_widget.setLayout(obs_layout)
        parent_layout.addWidget(obs_widget)
        

        if pattern_id < len(self.custom_patterns):
            self.custom_patterns[pattern_id]['obs_inputs'].append(obs_input)
    
    def add_another_observation(self, pattern_id, parent_layout):
        """Add another observation input to an existing custom pattern"""
        if pattern_id < len(self.custom_patterns):
            obs_count = self.custom_patterns[pattern_id]['obs_count']
            self.add_custom_observation_input(parent_layout, pattern_id, obs_count)
            self.custom_patterns[pattern_id]['obs_count'] += 1
    
    def toggle_status_checkbox(self, pattern_name, state):
        """Handle Status: OK / Status: NOT OK mutual exclusivity and visibility"""
        if pattern_name in self.observation_containers:
            self.observation_containers[pattern_name].setVisible(state == 2)
        

        if state == 2:
            if pattern_name == "Status: OK" and "Status: NOT OK" in self.checkboxes:
                self.checkboxes["Status: NOT OK"].setChecked(False)
            elif pattern_name == "Status: NOT OK" and "Status: OK" in self.checkboxes:
                self.checkboxes["Status: OK"].setChecked(False)
    
    def toggle_pattern_section(self, pattern_name, state):
        """Show/hide observation checkboxes and image upload when main checkbox is toggled"""
        if pattern_name in self.observation_containers:
            self.observation_containers[pattern_name].setVisible(state == 2)
    
    def get_selected_patterns(self):
        selected = []
        for pattern_name, checkbox in self.checkboxes.items():
            if checkbox.isChecked():
                selected.append(pattern_name)
        return selected
    
    def get_custom_patterns(self):
        """Get all custom patterns with their data"""
        custom_list = []
        is_special = self.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection"]
        
        for custom_data in self.custom_patterns:
            if is_special:
                text = custom_data.get('text_input')
                if text:

                    if hasattr(text, 'toPlainText'):
                        text_content = text.toPlainText().strip()
                    else:
                        text_content = text.text().strip()
                    
                    if text_content:
                        images = []
                        img_widget = custom_data.get('image_widget')
                        if img_widget and hasattr(img_widget, 'image_paths'):
                            images = img_widget.image_paths
                        
                        custom_list.append({
                            'text': text_content,
                            'images': images
                        })
            else:
                name_input = custom_data.get('name_input')
                if name_input:
                    pattern_name = name_input.text().strip()
                    if pattern_name:
                        observations = []
                        obs_inputs = custom_data.get('obs_inputs', [])
                        for obs_input in obs_inputs:
                            obs_text = obs_input.text().strip()
                            if obs_text:
                                observations.append(obs_text)
                        
                        if observations:
                            img_widget = custom_data.get('image_widget')
                            custom_list.append({
                                'name': pattern_name,
                                'observations': observations,
                                'images': img_widget.image_paths if img_widget else []
                            })
        return custom_list
    
    def get_selected_observations(self, pattern_name):
        """Get selected observation sentences for a specific wear pattern"""

        if pattern_name in self.wear_patterns:
            pattern_data = self.wear_patterns[pattern_name]
            is_status_ok = (len(pattern_data.get("observations", [])) == 1 and 
                          "sub_observations" not in pattern_data)
            
            if is_status_ok:
                return pattern_data["observations"]
        
        if pattern_name not in self.observation_checkboxes:
            return []
        
        selected = []
        for obs_cb in self.observation_checkboxes[pattern_name]:
            if obs_cb.isChecked():
                selected.append(obs_cb.text())
        
        return selected
    
    def add_first_custom_btn(self):
        """Add the initial 'Add Custom Wear Pattern' button"""

        if not hasattr(self, 'custom_section_layout') or self.custom_section_layout is None:
            return
            
        self.first_custom_btn = QPushButton("+ Add Custom Wear Pattern")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)
    
    def get_observation_images(self, pattern_name, observation_text):
        """Get images for a specific observation within a pattern"""
        if pattern_name in self.pattern_image_widgets:
            widget_data = self.pattern_image_widgets[pattern_name]

            if isinstance(widget_data, dict):
                if observation_text in widget_data:
                    images = widget_data[observation_text].image_paths
                    print(f"Found images for '{observation_text}': {len(images)} images")
                    return images

                for key, widget in widget_data.items():
                    if observation_text.strip() == key.strip():
                        images = widget.image_paths
                        print(f"Found images (stripped match) for '{observation_text}': {len(images)} images")
                        return images
                print(f"No images found for observation: '{observation_text}'")
                print(f"Available keys: {list(widget_data.keys())}")
        return []
    
    def get_pattern_images(self, pattern_name):
        """Get images for a specific wear pattern"""
        if pattern_name in self.pattern_image_widgets:
            widget = self.pattern_image_widgets[pattern_name]

            if isinstance(widget, dict):

                if '__pattern_level__' in widget:
                    return widget['__pattern_level__'].image_paths
                return []
            else:

                return widget.image_paths
        return []
    
    def get_notes(self):
        return self.notes_text.toPlainText()
    
    def get_images(self):
        return self.image_widget.image_paths

class TestResultsSection(QWidget):
    """Widget for test results including Audio, Vibration, Temperature, and Electrical Resistance"""
    def __init__(self):
        super().__init__()
        self.title = "Test Results"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        

        title_label = QLabel("Test Results")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        

        audio_group = QGroupBox("Audio Testing")
        audio_layout = QVBoxLayout()
        
        audio_info = QLabel("Limit: 72 dB (values over 72 are NOT OK)")
        audio_info.setStyleSheet("color: #666; font-style: italic;")
        audio_layout.addWidget(audio_info)
        
        audio_value_layout = QHBoxLayout()
        audio_value_layout.addWidget(QLabel("Result Value (dB):"))
        self.audio_value = QLineEdit()
        self.audio_value.setPlaceholderText("Enter audio test result in dB...")
        audio_value_layout.addWidget(self.audio_value)
        audio_layout.addLayout(audio_value_layout)
        
        self.audio_images = ImageUploadWidget("Audio Testing")
        audio_layout.addWidget(self.audio_images)
        
        audio_group.setLayout(audio_layout)
        layout.addWidget(audio_group)
        

        vibration_group = QGroupBox("Vibration Testing")
        vibration_layout = QVBoxLayout()
        
        vibration_info = QLabel("Enter RPM and vibration value to determine status per ISO 10816")
        vibration_info.setStyleSheet("color: #666; font-style: italic;")
        vibration_layout.addWidget(vibration_info)
        
        vibration_rpm_layout = QHBoxLayout()
        vibration_rpm_layout.addWidget(QLabel("RPM (r/min):"))
        self.vibration_rpm = QLineEdit()
        self.vibration_rpm.setPlaceholderText("Enter RPM...")
        vibration_rpm_layout.addWidget(self.vibration_rpm)
        vibration_layout.addLayout(vibration_rpm_layout)
        
        vibration_value_layout = QHBoxLayout()
        vibration_value_layout.addWidget(QLabel("Vibration Value (mm/s):"))
        self.vibration_value = QLineEdit()
        self.vibration_value.setPlaceholderText("Enter vibration value in mm/s...")
        vibration_value_layout.addWidget(self.vibration_value)
        vibration_layout.addLayout(vibration_value_layout)
        
        self.vibration_images = ImageUploadWidget("Vibration Testing")
        vibration_layout.addWidget(self.vibration_images)
        
        vibration_group.setLayout(vibration_layout)
        layout.addWidget(vibration_group)
        

        temperature_group = QGroupBox("Temperature Testing")
        temperature_layout = QVBoxLayout()
        
        temperature_info = QLabel("Limit: 95°F (values over 95 are NOT OK)")
        temperature_info.setStyleSheet("color: #666; font-style: italic;")
        temperature_layout.addWidget(temperature_info)
        
        temperature_value_layout = QHBoxLayout()
        temperature_value_layout.addWidget(QLabel("Result Value (°F):"))
        self.temperature_value = QLineEdit()
        self.temperature_value.setPlaceholderText("Enter temperature test result in °F...")
        temperature_value_layout.addWidget(self.temperature_value)
        temperature_layout.addLayout(temperature_value_layout)
        
        self.temperature_images = ImageUploadWidget("Temperature Testing")
        temperature_layout.addWidget(self.temperature_images)
        
        temperature_group.setLayout(temperature_layout)
        layout.addWidget(temperature_group)
        
        resistance_group = QGroupBox("Electrical Resistance Testing")
        resistance_layout = QVBoxLayout()

        resistance_status_layout = QHBoxLayout()
        self.resistance_ok_cb = QCheckBox("Status: OK")
        resistance_status_layout.addWidget(self.resistance_ok_cb)
        resistance_status_layout.addStretch()
        resistance_layout.addLayout(resistance_status_layout)

        resistance_value_layout = QGridLayout()
        resistance_value_layout.addWidget(QLabel("Black / Red:"), 0, 0)
        self.res_black_red = QLineEdit()
        self.res_black_red.setPlaceholderText("Ω")
        resistance_value_layout.addWidget(self.res_black_red, 0, 1)

        resistance_value_layout.addWidget(QLabel("Black / White:"), 1, 0)
        self.res_black_white = QLineEdit()
        self.res_black_white.setPlaceholderText("Ω")
        resistance_value_layout.addWidget(self.res_black_white, 1, 1)

        resistance_layout.addLayout(resistance_value_layout)
        
        self.five_wire_cb = QCheckBox("5 wire motor")
        self.five_wire_cb.setStyleSheet("font-weight: bold; color: #2196F3; font-size: 10pt;")
        self.five_wire_cb.stateChanged.connect(self.toggle_five_wire_fields)
        resistance_layout.addWidget(self.five_wire_cb)
        
        # Hidden until checkbox enables them
        self.five_wire_layout = QGridLayout()
        
        self.black_yellow_label = QLabel("Black / Yellow:")
        self.five_wire_layout.addWidget(self.black_yellow_label, 0, 0)
        self.res_black_yellow = QLineEdit()
        self.res_black_yellow.setPlaceholderText("Ω")
        self.five_wire_layout.addWidget(self.res_black_yellow, 0, 1)
        
        self.black_blue_label = QLabel("Black / Blue:")
        self.five_wire_layout.addWidget(self.black_blue_label, 1, 0)
        self.res_black_blue = QLineEdit()
        self.res_black_blue.setPlaceholderText("Ω")
        self.five_wire_layout.addWidget(self.res_black_blue, 1, 1)
        
        resistance_layout.addLayout(self.five_wire_layout)
        
        self.black_yellow_label.hide()
        self.res_black_yellow.hide()
        self.black_blue_label.hide()
        self.res_black_blue.hide()

        self.resistance_images = ImageUploadWidget("Electrical Resistance Testing")
        resistance_layout.addWidget(self.resistance_images)

        resistance_group.setLayout(resistance_layout)
        layout.addWidget(resistance_group)

        layout.addStretch()
        self.setLayout(layout)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def extract_numeric_value(self, text):
        """Extract numeric value from text, ignoring characters"""
        import re
        if not text:
            return None
        match = re.search(r'[-+]?\d*\.?\d+', text.strip())
        if match:
            try:
                return float(match.group())
            except ValueError:
                return None
        return None
    
    def get_vibration_classification(self, rpm, vibration_value):
        """Classify vibration based on General Machinery Vibration Severity Chart"""
            # Vibration classification using ISO 10816 boundary thresholds
        
        if rpm is None or vibration_value is None:
            return "Unable to classify", False
        
        import math
        
        try:
            # Calibrated using ISO 10816 standards at reference point: 1400 RPM, 1.1 mm/s = Fair
            
            
            
            rpm_factor = math.pow(rpm / 1000.0, 0.5)
            
            # Threshold values in mm/s RMS
            extremely_smooth_threshold = 0.28 * rpm_factor
            very_smooth_threshold = 0.45 * rpm_factor
            smooth_threshold = 0.71 * rpm_factor
            good_threshold = 1.12 * rpm_factor
            fair_threshold = 1.8 * rpm_factor
            slightly_rough_threshold = 2.8 * rpm_factor
            rough_threshold = 4.5 * rpm_factor
            very_rough_threshold = 7.1 * rpm_factor
            
            if vibration_value <= extremely_smooth_threshold:
                return "Extremely Smooth", True
            elif vibration_value <= very_smooth_threshold:
                return "Very Smooth", True
            elif vibration_value <= smooth_threshold:
                return "Smooth", True
            elif vibration_value <= good_threshold:
                return "Good", True
            elif vibration_value <= fair_threshold:
                return "Fair", False
            elif vibration_value <= slightly_rough_threshold:
                return "Slightly Rough", False
            elif vibration_value <= rough_threshold:
                return "Rough", False
            elif vibration_value <= very_rough_threshold:
                return "Very Rough", False
            else:
                return "Extremely Rough", False
                
        except (ValueError, ZeroDivisionError):
            return "Unable to classify", False
    
    def get_test_data(self):
        """Return all test result data with automatic OK/NOT OK determination"""

        audio_value_raw = self.audio_value.text().strip()
        audio_numeric = self.extract_numeric_value(audio_value_raw)
        audio_ok = audio_numeric is not None and audio_numeric <= 72
        

        temp_value_raw = self.temperature_value.text().strip()
        temp_numeric = self.extract_numeric_value(temp_value_raw)
        temp_ok = temp_numeric is not None and temp_numeric <= 95
        

        vibration_rpm_raw = self.vibration_rpm.text().strip()
        vibration_value_raw = self.vibration_value.text().strip()
        vibration_rpm_numeric = self.extract_numeric_value(vibration_rpm_raw)
        vibration_value_numeric = self.extract_numeric_value(vibration_value_raw)
        
        vibration_classification = "Not measured"
        vibration_ok = None
        if vibration_rpm_numeric and vibration_value_numeric:
            vibration_classification, vibration_ok = self.get_vibration_classification(
                vibration_rpm_numeric, vibration_value_numeric
            )
            if vibration_ok is None:
                vibration_ok = True
        
        return {
            'audio': {
                'ok': audio_ok,
                'value': audio_value_raw,
                'numeric': audio_numeric,
                'limit': 72,
                'images': self.audio_images.image_paths
            },
            'vibration': {
                'ok': vibration_ok,
                'rpm': vibration_rpm_raw,
                'value': vibration_value_raw,
                'rpm_numeric': vibration_rpm_numeric,
                'numeric': vibration_value_numeric,
                'classification': vibration_classification,
                'images': self.vibration_images.image_paths
            },
            'temperature': {
                'ok': temp_ok,
                'value': temp_value_raw,
                'numeric': temp_numeric,
                'limit': 95,
                'images': self.temperature_images.image_paths
            },
            'resistance': {
                'ok': self.resistance_ok_cb.isChecked(),
                'black_red': self.res_black_red.text().strip(),
                'black_white': self.res_black_white.text().strip(),
                'black_yellow': self.res_black_yellow.text().strip() if self.five_wire_cb.isChecked() else '',
                'black_blue': self.res_black_blue.text().strip() if self.five_wire_cb.isChecked() else '',
                'is_five_wire': self.five_wire_cb.isChecked(),
                'images': self.resistance_images.image_paths
            }
        }
    
    def toggle_five_wire_fields(self, state):
        """Show/hide 5-wire motor fields"""
        if state:
            self.black_yellow_label.show()
            self.res_black_yellow.show()
            self.black_blue_label.show()
            self.res_black_blue.show()
        else:
            self.black_yellow_label.hide()
            self.res_black_yellow.hide()
            self.black_blue_label.hide()
            self.res_black_blue.hide()

class MotorTab(QWidget):
    """Widget for a single motor containing housing and bearing sections"""
    def __init__(self, motor_number):
        super().__init__()
        self.motor_number = motor_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)
        
        housing_shaft_widget = QWidget()
        housing_shaft_layout = QVBoxLayout()
        
        self.housing_shaft_subtabs = QTabWidget()
        subtab_font = self.housing_shaft_subtabs.font()
        subtab_font.setPointSize(subtab_font.pointSize() + 1)
        self.housing_shaft_subtabs.setFont(subtab_font)
        
        scroll_housing = QScrollArea()
        scroll_housing.setWidgetResizable(True)
        self.housing_section = ComponentSection("Motor Housing", HOUSING_WEAR_PATTERNS)
        scroll_housing.setWidget(self.housing_section)
        self.housing_shaft_subtabs.addTab(scroll_housing, "Motor Housing")
        
        scroll_shaft = QScrollArea()
        scroll_shaft.setWidgetResizable(True)
        self.shaft_section = ComponentSection("Motor Shaft", MOTOR_SHAFT_PATTERNS)
        scroll_shaft.setWidget(self.shaft_section)
        self.housing_shaft_subtabs.addTab(scroll_shaft, "Motor Shaft")
        
        scroll_electrical = QScrollArea()
        scroll_electrical.setWidgetResizable(True)
        self.electrical_section = ComponentSection("Electrical Connection", ELECTRICAL_CONNECTION_PATTERNS)
        scroll_electrical.setWidget(self.electrical_section)
        self.housing_shaft_subtabs.addTab(scroll_electrical, "Electrical Connection")
        
        housing_shaft_layout.addWidget(self.housing_shaft_subtabs)
        housing_shaft_widget.setLayout(housing_shaft_layout)
        self.component_tabs.addTab(housing_shaft_widget, "Motor Housing / Shaft")
        
        scroll_tests = QScrollArea()
        scroll_tests.setWidgetResizable(True)
        self.test_results_section = TestResultsSection()
        scroll_tests.setWidget(self.test_results_section)
        self.component_tabs.addTab(scroll_tests, "Test Results")
        
        scroll_de = QScrollArea()
        scroll_de.setWidgetResizable(True)
        self.bearing_de_section = ComponentSection("Drive End Bearing", BEARING_WEAR_PATTERNS)
        scroll_de.setWidget(self.bearing_de_section)
        self.component_tabs.addTab(scroll_de, "Drive End Bearing")
        
        scroll_nde = QScrollArea()
        scroll_nde.setWidgetResizable(True)
        self.bearing_nde_section = ComponentSection("Non-Drive End Bearing", BEARING_WEAR_PATTERNS)
        scroll_nde.setWidget(self.bearing_nde_section)
        self.component_tabs.addTab(scroll_nde, "Non-Drive End Bearing")
        
        layout.addWidget(self.component_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all component sections in order"""
        return [
            (self.housing_section, HOUSING_WEAR_PATTERNS, 'component'),
            (self.shaft_section, MOTOR_SHAFT_PATTERNS, 'component'),
            (self.electrical_section, ELECTRICAL_CONNECTION_PATTERNS, 'component'),
            (self.test_results_section, None, 'test'),
            (self.bearing_de_section, BEARING_WEAR_PATTERNS, 'component'),
            (self.bearing_nde_section, BEARING_WEAR_PATTERNS, 'component')
        ]

class CompressorElectricalTestingSection(QWidget):
    """Widget for compressor electrical testing with resistance value entry"""
    def __init__(self):
        super().__init__()
        self.title = "Electrical Testing"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        

        title_label = QLabel("Electrical Testing")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        

        status_group = QGroupBox("Select Status:")
        status_layout = QVBoxLayout()
        
        ok_container = QWidget()
        ok_layout = QVBoxLayout()
        ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_ok_cb = QCheckBox("Status: OK")
        self.status_ok_cb.setStyleSheet("font-weight: bold;")
        
        ok_font = self.status_ok_cb.font()
        ok_font.setPointSize(10)
        self.status_ok_cb.setFont(ok_font)
        
        self.status_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, True))
        ok_layout.addWidget(self.status_ok_cb)
        
        self.ok_image_container = QWidget()
        ok_image_layout = QVBoxLayout()
        ok_image_layout.setContentsMargins(30, 5, 0, 5)
        self.ok_image_widget = ImageUploadWidget("Status: OK")
        ok_image_layout.addWidget(self.ok_image_widget)
        self.ok_image_container.setLayout(ok_image_layout)
        self.ok_image_container.setVisible(False)
        ok_layout.addWidget(self.ok_image_container)
        
        ok_container.setLayout(ok_layout)
        status_layout.addWidget(ok_container)
        
        not_ok_container = QWidget()
        not_ok_layout = QVBoxLayout()
        not_ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_not_ok_cb = QCheckBox("Status: NOT OK")
        self.status_not_ok_cb.setStyleSheet("font-weight: bold;")
        
        not_ok_font = self.status_not_ok_cb.font()
        not_ok_font.setPointSize(10)
        self.status_not_ok_cb.setFont(not_ok_font)
        
        self.status_not_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, False))
        not_ok_layout.addWidget(self.status_not_ok_cb)
        
        not_ok_container.setLayout(not_ok_layout)
        status_layout.addWidget(not_ok_container)
        
        status_group.setLayout(status_layout)
        layout.addWidget(status_group)
        
        self.issues_container = QWidget()
        issues_layout = QVBoxLayout()
        issues_layout.setContentsMargins(30, 5, 0, 5)
        
        issues_group = QGroupBox("Specific Issues:")
        issues_group_layout = QVBoxLayout()
        
        self.grounded_cb = QCheckBox("Grounded windings detected.")
        self.open_cb = QCheckBox("Open windings detected.")
        self.resistance_cb = QCheckBox("Resistance values out of acceptable range.")
        
        issues_group_layout.addWidget(self.grounded_cb)
        issues_group_layout.addWidget(self.open_cb)
        issues_group_layout.addWidget(self.resistance_cb)
        
        issues_group.setLayout(issues_group_layout)
        issues_layout.addWidget(issues_group)
        

        self.custom_section_container = QWidget()
        self.custom_section_layout = QVBoxLayout()
        self.custom_section_layout.setContentsMargins(0, 10, 0, 5)
        self.custom_section_container.setLayout(self.custom_section_layout)
        

        self.first_custom_btn = QPushButton("+ Add Custom Issue")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)
        
        issues_layout.addWidget(self.custom_section_container)
        
        self.not_ok_image_widget = ImageUploadWidget("Electrical Testing - Issues")
        issues_layout.addWidget(self.not_ok_image_widget)
        
        self.issues_container.setLayout(issues_layout)
        self.issues_container.setVisible(False)
        layout.addWidget(self.issues_container)
        
        self.custom_patterns = []
        
        resistance_group = QGroupBox("Resistance Values")
        resistance_layout = QVBoxLayout()
        
        # First resistance value
        resistance1_layout = QHBoxLayout()
        resistance1_layout.addWidget(QLabel("Resistance 1 (Ω):"))
        self.resistance_value = QLineEdit()
        self.resistance_value.setPlaceholderText("Enter resistance value...")
        resistance1_layout.addWidget(self.resistance_value)
        resistance_layout.addLayout(resistance1_layout)
        
        # Second resistance value
        resistance2_layout = QHBoxLayout()
        resistance2_layout.addWidget(QLabel("Resistance 2 (Ω):"))
        self.resistance_value_2 = QLineEdit()
        self.resistance_value_2.setPlaceholderText("Enter resistance value...")
        resistance2_layout.addWidget(self.resistance_value_2)
        resistance_layout.addLayout(resistance2_layout)
        
        resistance_group.setLayout(resistance_layout)
        layout.addWidget(resistance_group)
        
        # Resistance testing images
        resistance_images_label = QLabel("Resistance Testing Images:")
        resistance_images_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(resistance_images_label)
        
        self.resistance_image_widget = ImageUploadWidget("Resistance Testing")
        layout.addWidget(self.resistance_image_widget)
        
        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Enter any additional observations...")
        self.notes_text.setMaximumHeight(100)
        layout.addWidget(self.notes_text)
        

        general_images_label = QLabel("General Section Images:")
        general_images_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(general_images_label)
        
        self.general_image_widget = ImageUploadWidget("Electrical Testing - General")
        layout.addWidget(self.general_image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_status(self, state, is_ok):
        """Handle mutual exclusivity of status checkboxes and visibility"""
        if state == 2:  # Checked
            if is_ok:
                self.status_not_ok_cb.setChecked(False)
                self.issues_container.setVisible(False)
                self.ok_image_container.setVisible(True)
            else:
                self.status_ok_cb.setChecked(False)
                self.ok_image_container.setVisible(False)
                self.issues_container.setVisible(True)
        else:  # Unchecked
            if is_ok:
                self.ok_image_container.setVisible(False)
            else:
                self.issues_container.setVisible(False)
    
    
    def add_custom_pattern(self):
        """Add a new custom issue entry"""

        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)
        
        custom_id = len(self.custom_patterns)
        

        custom_container = QWidget()
        custom_layout = QVBoxLayout()
        custom_layout.setContentsMargins(0, 5, 0, 5)
        
        inline_layout = QHBoxLayout()
        
        text_input = QLineEdit()
        text_input.setPlaceholderText("Enter custom issue description...")
        inline_layout.addWidget(text_input)
        
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 5px;")
        delete_btn.clicked.connect(lambda: self.delete_custom_pattern(custom_id, custom_container))
        inline_layout.addWidget(delete_btn)
        
        custom_layout.addLayout(inline_layout)
        

        
        custom_container.setLayout(custom_layout)
        

        custom_data = {
            'container': custom_container,
            'text_input': text_input,
            'image_widget': None,
            'id': custom_id
        }
        self.custom_patterns.append(custom_data)
        

        if not hasattr(self, 'shared_add_another_btn'):
            self.shared_add_another_btn = QPushButton("+ Add Another Custom Issue")
            self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)

            self.custom_section_layout.addWidget(self.shared_add_another_btn)
        else:

            self.shared_add_another_btn.setVisible(True)
        

        button_position = self.custom_section_layout.indexOf(self.shared_add_another_btn)
        if button_position >= 0:
            self.custom_section_layout.insertWidget(button_position, custom_container)
        else:

            self.custom_section_layout.insertWidget(self.custom_section_layout.count(), custom_container)
    
    def delete_custom_pattern(self, pattern_id, container):
        """Delete a custom wear pattern"""
        container.setParent(None)
        container.deleteLater()
        
        self.custom_patterns = [cp for cp in self.custom_patterns if cp.get('id') != pattern_id]
        

        if len(self.custom_patterns) == 0:
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)
            if hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn.setVisible(False)
    
    def get_test_data(self):
        """Return electrical testing data"""
        issues = []
        if self.grounded_cb.isChecked():
            issues.append("Grounded windings detected.")
        if self.open_cb.isChecked():
            issues.append("Open windings detected.")
        if self.resistance_cb.isChecked():
            issues.append("Resistance values out of acceptable range.")
        

        custom_list = []
        for custom_data in self.custom_patterns:
            text = custom_data['text_input'].text().strip()
            if text:
                custom_list.append({
                    'text': text,
                    'images': []
                })
        
        return {
            'ok': self.status_ok_cb.isChecked(),
            'not_ok': self.status_not_ok_cb.isChecked(),
            'ok_images': self.ok_image_widget.image_paths if hasattr(self, 'ok_image_widget') else [],
            'issues': issues,
            'custom_patterns': custom_list,
            'not_ok_images': self.not_ok_image_widget.image_paths if hasattr(self, 'not_ok_image_widget') else [],
            'resistance': self.resistance_value.text().strip(),
            'resistance_2': self.resistance_value_2.text().strip() if hasattr(self, 'resistance_value_2') else '',
            'resistance_images': self.resistance_image_widget.image_paths if hasattr(self, 'resistance_image_widget') else [],
            'general_images': self.general_image_widget.image_paths if hasattr(self, 'general_image_widget') else [],
            'notes': self.notes_text.toPlainText()
        }

class PressureTestingSection(QWidget):
    """Widget for coil pressure testing with 250psi and 500psi sections"""
    def __init__(self):
        super().__init__()
        self.title = "Pressure Testing"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        title_label = QLabel("Pressure Testing")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        psi_250_group = QGroupBox("250 PSI Test")
        psi_250_layout = QVBoxLayout()
        
        self.psi_250_leak_cb = QCheckBox("Leak")
        self.psi_250_leak_cb.stateChanged.connect(lambda state: self.toggle_250_checkbox(state))
        psi_250_layout.addWidget(self.psi_250_leak_cb)
        
        self.psi_250_leak_location_container = QWidget()
        leak_location_250_layout = QVBoxLayout()
        leak_location_250_layout.setContentsMargins(30, 0, 0, 0)
        
        self.psi_250_leak_face_cb = QCheckBox("Face")
        leak_location_250_layout.addWidget(self.psi_250_leak_face_cb)
        
        self.psi_250_leak_ubends_cb = QCheckBox("U-bends")
        leak_location_250_layout.addWidget(self.psi_250_leak_ubends_cb)
        
        self.psi_250_leak_location_container.setLayout(leak_location_250_layout)
        self.psi_250_leak_location_container.setVisible(False)
        psi_250_layout.addWidget(self.psi_250_leak_location_container)
        
        self.psi_250_no_leak_cb = QCheckBox("No Leak")
        self.psi_250_no_leak_cb.stateChanged.connect(lambda state: self.toggle_250_checkbox(state))
        psi_250_layout.addWidget(self.psi_250_no_leak_cb)
        
        self.psi_250_image_widget = ImageUploadWidget("250 PSI Test")
        self.psi_250_image_widget.setVisible(False)
        psi_250_layout.addWidget(self.psi_250_image_widget)
        
        psi_250_group.setLayout(psi_250_layout)
        layout.addWidget(psi_250_group)
        
        psi_500_group = QGroupBox("500 PSI Test")
        psi_500_layout = QVBoxLayout()
        
        self.psi_500_leak_cb = QCheckBox("Leak")
        self.psi_500_leak_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_leak_cb)
        
        self.psi_500_leak_location_container = QWidget()
        leak_location_500_layout = QVBoxLayout()
        leak_location_500_layout.setContentsMargins(30, 0, 0, 0)
        
        self.psi_500_leak_face_cb = QCheckBox("Face")
        leak_location_500_layout.addWidget(self.psi_500_leak_face_cb)
        
        self.psi_500_leak_ubends_cb = QCheckBox("U-bends")
        leak_location_500_layout.addWidget(self.psi_500_leak_ubends_cb)
        
        self.psi_500_leak_location_container.setLayout(leak_location_500_layout)
        self.psi_500_leak_location_container.setVisible(False)
        psi_500_layout.addWidget(self.psi_500_leak_location_container)
        
        self.psi_500_no_leak_cb = QCheckBox("No Leak")
        self.psi_500_no_leak_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_no_leak_cb)
        
        self.psi_500_not_performed_cb = QCheckBox("Not Performed")
        self.psi_500_not_performed_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_not_performed_cb)
        
        self.psi_500_image_widget = ImageUploadWidget("500 PSI Test")
        self.psi_500_image_widget.setVisible(False)
        psi_500_layout.addWidget(self.psi_500_image_widget)
        
        psi_500_group.setLayout(psi_500_layout)
        layout.addWidget(psi_500_group)
        
        notes_label = QLabel("Additional Notes:")
        layout.addWidget(notes_label)
        self.notes_field = QTextEdit()
        self.notes_field.setMaximumHeight(100)
        layout.addWidget(self.notes_field)
        
        self.image_widget = ImageUploadWidget("General Section Images")
        layout.addWidget(self.image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_250_checkbox(self, state):
        """Handle 250 PSI checkbox toggling and image/location visibility"""
        if state:
            if self.sender() == self.psi_250_leak_cb:
                self.psi_250_no_leak_cb.setChecked(False)
            else:
                self.psi_250_leak_cb.setChecked(False)
        
        self.psi_250_leak_location_container.setVisible(self.psi_250_leak_cb.isChecked())
        
        either_checked = self.psi_250_leak_cb.isChecked() or self.psi_250_no_leak_cb.isChecked()
        self.psi_250_image_widget.setVisible(either_checked)
    
    def toggle_500_checkbox(self, state):
        """Handle 500 PSI checkbox toggling and image/location visibility"""
        if state:
            if self.sender() == self.psi_500_leak_cb:
                self.psi_500_no_leak_cb.setChecked(False)
                self.psi_500_not_performed_cb.setChecked(False)
            elif self.sender() == self.psi_500_no_leak_cb:
                self.psi_500_leak_cb.setChecked(False)
                self.psi_500_not_performed_cb.setChecked(False)
            else:  # Not Performed
                self.psi_500_leak_cb.setChecked(False)
                self.psi_500_no_leak_cb.setChecked(False)
        
        self.psi_500_leak_location_container.setVisible(self.psi_500_leak_cb.isChecked())
        
        show_images = self.psi_500_leak_cb.isChecked() or self.psi_500_no_leak_cb.isChecked()
        self.psi_500_image_widget.setVisible(show_images)
    
    def get_pressure_data(self):
        """Return pressure testing data"""
        return {
            '250psi_leak': self.psi_250_leak_cb.isChecked(),
            '250psi_leak_face': self.psi_250_leak_face_cb.isChecked(),
            '250psi_leak_ubends': self.psi_250_leak_ubends_cb.isChecked(),
            '250psi_no_leak': self.psi_250_no_leak_cb.isChecked(),
            '250psi_images': self.psi_250_image_widget.image_paths,
            '500psi_leak': self.psi_500_leak_cb.isChecked(),
            '500psi_leak_face': self.psi_500_leak_face_cb.isChecked(),
            '500psi_leak_ubends': self.psi_500_leak_ubends_cb.isChecked(),
            '500psi_no_leak': self.psi_500_no_leak_cb.isChecked(),
            '500psi_not_performed': self.psi_500_not_performed_cb.isChecked(),
            '500psi_images': self.psi_500_image_widget.image_paths,
            'images': self.image_widget.image_paths,
            'notes': self.notes_field.toPlainText()
        }
    
    def get_notes(self):
        """Return notes text"""
        return self.notes_field.toPlainText()
    
    def get_images(self):
        """Return image paths"""
        return self.image_widget.image_paths

class ValveElectricalInspectionSection(QWidget):
    """Widget for valve electrical inspection with resistance measurement"""
    def __init__(self):
        super().__init__()
        self.title = "Electrical Inspection"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        title_label = QLabel("Electrical Inspection")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        status_group = QGroupBox("Select Status:")
        status_layout = QVBoxLayout()
        
        ok_container = QWidget()
        ok_layout = QVBoxLayout()
        ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_ok_cb = QCheckBox("Status: OK")
        self.status_ok_cb.setStyleSheet("font-weight: bold;")
        ok_font = self.status_ok_cb.font()
        ok_font.setPointSize(10)
        self.status_ok_cb.setFont(ok_font)
        self.status_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, True))
        ok_layout.addWidget(self.status_ok_cb)
        
        self.ok_image_container = QWidget()
        ok_image_layout = QVBoxLayout()
        ok_image_layout.setContentsMargins(30, 5, 0, 5)
        self.ok_image_widget = ImageUploadWidget("Status: OK")
        ok_image_layout.addWidget(self.ok_image_widget)
        self.ok_image_container.setLayout(ok_image_layout)
        self.ok_image_container.setVisible(False)
        ok_layout.addWidget(self.ok_image_container)
        
        ok_container.setLayout(ok_layout)
        status_layout.addWidget(ok_container)
        
        not_ok_container = QWidget()
        not_ok_layout = QVBoxLayout()
        not_ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_not_ok_cb = QCheckBox("Status: NOT OK")
        self.status_not_ok_cb.setStyleSheet("font-weight: bold;")
        not_ok_font = self.status_not_ok_cb.font()
        not_ok_font.setPointSize(10)
        self.status_not_ok_cb.setFont(not_ok_font)
        self.status_not_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, False))
        not_ok_layout.addWidget(self.status_not_ok_cb)
        
        not_ok_container.setLayout(not_ok_layout)
        status_layout.addWidget(not_ok_container)
        
        status_group.setLayout(status_layout)
        layout.addWidget(status_group)
        
        self.issues_container = QWidget()
        issues_layout = QVBoxLayout()
        issues_layout.setContentsMargins(30, 5, 0, 5)
        
        issues_group = QGroupBox("Specific Issues:")
        issues_group_layout = QVBoxLayout()
        
        self.issue_checkboxes = []
        self.issue_image_widgets = {}
        
        for issue in VALVE_ELECTRICAL_INSPECTION_PATTERNS["Status: NOT OK"]["sub_observations"][""]:
            obs_container = QWidget()
            obs_layout = QVBoxLayout()
            obs_layout.setContentsMargins(0, 2, 0, 2)
            
            checkbox_row = QHBoxLayout()
            
            cb = QCheckBox(issue)
            checkbox_row.addWidget(cb)
            self.issue_checkboxes.append(cb)
            
            toggle_btn = QPushButton("📷 Upload Images")
            toggle_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px 8px; font-size: 10px;")
            toggle_btn.setMaximumWidth(120)
            toggle_btn.setVisible(False)
            checkbox_row.addWidget(toggle_btn)
            checkbox_row.addStretch()
            
            obs_layout.addLayout(checkbox_row)
            
            obs_img_widget = ImageUploadWidget(f"Electrical - {issue[:30]}...")
            obs_img_widget.setVisible(False)
            obs_layout.addWidget(obs_img_widget)
            
            self.issue_image_widgets[issue] = obs_img_widget
            
            cb.stateChanged.connect(
                lambda state, btn=toggle_btn: btn.setVisible(state == 2)
            )
            
            toggle_btn.clicked.connect(
                lambda checked=False, img_w=obs_img_widget, btn=toggle_btn: (
                    img_w.setVisible(not img_w.isVisible()),
                    btn.setText("📷 Hide Images" if img_w.isVisible() else "📷 Upload Images")
                )
            )
            
            obs_container.setLayout(obs_layout)
            issues_group_layout.addWidget(obs_container)
        
        issues_group.setLayout(issues_group_layout)
        issues_layout.addWidget(issues_group)
        
        self.custom_section_container = QWidget()
        self.custom_section_layout = QVBoxLayout()
        self.custom_section_layout.setContentsMargins(0, 10, 0, 5)
        self.custom_section_container.setLayout(self.custom_section_layout)
        
        self.first_custom_btn = QPushButton("+ Add Custom Issue")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)
        
        issues_layout.addWidget(self.custom_section_container)
        
        self.not_ok_image_widget = ImageUploadWidget("Electrical Inspection - Issues")
        issues_layout.addWidget(self.not_ok_image_widget)
        
        self.issues_container.setLayout(issues_layout)
        self.issues_container.setVisible(False)
        layout.addWidget(self.issues_container)
        
        self.custom_patterns = []
        
        resistance_group = QGroupBox("Resistance Value")
        resistance_layout = QHBoxLayout()
        resistance_layout.addWidget(QLabel("Resistance (Ω):"))
        self.resistance_value = QLineEdit()
        self.resistance_value.setPlaceholderText("Enter resistance value...")
        resistance_layout.addWidget(self.resistance_value)
        resistance_group.setLayout(resistance_layout)
        layout.addWidget(resistance_group)
        
        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Enter any additional observations...")
        self.notes_text.setMaximumHeight(100)
        layout.addWidget(self.notes_text)
        
        general_images_label = QLabel("General Section Images:")
        general_images_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(general_images_label)
        
        self.general_image_widget = ImageUploadWidget("Electrical Inspection - General")
        layout.addWidget(self.general_image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_status(self, state, is_ok):
        """Handle mutual exclusivity of status checkboxes and visibility"""
        if state == 2:  # Checked
            if is_ok:
                self.status_not_ok_cb.setChecked(False)
                self.issues_container.setVisible(False)
                self.ok_image_container.setVisible(True)
            else:
                self.status_ok_cb.setChecked(False)
                self.ok_image_container.setVisible(False)
                self.issues_container.setVisible(True)
        else:  # Unchecked
            if is_ok:
                self.ok_image_container.setVisible(False)
            else:
                self.issues_container.setVisible(False)
    
    def add_custom_pattern(self):
        """Add a new custom issue entry"""
        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)
        
        custom_container = QWidget()
        custom_layout = QHBoxLayout()
        custom_layout.setContentsMargins(0, 2, 0, 2)
        
        text_input = QLineEdit()
        text_input.setPlaceholderText("Describe custom issue...")
        custom_layout.addWidget(text_input)
        
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 3px 8px;")
        delete_btn.clicked.connect(lambda: self.remove_custom_pattern(custom_container, text_input))
        custom_layout.addWidget(delete_btn)
        
        custom_container.setLayout(custom_layout)
        self.custom_section_layout.insertWidget(self.custom_section_layout.count() - 1, custom_container)
        
        self.custom_patterns.append({
            'container': custom_container,
            'text_input': text_input
        })
        
        if not hasattr(self, 'shared_add_another_btn'):
            self.shared_add_another_btn = QPushButton("+ Add Another Custom Issue")
            self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)
            self.custom_section_layout.addWidget(self.shared_add_another_btn)
    
    def remove_custom_pattern(self, container, text_input):
        """Remove a custom issue entry"""
        for custom in self.custom_patterns:
            if custom['text_input'] == text_input:
                self.custom_patterns.remove(custom)
                break
        container.deleteLater()
        
        if len(self.custom_patterns) == 0:
            if hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn.setVisible(False)
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)
    
    def get_test_data(self):
        """Return electrical testing data"""
        issues = []
        issue_images = {}
        
        for cb in self.issue_checkboxes:
            if cb.isChecked():
                issue_text = cb.text()
                issues.append(issue_text)
                if issue_text in self.issue_image_widgets:
                    issue_images[issue_text] = self.issue_image_widgets[issue_text].image_paths
        
        custom_list = []
        for custom_data in self.custom_patterns:
            text = custom_data['text_input'].text().strip()
            if text:
                custom_list.append({
                    'text': text,
                    'images': custom_data.get('images', []) if 'images' in custom_data else []
                })
        
        return {
            'ok': self.status_ok_cb.isChecked(),
            'not_ok': self.status_not_ok_cb.isChecked(),
            'ok_images': self.ok_image_widget.image_paths if hasattr(self, 'ok_image_widget') else [],
            'issues': issues,
            'issue_images': issue_images,
            'custom_patterns': custom_list,
            'not_ok_images': self.not_ok_image_widget.image_paths if hasattr(self, 'not_ok_image_widget') else [],
            'resistance': self.resistance_value.text().strip(),
            'general_images': self.general_image_widget.image_paths if hasattr(self, 'general_image_widget') else [],
            'notes': self.notes_text.toPlainText()
        }
    
    def get_notes(self):
        """Return notes text"""
        return self.notes_text.toPlainText()

class ValveMechanicalInspectionSection(QWidget):
    """Widget for valve mechanical inspection with bore measurements"""
    def __init__(self):
        super().__init__()
        self.title = "Mechanical Inspection"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        title_label = QLabel("Mechanical Inspection")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        bore_group = QGroupBox("Internal Diameter Measurements")
        bore_layout = QVBoxLayout()
        
        tb_layout = QHBoxLayout()
        tb_layout.addWidget(QLabel("Top to Bottom (mm):"))
        self.bore_top_bottom = QLineEdit()
        self.bore_top_bottom.setPlaceholderText("Enter measurement...")
        tb_layout.addWidget(self.bore_top_bottom)
        bore_layout.addLayout(tb_layout)
        
        lr_layout = QHBoxLayout()
        lr_layout.addWidget(QLabel("Left to Right (mm):"))
        self.bore_left_right = QLineEdit()
        self.bore_left_right.setPlaceholderText("Enter measurement...")
        lr_layout.addWidget(self.bore_left_right)
        bore_layout.addLayout(lr_layout)
        
        self.bore_consistent_cb = QCheckBox("Bore measurements are consistent")
        bore_layout.addWidget(self.bore_consistent_cb)
        
        # Add image upload widget for cylinder measurements
        cylinder_measurements_label = QLabel("Cylinder Measurements Images:")
        cylinder_measurements_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        bore_layout.addWidget(cylinder_measurements_label)
        
        self.cylinder_measurements_image_widget = ImageUploadWidget("Cylinder Measurements")
        bore_layout.addWidget(self.cylinder_measurements_image_widget)
        
        bore_group.setLayout(bore_layout)
        layout.addWidget(bore_group)
        
        movement_group = QGroupBox("Valve Movement")
        movement_layout = QVBoxLayout()
        
        # Left valve movement subsection
        left_group = QGroupBox("Left Valve Movement")
        left_layout = QVBoxLayout()
        
        self.left_smooth_cb = QCheckBox("Smooth movement")
        left_layout.addWidget(self.left_smooth_cb)
        
        self.left_not_smooth_cb = QCheckBox("Not smooth movement")
        left_layout.addWidget(self.left_not_smooth_cb)
        
        # Make left checkboxes mutually exclusive
        self.left_smooth_cb.stateChanged.connect(
            lambda state: self.left_not_smooth_cb.setChecked(False) if state == 2 else None
        )
        self.left_not_smooth_cb.stateChanged.connect(
            lambda state: self.left_smooth_cb.setChecked(False) if state == 2 else None
        )
        
        left_group.setLayout(left_layout)
        movement_layout.addWidget(left_group)
        
        # Right valve movement subsection
        right_group = QGroupBox("Right Valve Movement")
        right_layout = QVBoxLayout()
        
        self.right_smooth_cb = QCheckBox("Smooth movement")
        right_layout.addWidget(self.right_smooth_cb)
        
        self.right_not_smooth_cb = QCheckBox("Not smooth movement")
        right_layout.addWidget(self.right_not_smooth_cb)
        
        # Make right checkboxes mutually exclusive
        self.right_smooth_cb.stateChanged.connect(
            lambda state: self.right_not_smooth_cb.setChecked(False) if state == 2 else None
        )
        self.right_not_smooth_cb.stateChanged.connect(
            lambda state: self.right_smooth_cb.setChecked(False) if state == 2 else None
        )
        
        right_group.setLayout(right_layout)
        movement_layout.addWidget(right_group)
        
        movement_group.setLayout(movement_layout)
        layout.addWidget(movement_group)
        
        status_group = QGroupBox("Select Status:")
        status_layout = QVBoxLayout()
        
        ok_container = QWidget()
        ok_layout = QVBoxLayout()
        ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_ok_cb = QCheckBox("Status: OK")
        self.status_ok_cb.setStyleSheet("font-weight: bold;")
        ok_font = self.status_ok_cb.font()
        ok_font.setPointSize(10)
        self.status_ok_cb.setFont(ok_font)
        self.status_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, True))
        ok_layout.addWidget(self.status_ok_cb)
        
        self.ok_image_container = QWidget()
        ok_image_layout = QVBoxLayout()
        ok_image_layout.setContentsMargins(30, 5, 0, 5)
        self.ok_image_widget = ImageUploadWidget("Status: OK")
        ok_image_layout.addWidget(self.ok_image_widget)
        self.ok_image_container.setLayout(ok_image_layout)
        self.ok_image_container.setVisible(False)
        ok_layout.addWidget(self.ok_image_container)
        
        ok_container.setLayout(ok_layout)
        status_layout.addWidget(ok_container)
        
        not_ok_container = QWidget()
        not_ok_layout = QVBoxLayout()
        not_ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_not_ok_cb = QCheckBox("Status: NOT OK")
        self.status_not_ok_cb.setStyleSheet("font-weight: bold;")
        not_ok_font = self.status_not_ok_cb.font()
        not_ok_font.setPointSize(10)
        self.status_not_ok_cb.setFont(not_ok_font)
        self.status_not_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, False))
        not_ok_layout.addWidget(self.status_not_ok_cb)
        
        not_ok_container.setLayout(not_ok_layout)
        status_layout.addWidget(not_ok_container)
        
        status_group.setLayout(status_layout)
        layout.addWidget(status_group)
        
        self.issues_container = QWidget()
        issues_layout = QVBoxLayout()
        issues_layout.setContentsMargins(30, 5, 0, 5)
        
        issues_group = QGroupBox("Specific Issues:")
        issues_group_layout = QVBoxLayout()
        
        self.issue_checkboxes = []
        self.issue_image_widgets = {}
        
        for issue in VALVE_MECHANICAL_INSPECTION_PATTERNS["Status: NOT OK"]["sub_observations"][""]:
            obs_container = QWidget()
            obs_layout = QVBoxLayout()
            obs_layout.setContentsMargins(0, 2, 0, 2)
            
            checkbox_row = QHBoxLayout()
            
            cb = QCheckBox(issue)
            checkbox_row.addWidget(cb)
            self.issue_checkboxes.append(cb)
            
            toggle_btn = QPushButton("📷 Upload Images")
            toggle_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px 8px; font-size: 10px;")
            toggle_btn.setMaximumWidth(120)
            toggle_btn.setVisible(False)
            checkbox_row.addWidget(toggle_btn)
            checkbox_row.addStretch()
            
            obs_layout.addLayout(checkbox_row)
            
            obs_img_widget = ImageUploadWidget(f"Mechanical - {issue[:30]}...")
            obs_img_widget.setVisible(False)
            obs_layout.addWidget(obs_img_widget)
            
            self.issue_image_widgets[issue] = obs_img_widget
            
            cb.stateChanged.connect(
                lambda state, btn=toggle_btn: btn.setVisible(state == 2)
            )
            
            toggle_btn.clicked.connect(
                lambda checked=False, img_w=obs_img_widget, btn=toggle_btn: (
                    img_w.setVisible(not img_w.isVisible()),
                    btn.setText("📷 Hide Images" if img_w.isVisible() else "📷 Upload Images")
                )
            )
            
            obs_container.setLayout(obs_layout)
            issues_group_layout.addWidget(obs_container)
        
        issues_group.setLayout(issues_group_layout)
        issues_layout.addWidget(issues_group)
        
        self.custom_section_container = QWidget()
        self.custom_section_layout = QVBoxLayout()
        self.custom_section_layout.setContentsMargins(0, 10, 0, 5)
        self.custom_section_container.setLayout(self.custom_section_layout)
        
        self.first_custom_btn = QPushButton("+ Add Custom Issue")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)
        
        issues_layout.addWidget(self.custom_section_container)
        
        self.not_ok_image_widget = ImageUploadWidget("Mechanical Inspection - Issues")
        issues_layout.addWidget(self.not_ok_image_widget)
        
        self.issues_container.setLayout(issues_layout)
        self.issues_container.setVisible(False)
        layout.addWidget(self.issues_container)
        
        self.custom_patterns = []
        
        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Enter any additional observations...")
        self.notes_text.setMaximumHeight(100)
        layout.addWidget(self.notes_text)
        
        general_images_label = QLabel("General Section Images:")
        general_images_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(general_images_label)
        
        self.general_image_widget = ImageUploadWidget("Mechanical Inspection - General")
        layout.addWidget(self.general_image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_status(self, state, is_ok):
        """Handle mutual exclusivity of status checkboxes and visibility"""
        if state == 2:
            if is_ok:
                self.status_not_ok_cb.setChecked(False)
                self.issues_container.setVisible(False)
                self.ok_image_container.setVisible(True)
            else:
                self.status_ok_cb.setChecked(False)
                self.ok_image_container.setVisible(False)
                self.issues_container.setVisible(True)
        else:
            if is_ok:
                self.ok_image_container.setVisible(False)
            else:
                self.issues_container.setVisible(False)
    
    def add_custom_pattern(self):
        """Add a new custom issue entry"""
        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)
        
        custom_container = QWidget()
        custom_layout = QHBoxLayout()
        custom_layout.setContentsMargins(0, 2, 0, 2)
        
        text_input = QLineEdit()
        text_input.setPlaceholderText("Describe custom issue...")
        custom_layout.addWidget(text_input)
        
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 3px 8px;")
        delete_btn.clicked.connect(lambda: self.remove_custom_pattern(custom_container, text_input))
        custom_layout.addWidget(delete_btn)
        
        custom_container.setLayout(custom_layout)
        self.custom_section_layout.insertWidget(self.custom_section_layout.count() - 1, custom_container)
        
        self.custom_patterns.append({
            'container': custom_container,
            'text_input': text_input
        })
        
        if not hasattr(self, 'shared_add_another_btn'):
            self.shared_add_another_btn = QPushButton("+ Add Another Custom Issue")
            self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)
            self.custom_section_layout.addWidget(self.shared_add_another_btn)
    
    def remove_custom_pattern(self, container, text_input):
        """Remove a custom issue entry"""
        for custom in self.custom_patterns:
            if custom['text_input'] == text_input:
                self.custom_patterns.remove(custom)
                break
        container.deleteLater()
        
        if len(self.custom_patterns) == 0:
            if hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn.setVisible(False)
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)
    
    def get_test_data(self):
        """Return mechanical testing data"""
        issues = []
        issue_images = {}
        
        for cb in self.issue_checkboxes:
            if cb.isChecked():
                issue_text = cb.text()
                issues.append(issue_text)
                if issue_text in self.issue_image_widgets:
                    issue_images[issue_text] = self.issue_image_widgets[issue_text].image_paths
        
        custom_list = []
        for custom_data in self.custom_patterns:
            text = custom_data['text_input'].text().strip()
            if text:
                custom_list.append({
                    'text': text,
                    'images': custom_data.get('images', []) if 'images' in custom_data else []
                })
        
        return {
            'ok': self.status_ok_cb.isChecked(),
            'not_ok': self.status_not_ok_cb.isChecked(),
            'ok_images': self.ok_image_widget.image_paths if hasattr(self, 'ok_image_widget') else [],
            'issues': issues,
            'issue_images': issue_images,
            'custom_patterns': custom_list,
            'not_ok_images': self.not_ok_image_widget.image_paths if hasattr(self, 'not_ok_image_widget') else [],
            'bore_top_bottom': self.bore_top_bottom.text().strip(),
            'bore_left_right': self.bore_left_right.text().strip(),
            'bore_consistent': self.bore_consistent_cb.isChecked(),
            'cylinder_measurements_images': self.cylinder_measurements_image_widget.image_paths if hasattr(self, 'cylinder_measurements_image_widget') else [],
            'left_smooth': self.left_smooth_cb.isChecked(),
            'left_not_smooth': self.left_not_smooth_cb.isChecked(),
            'right_smooth': self.right_smooth_cb.isChecked(),
            'right_not_smooth': self.right_not_smooth_cb.isChecked(),
            'general_images': self.general_image_widget.image_paths if hasattr(self, 'general_image_widget') else [],
            'notes': self.notes_text.toPlainText()
        }
    
    def get_notes(self):
        """Return notes text"""
        return self.notes_text.toPlainText()

class CompressorTab(QWidget):
    """Widget for a single compressor with electrical testing, oil evaluation, scroll plate inspection, and sleeve bearing inspection"""
    def __init__(self, compressor_number):
        super().__init__()
        self.compressor_number = compressor_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)
        
        scroll_electrical = QScrollArea()
        scroll_electrical.setWidgetResizable(True)
        self.electrical_testing_section = CompressorElectricalTestingSection()
        scroll_electrical.setWidget(self.electrical_testing_section)
        self.component_tabs.addTab(scroll_electrical, "Electrical Testing")
        
        scroll_oil = QScrollArea()
        scroll_oil.setWidgetResizable(True)
        self.oil_evaluation_section = ComponentSection("Oil Evaluation", OIL_EVALUATION_PATTERNS)
        scroll_oil.setWidget(self.oil_evaluation_section)
        self.component_tabs.addTab(scroll_oil, "Oil Evaluation")
        
        scroll_plate = QScrollArea()
        scroll_plate.setWidgetResizable(True)
        self.scroll_plate_section = ComponentSection("Scroll Plate Inspection", SCROLL_PLATE_PATTERNS)
        scroll_plate.setWidget(self.scroll_plate_section)
        self.component_tabs.addTab(scroll_plate, "Scroll Plate Inspection")
        
        scroll_bearing = QScrollArea()
        scroll_bearing.setWidgetResizable(True)
        self.sleeve_bearing_section = ComponentSection("Sleeve Bearing Inspection", SLEEVE_BEARING_PATTERNS)
        scroll_bearing.setWidget(self.sleeve_bearing_section)
        self.component_tabs.addTab(scroll_bearing, "Sleeve Bearing Inspection")
        
        layout.addWidget(self.component_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all component sections in order"""
        return [
            (self.electrical_testing_section, None, 'compressor_electrical'),
            (self.oil_evaluation_section, OIL_EVALUATION_PATTERNS, 'component'),
            (self.scroll_plate_section, SCROLL_PLATE_PATTERNS, 'component'),
            (self.sleeve_bearing_section, SLEEVE_BEARING_PATTERNS, 'component')
        ]


class ReceiverPressureTestingSection(QWidget):
    """Widget for receiver pressure testing with 250psi and 500psi sections (no leak location options)"""
    def __init__(self):
        super().__init__()
        self.title = "Pressure Testing"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        title_label = QLabel("Pressure Testing")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        psi_250_group = QGroupBox("250 PSI Test")
        psi_250_layout = QVBoxLayout()
        
        self.psi_250_leak_cb = QCheckBox("Leak")
        self.psi_250_leak_cb.stateChanged.connect(lambda state: self.toggle_250_checkbox(state))
        psi_250_layout.addWidget(self.psi_250_leak_cb)
        
        self.psi_250_no_leak_cb = QCheckBox("No Leak")
        self.psi_250_no_leak_cb.stateChanged.connect(lambda state: self.toggle_250_checkbox(state))
        psi_250_layout.addWidget(self.psi_250_no_leak_cb)
        
        self.psi_250_image_widget = ImageUploadWidget("250 PSI Test")
        self.psi_250_image_widget.setVisible(False)
        psi_250_layout.addWidget(self.psi_250_image_widget)
        
        psi_250_group.setLayout(psi_250_layout)
        layout.addWidget(psi_250_group)
        
        psi_500_group = QGroupBox("500 PSI Test")
        psi_500_layout = QVBoxLayout()
        
        self.psi_500_leak_cb = QCheckBox("Leak")
        self.psi_500_leak_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_leak_cb)
        
        self.psi_500_no_leak_cb = QCheckBox("No Leak")
        self.psi_500_no_leak_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_no_leak_cb)
        
        self.psi_500_not_performed_cb = QCheckBox("Not Performed")
        self.psi_500_not_performed_cb.stateChanged.connect(lambda state: self.toggle_500_checkbox(state))
        psi_500_layout.addWidget(self.psi_500_not_performed_cb)
        
        self.psi_500_image_widget = ImageUploadWidget("500 PSI Test")
        self.psi_500_image_widget.setVisible(False)
        psi_500_layout.addWidget(self.psi_500_image_widget)
        
        psi_500_group.setLayout(psi_500_layout)
        layout.addWidget(psi_500_group)
        
        notes_label = QLabel("Additional Notes:")
        layout.addWidget(notes_label)
        self.notes_field = QTextEdit()
        self.notes_field.setMaximumHeight(100)
        layout.addWidget(self.notes_field)
        
        self.image_widget = ImageUploadWidget("General Section Images")
        layout.addWidget(self.image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_250_checkbox(self, state):
        """Handle 250 PSI checkbox toggling and image visibility"""
        if state:
            if self.sender() == self.psi_250_leak_cb:
                self.psi_250_no_leak_cb.setChecked(False)
            else:
                self.psi_250_leak_cb.setChecked(False)
        
        either_checked = self.psi_250_leak_cb.isChecked() or self.psi_250_no_leak_cb.isChecked()
        self.psi_250_image_widget.setVisible(either_checked)
    
    def toggle_500_checkbox(self, state):
        """Handle 500 PSI checkbox toggling and image visibility"""
        if state:
            if self.sender() == self.psi_500_leak_cb:
                self.psi_500_no_leak_cb.setChecked(False)
                self.psi_500_not_performed_cb.setChecked(False)
            elif self.sender() == self.psi_500_no_leak_cb:
                self.psi_500_leak_cb.setChecked(False)
                self.psi_500_not_performed_cb.setChecked(False)
            else:  # Not Performed
                self.psi_500_leak_cb.setChecked(False)
                self.psi_500_no_leak_cb.setChecked(False)
        
        show_images = self.psi_500_leak_cb.isChecked() or self.psi_500_no_leak_cb.isChecked()
        self.psi_500_image_widget.setVisible(show_images)
    
    def get_pressure_data(self):
        """Return pressure testing data (without leak location fields)"""
        return {
            '250psi_leak': self.psi_250_leak_cb.isChecked(),
            '250psi_no_leak': self.psi_250_no_leak_cb.isChecked(),
            '250psi_images': self.psi_250_image_widget.image_paths,
            '500psi_leak': self.psi_500_leak_cb.isChecked(),
            '500psi_no_leak': self.psi_500_no_leak_cb.isChecked(),
            '500psi_not_performed': self.psi_500_not_performed_cb.isChecked(),
            '500psi_images': self.psi_500_image_widget.image_paths,
            'notes': self.notes_field.toPlainText(),
            'images': self.image_widget.image_paths
        }
    
    def get_notes(self):
        """Return notes text"""
        return self.notes_field.toPlainText()
    
    def get_images(self):
        """Return general section images"""
        return self.image_widget.image_paths

class CoilTab(QWidget):
    """Widget for a single coil with visual inspection and pressure testing"""
    def __init__(self, coil_number):
        super().__init__()
        self.coil_number = coil_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)
        
        scroll_visual = QScrollArea()
        scroll_visual.setWidgetResizable(True)
        self.visual_inspection_section = ComponentSection("Visual Inspection", COIL_VISUAL_INSPECTION_PATTERNS)
        scroll_visual.setWidget(self.visual_inspection_section)
        self.component_tabs.addTab(scroll_visual, "Visual Inspection")
        
        scroll_pressure = QScrollArea()
        scroll_pressure.setWidgetResizable(True)
        self.pressure_testing_section = PressureTestingSection()
        scroll_pressure.setWidget(self.pressure_testing_section)
        self.component_tabs.addTab(scroll_pressure, "Pressure Testing")
        
        layout.addWidget(self.component_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all component sections in order"""
        return [
            (self.visual_inspection_section, COIL_VISUAL_INSPECTION_PATTERNS, 'component'),
            (self.pressure_testing_section, None, 'coil_pressure')
        ]

class ReceiverTab(QWidget):
    """Widget for a single receiver with visual inspection and pressure testing"""
    def __init__(self, receiver_number):
        super().__init__()
        self.receiver_number = receiver_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)
        
        scroll_visual = QScrollArea()
        scroll_visual.setWidgetResizable(True)
        self.visual_inspection_section = ComponentSection("Visual Inspection", RECEIVER_VISUAL_INSPECTION_PATTERNS)
        scroll_visual.setWidget(self.visual_inspection_section)
        self.component_tabs.addTab(scroll_visual, "Visual Inspection")
        
        scroll_pressure = QScrollArea()
        scroll_pressure.setWidgetResizable(True)
        self.pressure_testing_section = ReceiverPressureTestingSection()
        scroll_pressure.setWidget(self.pressure_testing_section)
        self.component_tabs.addTab(scroll_pressure, "Pressure Testing")
        
        layout.addWidget(self.component_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all component sections in order"""
        return [
            (self.visual_inspection_section, RECEIVER_VISUAL_INSPECTION_PATTERNS, 'component'),
            (self.pressure_testing_section, None, 'receiver_pressure')
        ]

class PCBElectricalTestingSection(QWidget):
    """Widget for PCB/Inverter board diode electrical testing"""
    def __init__(self):
        super().__init__()
        self.title = "Electrical Testing"
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        title_label = QLabel("Electrical Testing")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        status_group = QGroupBox("Select Status:")
        status_layout = QVBoxLayout()
        
        ok_container = QWidget()
        ok_layout = QVBoxLayout()
        ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_ok_cb = QCheckBox("Status: OK")
        self.status_ok_cb.setStyleSheet("font-weight: bold;")
        ok_font = self.status_ok_cb.font()
        ok_font.setPointSize(10)
        self.status_ok_cb.setFont(ok_font)
        self.status_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, True))
        ok_layout.addWidget(self.status_ok_cb)
        
        self.ok_image_container = QWidget()
        ok_image_layout = QVBoxLayout()
        ok_image_layout.setContentsMargins(30, 5, 0, 5)
        self.ok_image_widget = ImageUploadWidget("Status: OK")
        ok_image_layout.addWidget(self.ok_image_widget)
        self.ok_image_container.setLayout(ok_image_layout)
        self.ok_image_container.setVisible(False)
        ok_layout.addWidget(self.ok_image_container)
        
        ok_container.setLayout(ok_layout)
        status_layout.addWidget(ok_container)
        
        not_ok_container = QWidget()
        not_ok_layout = QVBoxLayout()
        not_ok_layout.setContentsMargins(0, 5, 0, 5)
        
        self.status_not_ok_cb = QCheckBox("Status: NOT OK")
        self.status_not_ok_cb.setStyleSheet("font-weight: bold;")
        not_ok_font = self.status_not_ok_cb.font()
        not_ok_font.setPointSize(10)
        self.status_not_ok_cb.setFont(not_ok_font)
        self.status_not_ok_cb.stateChanged.connect(lambda state: self.toggle_status(state, False))
        not_ok_layout.addWidget(self.status_not_ok_cb)
        
        not_ok_container.setLayout(not_ok_layout)
        status_layout.addWidget(not_ok_container)
        
        status_group.setLayout(status_layout)
        layout.addWidget(status_group)
        
        self.issues_container = QWidget()
        issues_layout = QVBoxLayout()
        issues_layout.setContentsMargins(30, 5, 0, 5)
        
        issues_group = QGroupBox("Diode Failure Issues:")
        issues_group_layout = QVBoxLayout()
        
        self.issue_checkboxes = []
        self.issue_image_widgets = {}
        
        for issue in PCB_ELECTRICAL_TESTING_PATTERNS["Status: NOT OK"]["sub_observations"][""]:
            obs_container = QWidget()
            obs_layout = QVBoxLayout()
            obs_layout.setContentsMargins(0, 2, 0, 2)
            
            checkbox_row = QHBoxLayout()
            
            cb = QCheckBox(issue)
            checkbox_row.addWidget(cb)
            self.issue_checkboxes.append(cb)
            
            toggle_btn = QPushButton("📷 Upload Images")
            toggle_btn.setStyleSheet("background-color: #4CAF50; color: white; padding: 3px 8px; font-size: 10px;")
            toggle_btn.setMaximumWidth(120)
            toggle_btn.setVisible(False)
            checkbox_row.addWidget(toggle_btn)
            checkbox_row.addStretch()
            
            obs_layout.addLayout(checkbox_row)
            
            obs_img_widget = ImageUploadWidget(f"Electrical - {issue[:30]}...")
            obs_img_widget.setVisible(False)
            obs_layout.addWidget(obs_img_widget)
            
            self.issue_image_widgets[issue] = obs_img_widget
            
            cb.stateChanged.connect(
                lambda state, btn=toggle_btn: btn.setVisible(state == 2)
            )
            
            toggle_btn.clicked.connect(
                lambda checked=False, img_w=obs_img_widget, btn=toggle_btn: (
                    img_w.setVisible(not img_w.isVisible()),
                    btn.setText("📷 Hide Images" if img_w.isVisible() else "📷 Upload Images")
                )
            )
            
            obs_container.setLayout(obs_layout)
            issues_group_layout.addWidget(obs_container)
        
        issues_group.setLayout(issues_group_layout)
        issues_layout.addWidget(issues_group)
        
        self.custom_section_container = QWidget()
        self.custom_section_layout = QVBoxLayout()
        self.custom_section_layout.setContentsMargins(0, 10, 0, 5)
        self.custom_section_container.setLayout(self.custom_section_layout)
        
        self.first_custom_btn = QPushButton("+ Add Custom Issue")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)
        
        issues_layout.addWidget(self.custom_section_container)
        
        self.not_ok_image_widget = ImageUploadWidget("Electrical Testing - Issues")
        issues_layout.addWidget(self.not_ok_image_widget)
        
        self.issues_container.setLayout(issues_layout)
        self.issues_container.setVisible(False)
        layout.addWidget(self.issues_container)
        
        self.custom_patterns = []
        
        # Korea testing checkbox - between status and notes
        self.korea_testing_cb = QCheckBox("Sent for live testing in Korea")
        self.korea_testing_cb.setStyleSheet("font-weight: bold; font-size: 15px; margin-top: 10px;")
        layout.addWidget(self.korea_testing_cb)
        
        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Enter any additional observations...")
        self.notes_text.setMaximumHeight(100)
        layout.addWidget(self.notes_text)
        
        general_images_label = QLabel("General Section Images:")
        general_images_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(general_images_label)
        
        self.general_image_widget = ImageUploadWidget("Electrical Testing - General")
        layout.addWidget(self.general_image_widget)
        
        layout.addStretch()
        self.setLayout(layout)
    
    def toggle_status(self, state, is_ok):
        """Handle mutual exclusivity of status checkboxes and visibility"""
        if state == 2:  # Checked
            if is_ok:
                self.status_not_ok_cb.setChecked(False)
                self.issues_container.setVisible(False)
                self.ok_image_container.setVisible(True)
            else:
                self.status_ok_cb.setChecked(False)
                self.ok_image_container.setVisible(False)
                self.issues_container.setVisible(True)
        else:  # Unchecked
            if is_ok:
                self.ok_image_container.setVisible(False)
            else:
                self.issues_container.setVisible(False)
    
    def add_custom_pattern(self):
        """Add a new custom issue entry"""
        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)
        
        custom_container = QWidget()
        custom_layout = QHBoxLayout()
        custom_layout.setContentsMargins(0, 2, 0, 2)
        
        text_input = QLineEdit()
        text_input.setPlaceholderText("Describe custom issue...")
        custom_layout.addWidget(text_input)
        
        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 3px 8px;")
        delete_btn.clicked.connect(lambda: self.remove_custom_pattern(custom_container, text_input))
        custom_layout.addWidget(delete_btn)
        
        custom_container.setLayout(custom_layout)
        self.custom_section_layout.insertWidget(self.custom_section_layout.count() - 1, custom_container)
        
        self.custom_patterns.append({
            'container': custom_container,
            'text_input': text_input
        })
        
        if not hasattr(self, 'shared_add_another_btn'):
            self.shared_add_another_btn = QPushButton("+ Add Another Custom Issue")
            self.shared_add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.shared_add_another_btn.clicked.connect(self.add_custom_pattern)
            self.custom_section_layout.addWidget(self.shared_add_another_btn)
    
    def remove_custom_pattern(self, container, text_input):
        """Remove a custom issue entry"""
        for custom in self.custom_patterns:
            if custom['text_input'] == text_input:
                self.custom_patterns.remove(custom)
                break
        container.deleteLater()
        
        if len(self.custom_patterns) == 0:
            if hasattr(self, 'shared_add_another_btn'):
                self.shared_add_another_btn.setVisible(False)
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)
    
    def get_test_data(self):
        """Return electrical testing data"""
        issues = []
        issue_images = {}
        
        for cb in self.issue_checkboxes:
            if cb.isChecked():
                issue_text = cb.text()
                issues.append(issue_text)
                if issue_text in self.issue_image_widgets:
                    issue_images[issue_text] = self.issue_image_widgets[issue_text].image_paths
        
        custom_list = []
        for custom_data in self.custom_patterns:
            text = custom_data['text_input'].text().strip()
            if text:
                custom_list.append({
                    'text': text,
                    'images': custom_data.get('images', [])
                })
        
        return {
            'ok': self.status_ok_cb.isChecked(),
            'not_ok': self.status_not_ok_cb.isChecked(),
            'ok_images': self.ok_image_widget.image_paths if hasattr(self, 'ok_image_widget') else [],
            'issues': issues,
            'issue_images': issue_images,
            'custom_patterns': custom_list,
            'not_ok_images': self.not_ok_image_widget.image_paths if hasattr(self, 'not_ok_image_widget') else [],
            'general_images': self.general_image_widget.image_paths if hasattr(self, 'general_image_widget') else [],
            'notes': self.notes_text.toPlainText(),
            'korea_testing': self.korea_testing_cb.isChecked()
        }
    
    def get_notes(self):
        """Return notes text"""
        return self.notes_text.toPlainText()


class PCBTab(QWidget):
    """Widget for a PCB/Inverter Board with visual inspection and electrical testing"""
    def __init__(self, pcb_number):
        super().__init__()
        self.pcb_number = pcb_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)
        
        scroll_visual = QScrollArea()
        scroll_visual.setWidgetResizable(True)
        self.visual_inspection_section = ComponentSection("Visual Inspection", PCB_VISUAL_INSPECTION_PATTERNS)
        scroll_visual.setWidget(self.visual_inspection_section)
        self.component_tabs.addTab(scroll_visual, "Visual Inspection")
        
        scroll_electrical = QScrollArea()
        scroll_electrical.setWidgetResizable(True)
        self.electrical_testing_section = PCBElectricalTestingSection()
        scroll_electrical.setWidget(self.electrical_testing_section)
        self.component_tabs.addTab(scroll_electrical, "Electrical Testing")
        
        layout.addWidget(self.component_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all sections in order"""
        return [
            (self.visual_inspection_section, PCB_VISUAL_INSPECTION_PATTERNS, 'component'),
            (self.electrical_testing_section, None, 'pcb_electrical')
        ]


class GeneralSectionWidget(QWidget):
    """A sub-section inside the General component - handles Status OK/NOT OK, manual issues, and images"""
    def __init__(self, section_name):
        super().__init__()
        self.title = section_name
        self.custom_patterns = []
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        title_label = QLabel(self.title)
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)

        status_group = QGroupBox("Select Status:")
        status_layout = QVBoxLayout()

        # --- Status OK ---
        ok_outer = QWidget()
        ok_outer_layout = QVBoxLayout()
        ok_outer_layout.setContentsMargins(0, 5, 0, 5)

        self.status_ok_cb = QCheckBox("Status: OK")
        self.status_ok_cb.setStyleSheet("font-weight: bold;")
        ok_font = self.status_ok_cb.font()
        ok_font.setPointSize(10)
        self.status_ok_cb.setFont(ok_font)
        self.status_ok_cb.stateChanged.connect(lambda state: self._toggle_status(state, True))
        ok_outer_layout.addWidget(self.status_ok_cb)

        self.ok_image_container = QWidget()
        ok_img_layout = QVBoxLayout()
        ok_img_layout.setContentsMargins(30, 5, 0, 5)
        self.ok_image_widget = ImageUploadWidget(f"{self.title} - Status OK")
        ok_img_layout.addWidget(self.ok_image_widget)
        self.ok_image_container.setLayout(ok_img_layout)
        self.ok_image_container.setVisible(False)
        ok_outer_layout.addWidget(self.ok_image_container)

        ok_outer.setLayout(ok_outer_layout)
        status_layout.addWidget(ok_outer)

        # --- Status NOT OK ---
        not_ok_outer = QWidget()
        not_ok_outer_layout = QVBoxLayout()
        not_ok_outer_layout.setContentsMargins(0, 5, 0, 5)

        self.status_not_ok_cb = QCheckBox("Status: NOT OK")
        self.status_not_ok_cb.setStyleSheet("font-weight: bold;")
        not_ok_font = self.status_not_ok_cb.font()
        not_ok_font.setPointSize(10)
        self.status_not_ok_cb.setFont(not_ok_font)
        self.status_not_ok_cb.stateChanged.connect(lambda state: self._toggle_status(state, False))
        not_ok_outer_layout.addWidget(self.status_not_ok_cb)

        not_ok_outer.setLayout(not_ok_outer_layout)
        status_layout.addWidget(not_ok_outer)

        status_group.setLayout(status_layout)
        layout.addWidget(status_group)

        # --- Issues container (shown only when NOT OK) ---
        self.issues_container = QWidget()
        issues_layout = QVBoxLayout()
        issues_layout.setContentsMargins(30, 5, 0, 5)

        issues_group = QGroupBox("Issues:")
        self.custom_section_layout = QVBoxLayout()
        self.custom_section_layout.setContentsMargins(0, 0, 0, 5)

        self.first_custom_btn = QPushButton("+ Add Issue")
        self.first_custom_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
        self.first_custom_btn.clicked.connect(self.add_custom_pattern)
        self.custom_section_layout.addWidget(self.first_custom_btn)

        issues_group.setLayout(self.custom_section_layout)
        issues_layout.addWidget(issues_group)

        self.not_ok_image_widget = ImageUploadWidget(f"{self.title} - Issues")
        issues_layout.addWidget(self.not_ok_image_widget)

        self.issues_container.setLayout(issues_layout)
        self.issues_container.setVisible(False)
        layout.addWidget(self.issues_container)

        # --- Notes ---
        notes_label = QLabel("Additional Notes:")
        notes_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(notes_label)

        self.notes_field = QTextEdit()
        self.notes_field.setPlaceholderText("Enter any additional observations...")
        self.notes_field.setMaximumHeight(100)
        layout.addWidget(self.notes_field)

        # --- General images ---
        gen_img_label = QLabel("General Section Images:")
        gen_img_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        layout.addWidget(gen_img_label)

        self.image_widget = ImageUploadWidget(f"{self.title} - General")
        layout.addWidget(self.image_widget)

        layout.addStretch()
        self.setLayout(layout)

    def _toggle_status(self, state, is_ok):
        if state == 2:
            if is_ok:
                self.status_not_ok_cb.setChecked(False)
                self.issues_container.setVisible(False)
                self.ok_image_container.setVisible(True)
            else:
                self.status_ok_cb.setChecked(False)
                self.ok_image_container.setVisible(False)
                self.issues_container.setVisible(True)
        else:
            if is_ok:
                self.ok_image_container.setVisible(False)
            else:
                self.issues_container.setVisible(False)

    def add_custom_pattern(self):
        if hasattr(self, 'first_custom_btn') and self.first_custom_btn.isVisible():
            self.first_custom_btn.setVisible(False)

        custom_container = QWidget()
        custom_layout = QHBoxLayout()
        custom_layout.setContentsMargins(0, 2, 0, 2)

        text_input = QLineEdit()
        text_input.setPlaceholderText("Describe issue...")
        custom_layout.addWidget(text_input)

        delete_btn = QPushButton("Delete")
        delete_btn.setStyleSheet("background-color: #f44336; color: white; padding: 3px 8px;")
        delete_btn.clicked.connect(lambda: self._remove_custom_pattern(custom_container, text_input))
        custom_layout.addWidget(delete_btn)

        custom_container.setLayout(custom_layout)
        self.custom_section_layout.insertWidget(self.custom_section_layout.count() - 1, custom_container)

        self.custom_patterns.append({'container': custom_container, 'text_input': text_input})

        if not hasattr(self, 'add_another_btn'):
            self.add_another_btn = QPushButton("+ Add Another Issue")
            self.add_another_btn.setStyleSheet("background-color: #2196F3; color: white; padding: 5px; font-weight: bold;")
            self.add_another_btn.clicked.connect(self.add_custom_pattern)
            self.custom_section_layout.addWidget(self.add_another_btn)

    def _remove_custom_pattern(self, container, text_input):
        self.custom_patterns = [c for c in self.custom_patterns if c['text_input'] != text_input]
        container.deleteLater()
        if not self.custom_patterns:
            if hasattr(self, 'add_another_btn'):
                self.add_another_btn.setVisible(False)
            if hasattr(self, 'first_custom_btn'):
                self.first_custom_btn.setVisible(True)

    def get_notes(self):
        return self.notes_field.toPlainText()

    def get_images(self):
        return self.image_widget.image_paths

    def get_section_data(self):
        custom_list = [{'text': c['text_input'].text().strip()} for c in self.custom_patterns if c['text_input'].text().strip()]
        return {
            'ok': self.status_ok_cb.isChecked(),
            'not_ok': self.status_not_ok_cb.isChecked(),
            'ok_images': self.ok_image_widget.image_paths,
            'custom_patterns': custom_list,
            'not_ok_images': self.not_ok_image_widget.image_paths,
            'general_images': self.image_widget.image_paths,
            'notes': self.notes_field.toPlainText(),
        }


class GeneralTab(QWidget):
    """A standalone General component with Visual Inspection and Testing sub-tabs"""
    def __init__(self, general_number):
        super().__init__()
        self.general_number = general_number
        self.component_name = ""
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()

        self.component_tabs = QTabWidget()
        tab_font = self.component_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.component_tabs.setFont(tab_font)

        # --- Component Name tab ---
        name_widget = QWidget()
        name_layout = QVBoxLayout()
        name_layout.setContentsMargins(20, 20, 20, 20)

        name_title = QLabel("Component Name")
        name_title_font = QFont()
        name_title_font.setPointSize(12)
        name_title_font.setBold(True)
        name_title.setFont(name_title_font)
        name_layout.addWidget(name_title)

        name_desc = QLabel("Enter the component type below. The tab label will update automatically.")
        name_desc.setStyleSheet("color: #555; margin-bottom: 8px;")
        name_layout.addWidget(name_desc)

        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("e.g. Expansion Valve, Filter Drier, Sensor...")
        self.name_input.setStyleSheet("font-size: 14px; padding: 6px;")
        self.name_input.textChanged.connect(self._on_name_changed)
        name_layout.addWidget(self.name_input)

        name_layout.addStretch()
        name_widget.setLayout(name_layout)
        self.component_tabs.addTab(name_widget, "Component Name")

        scroll_visual = QScrollArea()
        scroll_visual.setWidgetResizable(True)
        self.visual_inspection_section = GeneralSectionWidget("Visual Inspection")
        scroll_visual.setWidget(self.visual_inspection_section)
        self.component_tabs.addTab(scroll_visual, "Visual Inspection")

        scroll_testing = QScrollArea()
        scroll_testing.setWidgetResizable(True)
        self.testing_section = GeneralSectionWidget("Testing")
        scroll_testing.setWidget(self.testing_section)
        self.component_tabs.addTab(scroll_testing, "Testing")

        layout.addWidget(self.component_tabs)
        self.setLayout(layout)

    def _on_name_changed(self, text):
        """Update the equipment tab label when the user types a component name"""
        self.component_name = text.strip()
        # Walk up to find the QTabWidget that holds this GeneralTab and update its label
        parent = self.parent()
        while parent is not None:
            if isinstance(parent, QTabWidget):
                idx = parent.indexOf(self)
                if idx != -1:
                    label = f"{self.component_name} {self.general_number}" if self.component_name else f"General {self.general_number}"
                    parent.setTabText(idx, label)
                break
            parent = parent.parent()

    def get_component_name(self):
        return self.component_name

    def get_all_sections(self):
        return [
            (self.visual_inspection_section, None, 'general_section'),
            (self.testing_section, None, 'general_section'),
        ]


class ValveTab(QWidget):
    """Widget for a single four-way valve with external, electrical/mechanical, and internal inspections"""
    def __init__(self, valve_number):
        super().__init__()
        self.valve_number = valve_number
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout()
        
        self.main_tabs = QTabWidget()
        tab_font = self.main_tabs.font()
        tab_font.setPointSize(tab_font.pointSize() + 1)
        self.main_tabs.setFont(tab_font)
        
        scroll_external = QScrollArea()
        scroll_external.setWidgetResizable(True)
        self.external_inspection_section = ComponentSection("External Inspection", VALVE_EXTERNAL_INSPECTION_PATTERNS)
        scroll_external.setWidget(self.external_inspection_section)
        self.main_tabs.addTab(scroll_external, "External Inspection")
        
        electrical_mechanical_widget = QWidget()
        em_layout = QVBoxLayout()
        
        self.em_tabs = QTabWidget()
        em_tab_font = self.em_tabs.font()
        em_tab_font.setPointSize(em_tab_font.pointSize() + 1)
        self.em_tabs.setFont(em_tab_font)
        
        scroll_electrical = QScrollArea()
        scroll_electrical.setWidgetResizable(True)
        self.electrical_inspection_section = ValveElectricalInspectionSection()
        scroll_electrical.setWidget(self.electrical_inspection_section)
        self.em_tabs.addTab(scroll_electrical, "Electrical Inspection")
        
        scroll_mechanical = QScrollArea()
        scroll_mechanical.setWidgetResizable(True)
        self.mechanical_inspection_section = ValveMechanicalInspectionSection()
        scroll_mechanical.setWidget(self.mechanical_inspection_section)
        self.em_tabs.addTab(scroll_mechanical, "Mechanical Inspection")
        
        em_layout.addWidget(self.em_tabs)
        electrical_mechanical_widget.setLayout(em_layout)
        self.main_tabs.addTab(electrical_mechanical_widget, "Electrical and Mechanical Inspection")
        
        internal_widget = QWidget()
        internal_layout = QVBoxLayout()
        
        self.internal_tabs = QTabWidget()
        internal_tab_font = self.internal_tabs.font()
        internal_tab_font.setPointSize(internal_tab_font.pointSize() + 1)
        self.internal_tabs.setFont(internal_tab_font)
        
        scroll_cylinder_a = QScrollArea()
        scroll_cylinder_a.setWidgetResizable(True)
        self.internal_cylinder_a_section = ComponentSection("Internal Cylinder Inspection A", VALVE_INTERNAL_CYLINDER_A_PATTERNS)
        scroll_cylinder_a.setWidget(self.internal_cylinder_a_section)
        self.internal_tabs.addTab(scroll_cylinder_a, "Internal Cylinder Inspection A")
        
        scroll_cylinder_b = QScrollArea()
        scroll_cylinder_b.setWidgetResizable(True)
        self.internal_cylinder_b_section = ComponentSection("Internal Cylinder Inspection B", VALVE_INTERNAL_CYLINDER_B_PATTERNS)
        scroll_cylinder_b.setWidget(self.internal_cylinder_b_section)
        self.internal_tabs.addTab(scroll_cylinder_b, "Internal Cylinder Inspection B")
        
        internal_layout.addWidget(self.internal_tabs)
        internal_widget.setLayout(internal_layout)
        self.main_tabs.addTab(internal_widget, "Internal Inspection")
        
        layout.addWidget(self.main_tabs)
        self.setLayout(layout)
    
    def get_all_sections(self):
        """Return all sections in order"""
        return [
            (self.external_inspection_section, VALVE_EXTERNAL_INSPECTION_PATTERNS, 'component'),
            (self.electrical_inspection_section, None, 'valve_electrical'),
            (self.mechanical_inspection_section, None, 'valve_mechanical'),
            (self.internal_cylinder_a_section, VALVE_INTERNAL_CYLINDER_A_PATTERNS, 'component'),
            (self.internal_cylinder_b_section, VALVE_INTERNAL_CYLINDER_B_PATTERNS, 'component')
        ]

# Auto-Update System
class ReportGeneratorApp(QMainWindow):
    def __init__(self):
        """Initialize application. Author: Daniel Fraser"""
        super().__init__()
        
        self.motor_tabs = []
        self.motor_count = 0
        self.compressor_tabs = []
        self.compressor_count = 0
        self.coil_tabs = []
        self.coil_count = 0
        self.receiver_tabs = []
        self.receiver_count = 0
        self.valve_tabs = []
        self.valve_count = 0
        self.pcb_tabs = []
        self.pcb_count = 0
        self.general_tabs = []
        self.general_count = 0
        
        self.last_save_path = None
        self.autosave_timer = None
        self.is_autosaving = False
        self.autosave_interval = 60000  # 1 minute in milliseconds
        
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Report Generator")
        self.setGeometry(100, 100, 1000, 750)


        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout()


        header_layout = QGridLayout()
        
        label_font = QFont()
        label_font.setPointSize(10)

        rma_label = QLabel("RMA:")
        rma_label.setFont(label_font)
        header_layout.addWidget(rma_label, 0, 0)
        self.equipment_id = QLineEdit()
        header_layout.addWidget(self.equipment_id, 0, 1)

        inspection_date_label = QLabel("Inspection Date:")
        inspection_date_label.setFont(label_font)
        header_layout.addWidget(inspection_date_label, 0, 2)
        self.inspection_date = QLineEdit()
        self.inspection_date.setText(datetime.now().strftime("%Y-%m-%d"))
        header_layout.addWidget(self.inspection_date, 0, 3)


        unit_model_label = QLabel("Unit Model:")
        unit_model_label.setFont(label_font)
        header_layout.addWidget(unit_model_label, 1, 0)
        self.unit_model = QLineEdit()
        header_layout.addWidget(self.unit_model, 1, 1)

        unit_serial_label = QLabel("Unit Serial:")
        unit_serial_label.setFont(label_font)
        header_layout.addWidget(unit_serial_label, 1, 2)
        self.unit_serial = QLineEdit()
        header_layout.addWidget(self.unit_serial, 1, 3)


        part_number_label = QLabel("Part Number:")
        part_number_label.setFont(label_font)
        header_layout.addWidget(part_number_label, 2, 0)
        self.part_number = QLineEdit()
        header_layout.addWidget(self.part_number, 2, 1)

        warranty_claim_label = QLabel("Warranty Claim Number:")
        warranty_claim_label.setFont(label_font)
        header_layout.addWidget(warranty_claim_label, 2, 2)
        self.warranty_claim = QLineEdit()
        header_layout.addWidget(self.warranty_claim, 2, 3)


        customer_email_label = QLabel("Customer and Email:")
        customer_email_label.setFont(label_font)
        header_layout.addWidget(customer_email_label, 3, 0)
        self.customer_email = QLineEdit()
        header_layout.addWidget(self.customer_email, 3, 1)

        total_label = QLabel("Total:")
        total_label.setFont(label_font)
        header_layout.addWidget(total_label, 3, 2)
        self.total_field = QLineEdit()
        header_layout.addWidget(self.total_field, 3, 3)


        install_date_label = QLabel("Install Date:")
        install_date_label.setFont(label_font)
        header_layout.addWidget(install_date_label, 4, 0)
        self.install_date = QLineEdit()
        header_layout.addWidget(self.install_date, 4, 1)

        failure_date_label = QLabel("Failure Date:")
        failure_date_label.setFont(label_font)
        header_layout.addWidget(failure_date_label, 4, 2)
        self.failure_date = QLineEdit()
        header_layout.addWidget(self.failure_date, 4, 3)


        lab_tech_label = QLabel("Lab Technician:")
        lab_tech_label.setFont(label_font)
        header_layout.addWidget(lab_tech_label, 5, 0)
        self.lab_tech = QLineEdit()
        header_layout.addWidget(self.lab_tech, 5, 1)

        manager_label = QLabel("Manager:")
        manager_label.setFont(label_font)
        header_layout.addWidget(manager_label, 5, 2)
        self.manager_name = QLineEdit()
        header_layout.addWidget(self.manager_name, 5, 3)
        
        input_font = QFont()
        input_font.setPointSize(12)
        for field in [self.equipment_id, self.inspection_date, self.unit_model, self.unit_serial,
                      self.part_number, self.warranty_claim, self.customer_email, self.total_field,
                      self.install_date, self.failure_date, self.lab_tech, self.manager_name]:
            field.setFont(input_font)

        main_layout.addLayout(header_layout)
        

        content_layout = QHBoxLayout()
        
        buttons_layout = QVBoxLayout()
        
        add_motor_btn = QPushButton("+ Add Motor")
        add_motor_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #0b7dda;
            }
        """)
        add_motor_btn.clicked.connect(self.add_motor_tab)
        buttons_layout.addWidget(add_motor_btn)
        
        add_compressor_btn = QPushButton("+ Add Compressor")
        add_compressor_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
        """)
        add_compressor_btn.clicked.connect(self.add_compressor_tab)
        buttons_layout.addWidget(add_compressor_btn)
        
        add_coil_btn = QPushButton("+ Add Coil")
        add_coil_btn.setStyleSheet("""
            QPushButton {
                background-color: #9C27B0;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #7B1FA2;
            }
        """)
        add_coil_btn.clicked.connect(self.add_coil_tab)
        buttons_layout.addWidget(add_coil_btn)
        
        add_valve_btn = QPushButton("+ Add Valve")
        add_valve_btn.setStyleSheet("""
            QPushButton {
                background-color: #00BCD4;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #0097A7;
            }
        """)
        add_valve_btn.clicked.connect(self.add_valve_tab)
        buttons_layout.addWidget(add_valve_btn)
        
        add_accumulator_btn = QPushButton("+ Add Receiver")
        add_accumulator_btn.setStyleSheet("""
            QPushButton {
                background-color: #E91E63;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #C2185B;
            }
        """)
        add_accumulator_btn.clicked.connect(self.add_receiver_tab)
        buttons_layout.addWidget(add_accumulator_btn)
        
        add_pcb_btn = QPushButton("+ Add PCB")
        add_pcb_btn.setStyleSheet("""
            QPushButton {
                background-color: #607D8B;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #455A64;
            }
        """)
        add_pcb_btn.clicked.connect(self.add_pcb_tab)
        buttons_layout.addWidget(add_pcb_btn)

        add_general_btn = QPushButton("+ Add General")
        add_general_btn.setStyleSheet("""
            QPushButton {
                background-color: #795548;
                color: white;
                padding: 10px;
                font-size: 16px;
                font-weight: bold;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #5D4037;
            }
        """)
        add_general_btn.clicked.connect(self.add_general_tab)
        buttons_layout.addWidget(add_general_btn)
        
        content_layout.addLayout(buttons_layout)
        
        self.equipment_tabs_widget = QTabWidget()
        self.equipment_tabs_widget.setTabsClosable(True)
        self.equipment_tabs_widget.setMovable(True)
        self.equipment_tabs_widget.tabCloseRequested.connect(self.remove_equipment_tab)
        
        equipment_tab_font = self.equipment_tabs_widget.font()
        equipment_tab_font.setPointSize(10)
        self.equipment_tabs_widget.setFont(equipment_tab_font)
        
        content_layout.addWidget(self.equipment_tabs_widget, 1)
        
        main_layout.addLayout(content_layout)
        buttons_layout.addStretch()
        
        button_layout = QHBoxLayout()
        
        self.save_btn = QPushButton("💾 Save Progress")
        self.save_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                padding: 10px;
                font-size: 15px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #0b7dda;
            }
        """)
        self.save_btn.clicked.connect(self.save_progress)
        button_layout.addWidget(self.save_btn)
        
        self.load_btn = QPushButton("📂 Load Progress")
        self.load_btn.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                padding: 10px;
                font-size: 15px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #e68900;
            }
        """)
        self.load_btn.clicked.connect(self.load_progress)
        button_layout.addWidget(self.load_btn)

        button_layout.addStretch()

        help_btn = QPushButton("?")
        help_btn.setToolTip("Open User Guide")
        help_btn.setFixedSize(32, 32)
        help_btn.setStyleSheet("""
            QPushButton {
                background-color: #1976D2;
                color: white;
                font-size: 16px;
                font-weight: bold;
                border-radius: 16px;
                border: none;
            }
            QPushButton:hover {
                background-color: #1565C0;
            }
        """)
        help_btn.clicked.connect(self.open_user_guide)

        button_layout.addWidget(help_btn)

        
        button_layout.addStretch()

        self.generate_btn = QPushButton("Generate Report")
        self.generate_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 10px;
                font-size: 15px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        self.generate_btn.clicked.connect(self.generate_report)
        button_layout.addWidget(self.generate_btn)
        
        main_layout.addLayout(button_layout)
        
        main_widget.setLayout(main_layout)
    
    def add_motor_tab(self):
        """Add a new motor tab"""
        self.motor_count += 1
        motor_tab = MotorTab(self.motor_count)
        self.motor_tabs.append(motor_tab)
        self.equipment_tabs_widget.addTab(motor_tab, f"Motor {self.motor_count}")
        self.equipment_tabs_widget.setCurrentWidget(motor_tab)
    
    def add_compressor_tab(self):
        """Add a new compressor tab"""
        self.compressor_count += 1
        compressor_tab = CompressorTab(self.compressor_count)
        self.compressor_tabs.append(compressor_tab)
        self.equipment_tabs_widget.addTab(compressor_tab, f"Compressor {self.compressor_count}")
        self.equipment_tabs_widget.setCurrentWidget(compressor_tab)
    
    def add_coil_tab(self):
        """Add a new coil tab"""
        self.coil_count += 1
        coil_tab = CoilTab(self.coil_count)
        self.coil_tabs.append(coil_tab)
        self.equipment_tabs_widget.addTab(coil_tab, f"Coil {self.coil_count}")
        self.equipment_tabs_widget.setCurrentWidget(coil_tab)
    
    def add_receiver_tab(self):
        """Add a new receiver tab"""
        self.receiver_count += 1
        receiver_tab = ReceiverTab(self.receiver_count)
        self.receiver_tabs.append(receiver_tab)
        self.equipment_tabs_widget.addTab(receiver_tab, f"Receiver {self.receiver_count}")
        self.equipment_tabs_widget.setCurrentWidget(receiver_tab)
    
    def add_valve_tab(self):
        """Add a new valve tab"""
        self.valve_count += 1
        valve_tab = ValveTab(self.valve_count)
        self.valve_tabs.append(valve_tab)
        self.equipment_tabs_widget.addTab(valve_tab, f"Valve {self.valve_count}")
        self.equipment_tabs_widget.setCurrentWidget(valve_tab)
    
    def add_pcb_tab(self):
        """Add a new PCB/Inverter Board tab"""
        self.pcb_count += 1
        pcb_tab = PCBTab(self.pcb_count)
        self.pcb_tabs.append(pcb_tab)
        self.equipment_tabs_widget.addTab(pcb_tab, f"PCB {self.pcb_count}")
        self.equipment_tabs_widget.setCurrentWidget(pcb_tab)

    def add_general_tab(self):
        """Add a new General component tab"""
        self.general_count += 1
        general_tab = GeneralTab(self.general_count)
        self.general_tabs.append(general_tab)
        self.equipment_tabs_widget.addTab(general_tab, f"General {self.general_count}")
        self.equipment_tabs_widget.setCurrentWidget(general_tab)
    
    def open_user_guide(self):
        """Open the bundled user guide PDF"""
        import tempfile, shutil, subprocess
        if getattr(sys, 'frozen', False):
            pdf_src = os.path.join(sys._MEIPASS, '_assets', 'user_guide.pdf')
        else:
            pdf_src = os.path.join(os.path.dirname(__file__), '_assets', 'user_guide.pdf')
        
        if not os.path.exists(pdf_src):
            QMessageBox.warning(self, "User Guide Not Found",
                                "The user guide PDF could not be found.")
            return
        
        # Copy to a temp file so the exe doesn't lock the internal asset
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        tmp.close()
        shutil.copy2(pdf_src, tmp.name)
        
        try:
            if platform.system() == 'Windows':
                os.startfile(tmp.name)
            elif platform.system() == 'Darwin':
                subprocess.Popen(['open', tmp.name])
            else:
                subprocess.Popen(['xdg-open', tmp.name])
        except Exception as e:
            QMessageBox.warning(self, "Could Not Open PDF",
                                f"Unable to open user guide:\n{str(e)}")

    def remove_equipment_tab(self, index):
        """Remove a motor, compressor, coil, receiver, or valve tab"""
        widget = self.equipment_tabs_widget.widget(index)
        tab_text = self.equipment_tabs_widget.tabText(index)
        
        reply = QMessageBox.question(
            self, "Confirm Removal",
            f"Are you sure you want to remove {tab_text}?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.equipment_tabs_widget.removeTab(index)
            

            if isinstance(widget, MotorTab):
                self.motor_tabs.remove(widget)
                self.motor_count = len(self.motor_tabs)
            elif isinstance(widget, CompressorTab):
                self.compressor_tabs.remove(widget)
                self.compressor_count = len(self.compressor_tabs)
            elif isinstance(widget, CoilTab):
                self.coil_tabs.remove(widget)
                self.coil_count = len(self.coil_tabs)
            elif isinstance(widget, ReceiverTab):
                self.receiver_tabs.remove(widget)
                self.receiver_count = len(self.receiver_tabs)
            elif isinstance(widget, ValveTab):
                self.valve_tabs.remove(widget)
                self.valve_count = len(self.valve_tabs)
            elif isinstance(widget, PCBTab):
                self.pcb_tabs.remove(widget)
                self.pcb_count = len(self.pcb_tabs)
            elif isinstance(widget, GeneralTab):
                self.general_tabs.remove(widget)
                self.general_count = len(self.general_tabs)
            
            motor_index = 1
            compressor_index = 1
            coil_index = 1
            receiver_index = 1
            valve_index = 1
            pcb_index = 1
            general_index = 1
            for i in range(self.equipment_tabs_widget.count()):
                widget = self.equipment_tabs_widget.widget(i)
                if isinstance(widget, MotorTab):
                    self.equipment_tabs_widget.setTabText(i, f"Motor {motor_index}")
                    widget.motor_number = motor_index
                    motor_index += 1
                elif isinstance(widget, CompressorTab):
                    self.equipment_tabs_widget.setTabText(i, f"Compressor {compressor_index}")
                    widget.compressor_number = compressor_index
                    compressor_index += 1
                elif isinstance(widget, CoilTab):
                    self.equipment_tabs_widget.setTabText(i, f"Coil {coil_index}")
                    widget.coil_number = coil_index
                    coil_index += 1
                elif isinstance(widget, ReceiverTab):
                    self.equipment_tabs_widget.setTabText(i, f"Receiver {receiver_index}")
                    widget.receiver_number = receiver_index
                    receiver_index += 1
                elif isinstance(widget, ValveTab):
                    self.equipment_tabs_widget.setTabText(i, f"Valve {valve_index}")
                    widget.valve_number = valve_index
                    valve_index += 1
                elif isinstance(widget, PCBTab):
                    self.equipment_tabs_widget.setTabText(i, f"PCB {pcb_index}")
                    widget.pcb_number = pcb_index
                    pcb_index += 1
                elif isinstance(widget, GeneralTab):
                    widget.general_number = general_index
                    label = f"{widget.component_name} {general_index}" if widget.component_name else f"General {general_index}"
                    self.equipment_tabs_widget.setTabText(i, label)
                    general_index += 1
    
    def save_progress(self):
        """Save current report progress to JSON file with images folder"""
        try:
            if not self.equipment_id.text().strip():
                QMessageBox.warning(self, "RMA Required", "Please enter an RMA number before saving.")
                return
            
            default_filename = self.equipment_id.text().strip()
            default_path = os.path.join(os.path.expanduser("~"), default_filename + ".rpt")
            
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Progress",
                default_path,
                "Report Project (*.rpt);;JSON Files (*.json)"
            )
            
            if not save_path:
                return
            
            if not save_path.endswith('.rpt') and not save_path.endswith('.json'):
                save_path += '.rpt'
            
            base_name = os.path.splitext(save_path)[0]
            parent_dir = os.path.dirname(base_name)
            folder_name = os.path.basename(base_name)
            
            if platform.system() == 'Windows':
                project_folder = base_name + "_project"
            else:
                project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
            
            images_folder = os.path.join(project_folder, "images")
            os.makedirs(images_folder, exist_ok=True)
            
            if platform.system() == 'Windows':
                try:
                    FILE_ATTRIBUTE_HIDDEN = 0x02
                    ctypes.windll.kernel32.SetFileAttributesW(project_folder, FILE_ATTRIBUTE_HIDDEN)
                except:
                    pass  # If setting hidden fails, just continue
            
            data = {
                'version': '1.0',
                'header': {
                    'rma': self.equipment_id.text(),
                    'inspection_date': self.inspection_date.text(),
                    'unit_model': self.unit_model.text(),
                    'unit_serial': self.unit_serial.text(),
                    'customer_email': self.customer_email.text(),
                    'total_field': self.total_field.text(),
                    'lab_tech': self.lab_tech.text(),
                    'manager_name': self.manager_name.text(),
                    'install_date': self.install_date.text(),
                    'failure_date': self.failure_date.text(),
                    'warranty_claim': self.warranty_claim.text(),
                    'part_number': self.part_number.text()
                },
                'equipment': []
            }
            
            for i in range(self.equipment_tabs_widget.count()):
                widget = self.equipment_tabs_widget.widget(i)
                if isinstance(widget, MotorTab):
                    equipment_type = 'motor'
                elif isinstance(widget, CoilTab):
                    equipment_type = 'coil'
                elif isinstance(widget, ReceiverTab):
                    equipment_type = 'receiver'
                elif isinstance(widget, ValveTab):
                    equipment_type = 'valve'
                elif isinstance(widget, PCBTab):
                    equipment_type = 'pcb'
                elif isinstance(widget, GeneralTab):
                    equipment_type = 'general'
                else:
                    equipment_type = 'compressor'
                
                equipment_data = {
                    'type': equipment_type,
                    'name': self.equipment_tabs_widget.tabText(i)
                }
                
                if isinstance(widget, MotorTab):
                    equipment_data['sections'] = self._save_motor_tab(widget, images_folder)
                elif isinstance(widget, CoilTab):
                    equipment_data['sections'] = self._save_coil_tab(widget, images_folder)
                elif isinstance(widget, ReceiverTab):
                    equipment_data['sections'] = self._save_receiver_tab(widget, images_folder)
                elif isinstance(widget, ValveTab):
                    equipment_data['sections'] = self._save_valve_tab(widget, images_folder)
                elif isinstance(widget, PCBTab):
                    equipment_data['sections'] = self._save_pcb_tab(widget, images_folder)
                elif isinstance(widget, GeneralTab):
                    equipment_data['sections'] = self._save_general_tab(widget, images_folder)
                else:
                    equipment_data['sections'] = self._save_compressor_tab(widget, images_folder)
                
                data['equipment'].append(equipment_data)
            
            json_path = os.path.join(project_folder, "report_state.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            
            shutil.copy(json_path, save_path)
            
            self.last_save_path = save_path
            
            self.start_autosave()
            
            QMessageBox.information(self, "Success", f"Progress saved to:\n{save_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save progress:\n{str(e)}")
    
    def start_autosave(self):
        """Start the autosave timer"""
        if self.autosave_timer is None:
            self.autosave_timer = QTimer()
            self.autosave_timer.timeout.connect(self.perform_autosave)
        
        if not self.autosave_timer.isActive():
            self.autosave_timer.start(self.autosave_interval)
    
    def stop_autosave(self):
        """Stop the autosave timer"""
        if self.autosave_timer and self.autosave_timer.isActive():
            self.autosave_timer.stop()
    
    def has_matching_save_file(self):
        """Check if a save file has been created for this session"""
        return bool(self.last_save_path and os.path.exists(self.last_save_path))
    
    def perform_autosave(self):
        """Silently save progress to the last saved path"""
        if not self.has_matching_save_file() or self.is_autosaving:
            return
        
        try:
            self.is_autosaving = True
            
            original_text = self.save_btn.text()
            original_style = self.save_btn.styleSheet()
            self.save_btn.setText("💾 Autosaving...")
            self.save_btn.setStyleSheet("""
                QPushButton {
                    background-color: #9E9E9E;
                    color: white;
                    padding: 10px;
                    font-size: 15px;
                    font-weight: bold;
                    border-radius: 5px;
                }
            """)
            
            QApplication.processEvents()
            
            save_path = self.last_save_path
            
            base_name = os.path.splitext(save_path)[0]
            parent_dir = os.path.dirname(base_name)
            folder_name = os.path.basename(base_name)
            
            if platform.system() == 'Windows':
                project_folder = base_name + "_project"
            else:
                project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
            
            images_folder = os.path.join(project_folder, "images")
            os.makedirs(images_folder, exist_ok=True)
            
            data = {
                'version': '1.0',
                'header': {
                    'rma': self.equipment_id.text(),
                    'inspection_date': self.inspection_date.text(),
                    'unit_model': self.unit_model.text(),
                    'unit_serial': self.unit_serial.text(),
                    'customer_email': self.customer_email.text(),
                    'total_field': self.total_field.text(),
                    'lab_tech': self.lab_tech.text(),
                    'manager_name': self.manager_name.text(),
                    'install_date': self.install_date.text(),
                    'failure_date': self.failure_date.text(),
                    'warranty_claim': self.warranty_claim.text(),
                    'part_number': self.part_number.text()
                },
                'equipment': []
            }
            
            for i in range(self.equipment_tabs_widget.count()):
                widget = self.equipment_tabs_widget.widget(i)
                if isinstance(widget, MotorTab):
                    equipment_type = 'motor'
                elif isinstance(widget, CoilTab):
                    equipment_type = 'coil'
                elif isinstance(widget, ReceiverTab):
                    equipment_type = 'receiver'
                elif isinstance(widget, ValveTab):
                    equipment_type = 'valve'
                elif isinstance(widget, PCBTab):
                    equipment_type = 'pcb'
                elif isinstance(widget, GeneralTab):
                    equipment_type = 'general'
                else:
                    equipment_type = 'compressor'
                
                equipment_data = {
                    'type': equipment_type,
                    'name': self.equipment_tabs_widget.tabText(i)
                }
                
                if isinstance(widget, MotorTab):
                    equipment_data['sections'] = self._save_motor_tab(widget, images_folder)
                elif isinstance(widget, CoilTab):
                    equipment_data['sections'] = self._save_coil_tab(widget, images_folder)
                elif isinstance(widget, ReceiverTab):
                    equipment_data['sections'] = self._save_receiver_tab(widget, images_folder)
                elif isinstance(widget, ValveTab):
                    equipment_data['sections'] = self._save_valve_tab(widget, images_folder)
                elif isinstance(widget, PCBTab):
                    equipment_data['sections'] = self._save_pcb_tab(widget, images_folder)
                elif isinstance(widget, GeneralTab):
                    equipment_data['sections'] = self._save_general_tab(widget, images_folder)
                else:
                    equipment_data['sections'] = self._save_compressor_tab(widget, images_folder)
                
                data['equipment'].append(equipment_data)
            
            json_path = os.path.join(project_folder, "report_state.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            
            shutil.copy(json_path, save_path)
            
            QTimer.singleShot(1000, lambda: (
                self.save_btn.setText(original_text),
                self.save_btn.setStyleSheet(original_style)
            ))
            
        except Exception as e:
            print(f"Autosave failed: {str(e)}")
        finally:
            self.is_autosaving = False
    
    def closeEvent(self, event):
        """Handle application close - save and stop autosave timer"""
        if self.has_matching_save_file():
            # Save file exists for current RMA - auto-save silently
            try:
                save_path = self.last_save_path
                
                base_name = os.path.splitext(save_path)[0]
                parent_dir = os.path.dirname(base_name)
                folder_name = os.path.basename(base_name)
                
                if platform.system() == 'Windows':
                    project_folder = base_name + "_project"
                else:
                    project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
                
                images_folder = os.path.join(project_folder, "images")
                os.makedirs(images_folder, exist_ok=True)
                
                data = {
                    'version': '1.0',
                    'header': {
                        'rma': self.equipment_id.text(),
                        'inspection_date': self.inspection_date.text(),
                        'unit_model': self.unit_model.text(),
                        'unit_serial': self.unit_serial.text(),
                        'customer_email': self.customer_email.text(),
                        'total_field': self.total_field.text(),
                        'lab_tech': self.lab_tech.text(),
                        'manager_name': self.manager_name.text(),
                        'install_date': self.install_date.text(),
                        'failure_date': self.failure_date.text(),
                        'warranty_claim': self.warranty_claim.text(),
                        'part_number': self.part_number.text()
                    },
                    'equipment': []
                }
                
                for i in range(self.equipment_tabs_widget.count()):
                    widget = self.equipment_tabs_widget.widget(i)
                    if isinstance(widget, MotorTab):
                        equipment_type = 'motor'
                    elif isinstance(widget, CoilTab):
                        equipment_type = 'coil'
                    elif isinstance(widget, ReceiverTab):
                        equipment_type = 'receiver'
                    elif isinstance(widget, ValveTab):
                        equipment_type = 'valve'
                    elif isinstance(widget, PCBTab):
                        equipment_type = 'pcb'
                    elif isinstance(widget, GeneralTab):
                        equipment_type = 'general'
                    else:
                        equipment_type = 'compressor'
                    
                    equipment_data = {
                        'type': equipment_type,
                        'name': self.equipment_tabs_widget.tabText(i)
                    }
                    
                    if isinstance(widget, MotorTab):
                        equipment_data['sections'] = self._save_motor_tab(widget, images_folder)
                    elif isinstance(widget, CoilTab):
                        equipment_data['sections'] = self._save_coil_tab(widget, images_folder)
                    elif isinstance(widget, ReceiverTab):
                        equipment_data['sections'] = self._save_receiver_tab(widget, images_folder)
                    elif isinstance(widget, ValveTab):
                        equipment_data['sections'] = self._save_valve_tab(widget, images_folder)
                    elif isinstance(widget, PCBTab):
                        equipment_data['sections'] = self._save_pcb_tab(widget, images_folder)
                    elif isinstance(widget, GeneralTab):
                        equipment_data['sections'] = self._save_general_tab(widget, images_folder)
                    else:
                        equipment_data['sections'] = self._save_compressor_tab(widget, images_folder)
                    
                    data['equipment'].append(equipment_data)
                
                json_path = os.path.join(project_folder, "report_state.json")
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
                
                shutil.copy(json_path, save_path)
                
            except Exception as e:
                print(f"Save on close failed: {str(e)}")
        else:
            # No save file exists - prompt user using a dialog with no X button
            dialog = QDialog(self)
            dialog.setWindowTitle("Save Progress?")
            # CustomizeWindowHint removes the X button; only Yes/No buttons can dismiss it
            dialog.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
            dialog.setFixedSize(520, 150)
            
            dlg_layout = QVBoxLayout()
            dlg_layout.setContentsMargins(24, 24, 24, 18)
            dlg_layout.setSpacing(16)
            
            # Icon + text row
            icon_row = QHBoxLayout()
            icon_label = QLabel()
            icon_label.setPixmap(dialog.style().standardPixmap(dialog.style().SP_MessageBoxWarning))
            icon_row.addWidget(icon_label)
            icon_row.addSpacing(12)
            
            dlg_label = QLabel("Save progress before closing?\nUnsaved changes will be permanently lost.")
            dlg_label.setWordWrap(True)
            dlg_font = dlg_label.font()
            dlg_font.setPointSize(13)
            dlg_label.setFont(dlg_font)
            icon_row.addWidget(dlg_label)
            icon_row.addStretch()
            dlg_layout.addLayout(icon_row)
            
            btn_layout = QHBoxLayout()
            btn_layout.setSpacing(10)
            
            btn_style = "font-size: 14px; padding: 6px 18px;"
            
            cancel_btn = QPushButton("Cancel")
            cancel_btn.setMinimumWidth(120)
            cancel_btn.setMinimumHeight(38)
            cancel_btn.setStyleSheet(btn_style)
            
            discard_btn = QPushButton("Close without saving")
            discard_btn.setMinimumWidth(170)
            discard_btn.setMinimumHeight(38)
            discard_btn.setStyleSheet(btn_style)
            
            save_btn = QPushButton("Save and Close")
            save_btn.setDefault(True)
            save_btn.setMinimumWidth(150)
            save_btn.setMinimumHeight(38)
            save_btn.setStyleSheet(btn_style + " font-weight: bold;")
            
            btn_layout.addWidget(cancel_btn)
            btn_layout.addStretch()
            btn_layout.addWidget(discard_btn)
            btn_layout.addStretch()
            btn_layout.addWidget(save_btn)
            dlg_layout.addLayout(btn_layout)
            dialog.setLayout(dlg_layout)
            
            choice = [None]
            save_btn.clicked.connect(lambda: (choice.__setitem__(0, 'yes'), dialog.accept()))
            discard_btn.clicked.connect(lambda: (choice.__setitem__(0, 'no'), dialog.accept()))
            cancel_btn.clicked.connect(lambda: (choice.__setitem__(0, 'cancel'), dialog.accept()))
            dialog.exec_()
            
            if choice[0] == 'yes':
                # Save progress, then close if save was successful
                self.save_progress()
                if self.last_save_path:
                    self.stop_autosave()
                    event.accept()
                else:
                    # User cancelled the save dialog - don't close
                    event.ignore()
                return
            elif choice[0] == 'no':
                # User chose not to save - close the program
                self.stop_autosave()
                event.accept()
                return
            else:
                # Cancel - return to program without closing
                event.ignore()
                return
        
        self.stop_autosave()
        event.accept()
    
    # Auto-Update Methods
    
    def _save_motor_tab(self, motor_tab, images_folder):
        """Save motor tab data"""
        sections_data = []
        
        for section_data in motor_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type
            }
            
            if section_type == 'test':
                test_data = section.get_test_data()
                
                # Copy test images and update paths in test_data
                test_data_copy = {}
                for test_type in ['audio', 'vibration', 'temperature', 'resistance']:
                    if test_type in test_data:
                        test_data_copy[test_type] = test_data[test_type].copy()
                        if 'images' in test_data[test_type]:
                            # Copy images and store relative paths
                            test_data_copy[test_type]['images'] = self._copy_images(
                                test_data[test_type]['images'], 
                                images_folder
                            )
                
                section_info['test_data'] = test_data_copy
                
                # Save notes and images for test sections too (if they exist)
                section_info['notes'] = section.get_notes() if hasattr(section, 'get_notes') else ''
                section_info['images'] = self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                section_info['notes'] = section.get_notes()
                section_info['images'] = self._copy_images(section.get_images(), images_folder)
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list,
                        'images': {}
                    }
                    if section.title in ["Motor Housing", "Motor Shaft", "Electrical Connection"]:
                        for obs in obs_list:
                            obs_images = section.get_observation_images(pattern, obs)
                            if obs_images:
                                section_info['observations'][pattern]['images'][obs] = self._copy_images(obs_images, images_folder)
                    
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_compressor_tab(self, compressor_tab, images_folder):
        """Save compressor tab data"""
        sections_data = []
        
        for section_data in compressor_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type,
                'notes': section.get_notes() if hasattr(section, 'get_notes') else '',
                'images': self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
            }
            
            if section_type == 'compressor_electrical':
                test_data = section.get_test_data()
                test_data_copy = test_data.copy()
                test_data_copy['ok_images'] = self._copy_images(test_data.get('ok_images', []), images_folder)
                test_data_copy['not_ok_images'] = self._copy_images(test_data.get('not_ok_images', []), images_folder)
                test_data_copy['general_images'] = self._copy_images(test_data.get('general_images', []), images_folder)
                test_data_copy['resistance_images'] = self._copy_images(test_data.get('resistance_images', []), images_folder)
                section_info['test_data'] = test_data_copy
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list
                    }
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_coil_tab(self, coil_tab, images_folder):
        """Save coil tab data"""
        sections_data = []
        
        for section_data in coil_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type,
                'notes': section.get_notes() if hasattr(section, 'get_notes') else '',
                'images': self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
            }
            
            if section_type == 'coil_pressure':
                pressure_data = section.get_pressure_data()
                pressure_data_copy = pressure_data.copy()
                pressure_data_copy['250psi_images'] = self._copy_images(pressure_data.get('250psi_images', []), images_folder)
                pressure_data_copy['500psi_images'] = self._copy_images(pressure_data.get('500psi_images', []), images_folder)
                pressure_data_copy['images'] = self._copy_images(pressure_data.get('images', []), images_folder)
                section_info['pressure_data'] = pressure_data_copy
                section_info['images'] = []  # Already saved in pressure_data
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list
                    }
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_receiver_tab(self, receiver_tab, images_folder):
        """Save receiver tab data"""
        sections_data = []
        
        for section_data in receiver_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type,
                'notes': section.get_notes() if hasattr(section, 'get_notes') else '',
                'images': self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
            }
            
            if section_type in ('coil_pressure', 'receiver_pressure'):
                pressure_data = section.get_pressure_data()
                pressure_data_copy = pressure_data.copy()
                pressure_data_copy['250psi_images'] = self._copy_images(pressure_data.get('250psi_images', []), images_folder)
                pressure_data_copy['500psi_images'] = self._copy_images(pressure_data.get('500psi_images', []), images_folder)
                pressure_data_copy['images'] = self._copy_images(pressure_data.get('images', []), images_folder)
                section_info['pressure_data'] = pressure_data_copy
                section_info['images'] = []  # Already saved in pressure_data
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list
                    }
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_valve_tab(self, valve_tab, images_folder):
        """Save valve tab data"""
        sections_data = []
        
        for section_data in valve_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type,
                'notes': section.get_notes() if hasattr(section, 'get_notes') else ''
            }
            
            if section_type == 'valve_electrical':
                test_data = section.get_test_data()
                test_data_copy = test_data.copy()
                test_data_copy['ok_images'] = self._copy_images(test_data.get('ok_images', []), images_folder)
                test_data_copy['not_ok_images'] = self._copy_images(test_data.get('not_ok_images', []), images_folder)
                test_data_copy['general_images'] = self._copy_images(test_data.get('general_images', []), images_folder)
                
                issue_images_copy = {}
                for issue, imgs in test_data.get('issue_images', {}).items():
                    issue_images_copy[issue] = self._copy_images(imgs, images_folder)
                test_data_copy['issue_images'] = issue_images_copy
                
                section_info['test_data'] = test_data_copy
            elif section_type == 'valve_mechanical':
                test_data = section.get_test_data()
                test_data_copy = test_data.copy()
                test_data_copy['ok_images'] = self._copy_images(test_data.get('ok_images', []), images_folder)
                test_data_copy['not_ok_images'] = self._copy_images(test_data.get('not_ok_images', []), images_folder)
                test_data_copy['general_images'] = self._copy_images(test_data.get('general_images', []), images_folder)
                test_data_copy['cylinder_measurements_images'] = self._copy_images(test_data.get('cylinder_measurements_images', []), images_folder)
                
                issue_images_copy = {}
                for issue, imgs in test_data.get('issue_images', {}).items():
                    issue_images_copy[issue] = self._copy_images(imgs, images_folder)
                test_data_copy['issue_images'] = issue_images_copy
                
                section_info['test_data'] = test_data_copy
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                section_info['images'] = self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list
                    }
                    
                    obs_images = {}
                    for obs in obs_list:
                        obs_imgs = section.get_observation_images(pattern, obs)
                        if obs_imgs:
                            obs_images[obs] = self._copy_images(obs_imgs, images_folder)
                    if obs_images:
                        section_info['observations'][pattern]['observation_images'] = obs_images
                    
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_pcb_tab(self, pcb_tab, images_folder):
        """Save PCB tab data"""
        sections_data = []
        
        for section_data in pcb_tab.get_all_sections():
            section, patterns_dict, section_type = section_data
            
            section_info = {
                'title': section.title,
                'type': section_type,
                'notes': section.get_notes() if hasattr(section, 'get_notes') else ''
            }
            
            if section_type == 'pcb_electrical':
                test_data = section.get_test_data()
                test_data_copy = test_data.copy()
                test_data_copy['ok_images'] = self._copy_images(test_data.get('ok_images', []), images_folder)
                test_data_copy['not_ok_images'] = self._copy_images(test_data.get('not_ok_images', []), images_folder)
                test_data_copy['general_images'] = self._copy_images(test_data.get('general_images', []), images_folder)
                
                issue_images_copy = {}
                for issue, imgs in test_data.get('issue_images', {}).items():
                    issue_images_copy[issue] = self._copy_images(imgs, images_folder)
                test_data_copy['issue_images'] = issue_images_copy
                
                section_info['test_data'] = test_data_copy
            else:
                section_info['selected_patterns'] = section.get_selected_patterns()
                section_info['observations'] = {}
                section_info['custom_patterns'] = section.get_custom_patterns()
                section_info['images'] = self._copy_images(section.get_images(), images_folder) if hasattr(section, 'get_images') else []
                
                for pattern in section.get_selected_patterns():
                    obs_list = section.get_selected_observations(pattern)
                    section_info['observations'][pattern] = {
                        'selected': obs_list
                    }
                    
                    obs_images = {}
                    for obs in obs_list:
                        obs_imgs = section.get_observation_images(pattern, obs)
                        if obs_imgs:
                            obs_images[obs] = self._copy_images(obs_imgs, images_folder)
                    if obs_images:
                        section_info['observations'][pattern]['observation_images'] = obs_images
                    
                    pattern_images = section.get_pattern_images(pattern)
                    if pattern_images:
                        section_info['observations'][pattern]['pattern_images'] = self._copy_images(pattern_images, images_folder)
            
            sections_data.append(section_info)
        
        return sections_data
    
    def _save_general_tab(self, general_tab, images_folder):
        """Save General component tab data"""
        sections_data = []
        for section, _, section_type in general_tab.get_all_sections():
            sec_data = section.get_section_data()
            sec_data_copy = sec_data.copy()
            sec_data_copy['ok_images'] = self._copy_images(sec_data.get('ok_images', []), images_folder)
            sec_data_copy['not_ok_images'] = self._copy_images(sec_data.get('not_ok_images', []), images_folder)
            sec_data_copy['general_images'] = self._copy_images(sec_data.get('general_images', []), images_folder)
            sections_data.append({
                'title': section.title,
                'type': section_type,
                'section_data': sec_data_copy,
            })
        return {'component_name': general_tab.get_component_name(), 'sections': sections_data}

    def _load_general_tab(self, general_tab, sections_data, images_folder):
        """Load General component tab data"""
        # sections_data may be the new dict format or the old list format
        if isinstance(sections_data, dict):
            component_name = sections_data.get('component_name', '')
            sections_list = sections_data.get('sections', [])
        else:
            component_name = ''
            sections_list = sections_data

        if component_name:
            general_tab.name_input.setText(component_name)
            # Tab label update is triggered by textChanged signal automatically

        section_map = {s.title: s for s, _, _ in general_tab.get_all_sections()}
        for section_info in sections_list:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            sec_data = section_info.get('section_data', {})
            if hasattr(section, 'status_ok_cb'):
                section.status_ok_cb.setChecked(sec_data.get('ok', False))
            if hasattr(section, 'status_not_ok_cb'):
                section.status_not_ok_cb.setChecked(sec_data.get('not_ok', False))
            for custom in sec_data.get('custom_patterns', []):
                section.add_custom_pattern()
                if 'text' in custom and section.custom_patterns:
                    section.custom_patterns[-1]['text_input'].setText(custom['text'])
            for img_name in sec_data.get('ok_images', []):
                img_path = os.path.join(images_folder, img_name)
                if os.path.exists(img_path):
                    section.ok_image_widget.add_image(img_path)
            for img_name in sec_data.get('not_ok_images', []):
                img_path = os.path.join(images_folder, img_name)
                if os.path.exists(img_path):
                    section.not_ok_image_widget.add_image(img_path)
            for img_name in sec_data.get('general_images', []):
                img_path = os.path.join(images_folder, img_name)
                if os.path.exists(img_path):
                    section.image_widget.add_image(img_path)
            if hasattr(section, 'notes_field'):
                section.notes_field.setPlainText(sec_data.get('notes', ''))

    def _copy_images(self, image_paths, dest_folder):
        """Copy images to project folder and return relative paths"""
        import filecmp
        
        copied_paths = []
        for img_path in image_paths:
            if not os.path.exists(img_path):
                continue
            
            filename = os.path.basename(img_path)
            dest_path = os.path.join(dest_folder, filename)
            
            # Check if image already exists and is identical
            if os.path.exists(dest_path):
                try:
                    if filecmp.cmp(img_path, dest_path, shallow=False):
                        # Same file! Reuse it - no need to copy again
                        copied_paths.append(filename)
                        continue
                except:
                    pass  # If comparison fails, proceed with normal logic
            
            # File doesn't exist or is different - need to copy
            base, ext = os.path.splitext(filename)
            counter = 1
            new_filename = filename
            
            # Find unique filename if needed
            while os.path.exists(os.path.join(dest_folder, new_filename)):
                # Check if numbered version is the same file
                check_path = os.path.join(dest_folder, new_filename)
                try:
                    if filecmp.cmp(img_path, check_path, shallow=False):
                        # It's the same file with different name! Reuse it.
                        copied_paths.append(new_filename)
                        break
                except:
                    pass
                
                new_filename = f"{base}_{counter}{ext}"
                counter += 1
            else:
                # Copy the file (only if we didn't find a match above)
                dest_path = os.path.join(dest_folder, new_filename)
                shutil.copy2(img_path, dest_path)
                copied_paths.append(new_filename)
        
        return copied_paths
    
    def load_progress_from_path(self, load_path):
        """Load report progress from a specific file path (for double-click opening)"""
        try:
            if not os.path.exists(load_path):
                QMessageBox.warning(self, "File Not Found", f"Cannot find file: {load_path}")
                return False
            
            if not (load_path.endswith('.rpt') or load_path.endswith('.json')):
                QMessageBox.warning(self, "Invalid File", "Please select a .rpt or .json file")
                return False
            
            base_name = os.path.splitext(load_path)[0]
            parent_dir = os.path.dirname(base_name)
            folder_name = os.path.basename(base_name)
            
            if platform.system() == 'Windows':
                project_folder = base_name + "_project"
            else:
                project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
            
            if not os.path.exists(project_folder):
                project_folder = base_name + "_project"
            
            if not os.path.exists(project_folder):
                project_folder = os.path.dirname(load_path)
            
            json_path = load_path
            if os.path.exists(os.path.join(project_folder, "report_state.json")):
                json_path = os.path.join(project_folder, "report_state.json")
            
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            while self.equipment_tabs_widget.count() > 0:
                self.equipment_tabs_widget.removeTab(0)
            self.motor_tabs = []
            self.compressor_tabs = []
            self.coil_tabs = []
            self.valve_tabs = []
            self.pcb_tabs = []
            self.general_tabs = []
            self.motor_count = 0
            self.compressor_count = 0
            self.coil_count = 0
            self.valve_count = 0
            self.pcb_count = 0
            self.general_count = 0
            
            header = data.get('header', {})
            self.equipment_id.setText(header.get('rma', ''))
            self.inspection_date.setText(header.get('inspection_date', ''))
            self.unit_model.setText(header.get('unit_model', ''))
            self.unit_serial.setText(header.get('unit_serial', ''))
            self.customer_email.setText(header.get('customer_email', ''))
            self.total_field.setText(header.get('total_field', ''))
            self.lab_tech.setText(header.get('lab_tech', ''))
            self.manager_name.setText(header.get('manager_name', ''))
            self.install_date.setText(header.get('install_date', ''))
            self.failure_date.setText(header.get('failure_date', ''))
            self.warranty_claim.setText(header.get('warranty_claim', ''))
            self.part_number.setText(header.get('part_number', ''))
            
            images_folder = os.path.join(project_folder, "images")
            for equipment in data.get('equipment', []):
                if equipment['type'] == 'motor':
                    self.add_motor_tab()
                    tab = self.motor_tabs[-1]
                    self._load_motor_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'coil':
                    self.add_coil_tab()
                    tab = self.coil_tabs[-1]
                    self._load_coil_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'receiver':
                    self.add_receiver_tab()
                    tab = self.receiver_tabs[-1]
                    self._load_receiver_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'valve':
                    self.add_valve_tab()
                    tab = self.valve_tabs[-1]
                    self._load_valve_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'pcb':
                    self.add_pcb_tab()
                    tab = self.pcb_tabs[-1]
                    self._load_pcb_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'general':
                    self.add_general_tab()
                    tab = self.general_tabs[-1]
                    self._load_general_tab(tab, equipment['sections'], images_folder)
                else:
                    self.add_compressor_tab()
                    tab = self.compressor_tabs[-1]
                    self._load_compressor_tab(tab, equipment['sections'], images_folder)
            
            self.last_save_path = load_path
            self.start_autosave()
            
            QMessageBox.information(self, "Success", f"Progress loaded from {os.path.basename(load_path)}")
            return True
            
        except json.JSONDecodeError:
            QMessageBox.critical(self, "Error", "Failed to load progress: Invalid JSON file")
            return False
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load progress: {str(e)}")
            return False
    
    def load_progress(self):
        """Load report progress from JSON file"""
        try:
            load_path, _ = QFileDialog.getOpenFileName(
                self,
                "Load Progress",
                os.path.expanduser("~"),
                "Report Project (*.rpt *.json);;All Files (*)"
            )
            
            if not load_path:
                return
            
            if load_path.endswith('.rpt') or load_path.endswith('.json'):
                base_name = os.path.splitext(load_path)[0]
                parent_dir = os.path.dirname(base_name)
                folder_name = os.path.basename(base_name)
                
                if platform.system() == 'Windows':
                    project_folder = base_name + "_project"
                else:
                    project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
                
                if not os.path.exists(project_folder):
                    project_folder = base_name + "_project"
                
                if not os.path.exists(project_folder):
                    project_folder = os.path.dirname(load_path)
                
                json_path = load_path
                if os.path.exists(os.path.join(project_folder, "report_state.json")):
                    json_path = os.path.join(project_folder, "report_state.json")
            else:
                QMessageBox.warning(self, "Invalid File", "Please select a .rpt or .json file")
                return
            
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            while self.equipment_tabs_widget.count() > 0:
                self.equipment_tabs_widget.removeTab(0)
            self.motor_tabs = []
            self.compressor_tabs = []
            self.coil_tabs = []
            self.motor_count = 0
            self.compressor_count = 0
            self.coil_count = 0
            self.valve_tabs = []
            self.valve_count = 0
            self.pcb_tabs = []
            self.pcb_count = 0
            self.general_tabs = []
            self.general_count = 0
            header = data.get('header', {})
            self.equipment_id.setText(header.get('rma', ''))
            self.inspection_date.setText(header.get('inspection_date', ''))
            self.unit_model.setText(header.get('unit_model', ''))
            self.unit_serial.setText(header.get('unit_serial', ''))
            self.customer_email.setText(header.get('customer_email', ''))
            self.total_field.setText(header.get('total_field', ''))
            self.lab_tech.setText(header.get('lab_tech', ''))
            self.manager_name.setText(header.get('manager_name', ''))
            self.install_date.setText(header.get('install_date', ''))
            self.failure_date.setText(header.get('failure_date', ''))
            self.warranty_claim.setText(header.get('warranty_claim', ''))
            self.part_number.setText(header.get('part_number', ''))
            
            images_folder = os.path.join(project_folder, "images")
            for equipment in data.get('equipment', []):
                if equipment['type'] == 'motor':
                    self.add_motor_tab()
                    tab = self.motor_tabs[-1]
                    self._load_motor_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'coil':
                    self.add_coil_tab()
                    tab = self.coil_tabs[-1]
                    self._load_coil_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'receiver':
                    self.add_receiver_tab()
                    tab = self.receiver_tabs[-1]
                    self._load_receiver_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'valve':
                    self.add_valve_tab()
                    tab = self.valve_tabs[-1]
                    self._load_valve_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'pcb':
                    self.add_pcb_tab()
                    tab = self.pcb_tabs[-1]
                    self._load_pcb_tab(tab, equipment['sections'], images_folder)
                elif equipment['type'] == 'general':
                    self.add_general_tab()
                    tab = self.general_tabs[-1]
                    self._load_general_tab(tab, equipment['sections'], images_folder)
                else:
                    self.add_compressor_tab()
                    tab = self.compressor_tabs[-1]
                    self._load_compressor_tab(tab, equipment['sections'], images_folder)
            
            self.last_save_path = load_path
            
            self.start_autosave()
            
            QMessageBox.information(self, "Success", "Progress loaded successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load progress:\n{str(e)}")
    
    def _load_motor_tab(self, motor_tab, sections_data, images_folder):
        """Load motor tab data"""
        section_map = {s.title: s for s, _, _ in motor_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] == 'test':
                if 'test_data' in section_info:
                    test_data = section_info['test_data']
                    if hasattr(section, 'audio_value'):
                        section.audio_value.setText(test_data.get('audio', {}).get('value', ''))
                    if hasattr(section, 'vibration_rpm'):
                        section.vibration_rpm.setText(str(test_data.get('vibration', {}).get('rpm', '')))
                    if hasattr(section, 'vibration_value'):
                        section.vibration_value.setText(str(test_data.get('vibration', {}).get('value', '')))
                    if hasattr(section, 'temperature_value'):
                        section.temperature_value.setText(test_data.get('temperature', {}).get('value', ''))
                    if hasattr(section, 'res_black_red'):
                        section.res_black_red.setText(test_data.get('resistance', {}).get('black_red', ''))
                    if hasattr(section, 'res_black_white'):
                        section.res_black_white.setText(test_data.get('resistance', {}).get('black_white', ''))
                    if hasattr(section, 'resistance_ok_cb'):
                        section.resistance_ok_cb.setChecked(test_data.get('resistance', {}).get('ok', True))
                    if hasattr(section, 'five_wire_cb'):
                        is_five_wire = test_data.get('resistance', {}).get('is_five_wire', False)
                        section.five_wire_cb.setChecked(is_five_wire)
                        if is_five_wire:
                            if hasattr(section, 'res_black_yellow'):
                                section.res_black_yellow.setText(test_data.get('resistance', {}).get('black_yellow', ''))
                            if hasattr(section, 'res_black_blue'):
                                section.res_black_blue.setText(test_data.get('resistance', {}).get('black_blue', ''))
                    
                    # Load test images back into their respective image widgets
                    if hasattr(section, 'audio_images') and 'audio' in test_data:
                        for img_name in test_data['audio'].get('images', []):
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.audio_images.add_image(img_path)
                    
                    if hasattr(section, 'vibration_images') and 'vibration' in test_data:
                        for img_name in test_data['vibration'].get('images', []):
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.vibration_images.add_image(img_path)
                    
                    if hasattr(section, 'temperature_images') and 'temperature' in test_data:
                        for img_name in test_data['temperature'].get('images', []):
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.temperature_images.add_image(img_path)
                    
                    if hasattr(section, 'resistance_images') and 'resistance' in test_data:
                        for img_name in test_data['resistance'].get('images', []):
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.resistance_images.add_image(img_path)
                
                # Load notes and images for test sections (if they exist)
                if hasattr(section, 'notes_text') and 'notes' in section_info:
                    section.notes_text.setPlainText(section_info.get('notes', ''))
                
                if hasattr(section, 'image_widget') and 'images' in section_info:
                    for img_name in section_info.get('images', []):
                        img_path = os.path.join(images_folder, img_name)
                        if os.path.exists(img_path):
                            section.image_widget.add_image(img_path)
                
                continue
            
            for pattern in section_info.get('selected_patterns', []):
                if pattern in section.checkboxes:
                    section.checkboxes[pattern].setChecked(True)
                    
                    if pattern in section_info.get('observations', {}):
                        obs_data = section_info['observations'][pattern]
                        for obs in obs_data.get('selected', []):
                            if pattern in section.observation_checkboxes:
                                for obs_cb in section.observation_checkboxes[pattern]:
                                    if obs_cb.text() == obs:
                                        obs_cb.setChecked(True)
                                        
                                        if obs in obs_data.get('images', {}):
                                            img_widget = section.pattern_image_widgets.get(pattern, {}).get(obs)
                                            if img_widget:
                                                for img_name in obs_data['images'][obs]:
                                                    img_path = os.path.join(images_folder, img_name)
                                                    if os.path.exists(img_path):
                                                        img_widget.add_image(img_path)
                        
                        if 'pattern_images' in obs_data:
                            if isinstance(section.pattern_image_widgets.get(pattern), dict):
                                img_widget = section.pattern_image_widgets[pattern].get('__pattern_level__')
                            else:
                                img_widget = section.pattern_image_widgets.get(pattern)
                            
                            if img_widget:
                                for img_name in obs_data['pattern_images']:
                                    img_path = os.path.join(images_folder, img_name)
                                    if os.path.exists(img_path):
                                        img_widget.add_image(img_path)
            
            for custom in section_info.get('custom_patterns', []):
                section.add_custom_pattern()
                if 'text' in custom and section.custom_patterns:
                    last_custom = section.custom_patterns[-1]
                    text_input = last_custom.get('text_input')
                    if text_input:
                        text_input.setText(custom['text'])
                    
                    if 'images' in custom and custom['images']:
                        img_widget = last_custom.get('image_widget')
                        if img_widget:
                            for img_name in custom['images']:
                                img_path = os.path.join(images_folder, img_name)
                                if os.path.exists(img_path):
                                    img_widget.add_image(img_path)
            
            if hasattr(section, 'notes_text'):
                section.notes_text.setPlainText(section_info.get('notes', ''))
            
            for img_name in section_info.get('images', []):
                img_path = os.path.join(images_folder, img_name)
                if os.path.exists(img_path):
                    section.image_widget.add_image(img_path)
    
    def _load_compressor_tab(self, compressor_tab, sections_data, images_folder):
        """Load compressor tab data"""
        section_map = {s.title: s for s, _, _ in compressor_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] == 'compressor_electrical':
                test_data = section_info.get('test_data', {})
                if hasattr(section, 'status_ok_cb'):
                    section.status_ok_cb.setChecked(test_data.get('ok', True))
                if hasattr(section, 'status_not_ok_cb'):
                    section.status_not_ok_cb.setChecked(not test_data.get('ok', True))
                
                for issue in test_data.get('issues', []):
                    if issue == "Grounded windings detected." and hasattr(section, 'grounded_cb'):
                        section.grounded_cb.setChecked(True)
                    elif issue == "Open windings detected." and hasattr(section, 'open_cb'):
                        section.open_cb.setChecked(True)
                    elif issue == "Resistance values out of acceptable range." and hasattr(section, 'resistance_cb'):
                        section.resistance_cb.setChecked(True)
                
                if hasattr(section, 'resistance_value'):
                    section.resistance_value.setText(test_data.get('resistance', ''))
                if hasattr(section, 'resistance_value_2'):
                    section.resistance_value_2.setText(test_data.get('resistance_2', ''))
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(test_data.get('notes', ''))
                
                for custom in test_data.get('custom_patterns', []):
                    if hasattr(section, 'add_custom_pattern'):
                        section.add_custom_pattern()
                        if 'text' in custom and section.custom_patterns:
                            last_custom = section.custom_patterns[-1]
                            text_input = last_custom.get('text_input')
                            if text_input:
                                text_input.setText(custom['text'])
                
                if 'ok_images' in test_data and hasattr(section, 'ok_image_widget'):
                    for img_name in test_data['ok_images']:
                        img_path = os.path.join(images_folder, img_name)
                        if os.path.exists(img_path):
                            section.ok_image_widget.add_image(img_path)
                
                if 'not_ok_images' in test_data and hasattr(section, 'not_ok_image_widget'):
                    for img_name in test_data['not_ok_images']:
                        img_path = os.path.join(images_folder, img_name)
                        if os.path.exists(img_path):
                            section.not_ok_image_widget.add_image(img_path)
                
                if 'general_images' in test_data and hasattr(section, 'general_image_widget'):
                    for img_name in test_data['general_images']:
                        img_path = os.path.join(images_folder, img_name)
                        if os.path.exists(img_path):
                            section.general_image_widget.add_image(img_path)
                
                if 'resistance_images' in test_data and hasattr(section, 'resistance_image_widget'):
                    for img_name in test_data['resistance_images']:
                        img_path = os.path.join(images_folder, img_name)
                        if os.path.exists(img_path):
                            section.resistance_image_widget.add_image(img_path)
            else:
                for pattern in section_info.get('selected_patterns', []):
                    if pattern in section.checkboxes:
                        section.checkboxes[pattern].setChecked(True)
                        
                        if pattern in section_info.get('observations', {}):
                            obs_data = section_info['observations'][pattern]
                            for obs in obs_data.get('selected', []):
                                if pattern in section.observation_checkboxes:
                                    for obs_cb in section.observation_checkboxes[pattern]:
                                        if obs_cb.text() == obs:
                                            obs_cb.setChecked(True)
                
                for custom in section_info.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(section_info.get('notes', ''))
                
                for img_name in section_info.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
    
    def _load_coil_tab(self, coil_tab, sections_data, images_folder):
        """Load coil tab data"""
        section_map = {s.title: s for s, _, _ in coil_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] == 'coil_pressure':
                pressure_data = section_info.get('pressure_data', {})
                if hasattr(section, 'psi_250_leak_cb'):
                    section.psi_250_leak_cb.setChecked(pressure_data.get('250psi_leak', False))
                if hasattr(section, 'psi_250_no_leak_cb'):
                    section.psi_250_no_leak_cb.setChecked(pressure_data.get('250psi_no_leak', False))
                if hasattr(section, 'psi_500_leak_cb'):
                    section.psi_500_leak_cb.setChecked(pressure_data.get('500psi_leak', False))
                if hasattr(section, 'psi_500_no_leak_cb'):
                    section.psi_500_no_leak_cb.setChecked(pressure_data.get('500psi_no_leak', False))
                if hasattr(section, 'psi_500_not_performed_cb'):
                    section.psi_500_not_performed_cb.setChecked(pressure_data.get('500psi_not_performed', False))
                
                # Restore 250 PSI leak location
                if hasattr(section, 'psi_250_leak_face_cb'):
                    section.psi_250_leak_face_cb.setChecked(pressure_data.get('250psi_leak_face', False))
                if hasattr(section, 'psi_250_leak_ubends_cb'):
                    section.psi_250_leak_ubends_cb.setChecked(pressure_data.get('250psi_leak_ubends', False))
                
                # Restore 500 PSI leak location
                if hasattr(section, 'psi_500_leak_face_cb'):
                    section.psi_500_leak_face_cb.setChecked(pressure_data.get('500psi_leak_face', False))
                if hasattr(section, 'psi_500_leak_ubends_cb'):
                    section.psi_500_leak_ubends_cb.setChecked(pressure_data.get('500psi_leak_ubends', False))
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(pressure_data.get('notes', ''))
                
                # Restore 250 PSI images
                for img_name in pressure_data.get('250psi_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.psi_250_image_widget.add_image(img_path)
                
                # Restore 500 PSI images
                for img_name in pressure_data.get('500psi_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.psi_500_image_widget.add_image(img_path)
                
                for img_name in pressure_data.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
            else:
                for pattern in section_info.get('selected_patterns', []):
                    if pattern in section.checkboxes:
                        section.checkboxes[pattern].setChecked(True)
                        
                        if pattern in section_info.get('observations', {}):
                            obs_data = section_info['observations'][pattern]
                            for obs in obs_data.get('selected', []):
                                if pattern in section.observation_checkboxes:
                                    for obs_cb in section.observation_checkboxes[pattern]:
                                        if obs_cb.text() == obs:
                                            obs_cb.setChecked(True)
                
                for custom in section_info.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(section_info.get('notes', ''))
                
                for img_name in section_info.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
    
    def _load_receiver_tab(self, receiver_tab, sections_data, images_folder):
        """Load receiver tab data"""
        section_map = {s.title: s for s, _, _ in receiver_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] in ('coil_pressure', 'receiver_pressure'):
                pressure_data = section_info.get('pressure_data', {})
                if hasattr(section, 'psi_250_leak_cb'):
                    section.psi_250_leak_cb.setChecked(pressure_data.get('250psi_leak', False))
                if hasattr(section, 'psi_250_no_leak_cb'):
                    section.psi_250_no_leak_cb.setChecked(pressure_data.get('250psi_no_leak', False))
                if hasattr(section, 'psi_500_leak_cb'):
                    section.psi_500_leak_cb.setChecked(pressure_data.get('500psi_leak', False))
                if hasattr(section, 'psi_500_no_leak_cb'):
                    section.psi_500_no_leak_cb.setChecked(pressure_data.get('500psi_no_leak', False))
                if hasattr(section, 'psi_500_not_performed_cb'):
                    section.psi_500_not_performed_cb.setChecked(pressure_data.get('500psi_not_performed', False))
                
                # Restore 250 PSI leak location
                if hasattr(section, 'psi_250_leak_face_cb'):
                    section.psi_250_leak_face_cb.setChecked(pressure_data.get('250psi_leak_face', False))
                if hasattr(section, 'psi_250_leak_ubends_cb'):
                    section.psi_250_leak_ubends_cb.setChecked(pressure_data.get('250psi_leak_ubends', False))
                
                # Restore 500 PSI leak location
                if hasattr(section, 'psi_500_leak_face_cb'):
                    section.psi_500_leak_face_cb.setChecked(pressure_data.get('500psi_leak_face', False))
                if hasattr(section, 'psi_500_leak_ubends_cb'):
                    section.psi_500_leak_ubends_cb.setChecked(pressure_data.get('500psi_leak_ubends', False))
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(pressure_data.get('notes', ''))
                
                # Restore 250 PSI images
                for img_name in pressure_data.get('250psi_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.psi_250_image_widget.add_image(img_path)
                
                # Restore 500 PSI images
                for img_name in pressure_data.get('500psi_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.psi_500_image_widget.add_image(img_path)
                
                for img_name in pressure_data.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
            else:
                for pattern in section_info.get('selected_patterns', []):
                    if pattern in section.checkboxes:
                        section.checkboxes[pattern].setChecked(True)
                        
                        if pattern in section_info.get('observations', {}):
                            obs_data = section_info['observations'][pattern]
                            for obs in obs_data.get('selected', []):
                                if pattern in section.observation_checkboxes:
                                    for obs_cb in section.observation_checkboxes[pattern]:
                                        if obs_cb.text() == obs:
                                            obs_cb.setChecked(True)
                            
                            for obs, img_names in obs_data.get('observation_images', {}).items():
                                if hasattr(section, 'pattern_image_widgets') and pattern in section.pattern_image_widgets:
                                    if isinstance(section.pattern_image_widgets[pattern], dict) and obs in section.pattern_image_widgets[pattern]:
                                        for img_name in img_names:
                                            img_path = os.path.join(images_folder, img_name)
                                            if os.path.exists(img_path):
                                                section.pattern_image_widgets[pattern][obs].add_image(img_path)
                            
                            for img_name in obs_data.get('pattern_images', []):
                                img_path = os.path.join(images_folder, img_name)
                                if os.path.exists(img_path):
                                    pattern_widget = section.pattern_image_widgets.get(pattern) if hasattr(section, 'pattern_image_widgets') else None
                                    if pattern_widget and not isinstance(pattern_widget, dict):
                                        pattern_widget.add_image(img_path)
                
                for custom in section_info.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(section_info.get('notes', ''))
                
                for img_name in section_info.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
    
    def _load_valve_tab(self, valve_tab, sections_data, images_folder):
        """Load valve tab data"""
        section_map = {s.title: s for s, _, _ in valve_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] == 'valve_electrical':
                test_data = section_info.get('test_data', {})
                if hasattr(section, 'status_ok_cb'):
                    section.status_ok_cb.setChecked(test_data.get('ok', False))
                if hasattr(section, 'status_not_ok_cb'):
                    section.status_not_ok_cb.setChecked(test_data.get('not_ok', False))
                if hasattr(section, 'resistance_value'):
                    section.resistance_value.setText(test_data.get('resistance', ''))
                
                for cb in section.issue_checkboxes:
                    if cb.text() in test_data.get('issues', []):
                        cb.setChecked(True)
                
                for issue, img_names in test_data.get('issue_images', {}).items():
                    if issue in section.issue_image_widgets:
                        for img_name in img_names:
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.issue_image_widgets[issue].add_image(img_path)
                
                for custom in test_data.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                for img_name in test_data.get('ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('not_ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.not_ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('general_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.general_image_widget.add_image(img_path)
                
                if hasattr(section, 'notes_text'):
                    section.notes_text.setPlainText(test_data.get('notes', ''))
                    
            elif section_info['type'] == 'valve_mechanical':
                test_data = section_info.get('test_data', {})
                if hasattr(section, 'status_ok_cb'):
                    section.status_ok_cb.setChecked(test_data.get('ok', False))
                if hasattr(section, 'status_not_ok_cb'):
                    section.status_not_ok_cb.setChecked(test_data.get('not_ok', False))
                if hasattr(section, 'bore_top_bottom'):
                    section.bore_top_bottom.setText(test_data.get('bore_top_bottom', ''))
                if hasattr(section, 'bore_left_right'):
                    section.bore_left_right.setText(test_data.get('bore_left_right', ''))
                if hasattr(section, 'bore_consistent_cb'):
                    section.bore_consistent_cb.setChecked(test_data.get('bore_consistent', False))
                
                # New valve movement checkboxes
                if hasattr(section, 'left_smooth_cb'):
                    section.left_smooth_cb.setChecked(test_data.get('left_smooth', False))
                if hasattr(section, 'left_not_smooth_cb'):
                    section.left_not_smooth_cb.setChecked(test_data.get('left_not_smooth', False))
                if hasattr(section, 'right_smooth_cb'):
                    section.right_smooth_cb.setChecked(test_data.get('right_smooth', False))
                if hasattr(section, 'right_not_smooth_cb'):
                    section.right_not_smooth_cb.setChecked(test_data.get('right_not_smooth', False))
                
                # Backward compatibility - convert old format to new
                if 'smooth_left' in test_data and not ('left_smooth' in test_data or 'left_not_smooth' in test_data):
                    if test_data.get('smooth_left') and hasattr(section, 'left_smooth_cb'):
                        section.left_smooth_cb.setChecked(True)
                if 'smooth_right' in test_data and not ('right_smooth' in test_data or 'right_not_smooth' in test_data):
                    if test_data.get('smooth_right') and hasattr(section, 'right_smooth_cb'):
                        section.right_smooth_cb.setChecked(True)
                
                for cb in section.issue_checkboxes:
                    if cb.text() in test_data.get('issues', []):
                        cb.setChecked(True)
                
                for issue, img_names in test_data.get('issue_images', {}).items():
                    if issue in section.issue_image_widgets:
                        for img_name in img_names:
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.issue_image_widgets[issue].add_image(img_path)
                
                for custom in test_data.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                for img_name in test_data.get('ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('not_ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.not_ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('general_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.general_image_widget.add_image(img_path)
                
                for img_name in test_data.get('cylinder_measurements_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path) and hasattr(section, 'cylinder_measurements_image_widget'):
                        section.cylinder_measurements_image_widget.add_image(img_path)
                
                if hasattr(section, 'notes_text'):
                    section.notes_text.setPlainText(test_data.get('notes', ''))
                    
            else:
                for pattern in section_info.get('selected_patterns', []):
                    if pattern in section.checkboxes:
                        section.checkboxes[pattern].setChecked(True)
                        
                        if pattern in section_info.get('observations', {}):
                            obs_data = section_info['observations'][pattern]
                            for obs in obs_data.get('selected', []):
                                if pattern in section.observation_checkboxes:
                                    for obs_cb in section.observation_checkboxes[pattern]:
                                        if obs_cb.text() == obs:
                                            obs_cb.setChecked(True)
                            
                            for obs, img_names in obs_data.get('observation_images', {}).items():
                                obs_imgs = section.get_observation_images(pattern, obs)
                                if hasattr(section, 'pattern_image_widgets') and pattern in section.pattern_image_widgets:
                                    if isinstance(section.pattern_image_widgets[pattern], dict) and obs in section.pattern_image_widgets[pattern]:
                                        for img_name in img_names:
                                            img_path = os.path.join(images_folder, img_name)
                                            if os.path.exists(img_path):
                                                section.pattern_image_widgets[pattern][obs].add_image(img_path)
                            
                            for img_name in obs_data.get('pattern_images', []):
                                img_path = os.path.join(images_folder, img_name)
                                if os.path.exists(img_path):
                                    pattern_widget = section.pattern_image_widgets.get(pattern)
                                    if pattern_widget and not isinstance(pattern_widget, dict):
                                        pattern_widget.add_image(img_path)
                
                for custom in section_info.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                if hasattr(section, 'notes_text'):
                    section.notes_text.setPlainText(section_info.get('notes', ''))
                
                for img_name in section_info.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
    
    def _load_pcb_tab(self, pcb_tab, sections_data, images_folder):
        """Load PCB tab data"""
        section_map = {s.title: s for s, _, _ in pcb_tab.get_all_sections()}
        
        for section_info in sections_data:
            section = section_map.get(section_info['title'])
            if not section:
                continue
            
            if section_info['type'] == 'pcb_electrical':
                test_data = section_info.get('test_data', {})
                if hasattr(section, 'status_ok_cb'):
                    section.status_ok_cb.setChecked(test_data.get('ok', False))
                if hasattr(section, 'status_not_ok_cb'):
                    section.status_not_ok_cb.setChecked(test_data.get('not_ok', False))
                if hasattr(section, 'korea_testing_cb'):
                    section.korea_testing_cb.setChecked(test_data.get('korea_testing', False))
                
                for cb in section.issue_checkboxes:
                    if cb.text() in test_data.get('issues', []):
                        cb.setChecked(True)
                
                for issue, img_names in test_data.get('issue_images', {}).items():
                    if issue in section.issue_image_widgets:
                        for img_name in img_names:
                            img_path = os.path.join(images_folder, img_name)
                            if os.path.exists(img_path):
                                section.issue_image_widgets[issue].add_image(img_path)
                
                for custom in test_data.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                for img_name in test_data.get('ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('not_ok_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.not_ok_image_widget.add_image(img_path)
                
                for img_name in test_data.get('general_images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.general_image_widget.add_image(img_path)
                
                if hasattr(section, 'notes_text'):
                    section.notes_text.setPlainText(test_data.get('notes', ''))
            
            else:
                # component type (visual inspection)
                for pattern in section_info.get('selected_patterns', []):
                    if pattern in section.checkboxes:
                        section.checkboxes[pattern].setChecked(True)
                        
                        if pattern in section_info.get('observations', {}):
                            obs_data = section_info['observations'][pattern]
                            for obs in obs_data.get('selected', []):
                                if pattern in section.observation_checkboxes:
                                    for obs_cb in section.observation_checkboxes[pattern]:
                                        if obs_cb.text() == obs:
                                            obs_cb.setChecked(True)
                            
                            for obs, img_names in obs_data.get('observation_images', {}).items():
                                if hasattr(section, 'pattern_image_widgets') and pattern in section.pattern_image_widgets:
                                    if isinstance(section.pattern_image_widgets[pattern], dict) and obs in section.pattern_image_widgets[pattern]:
                                        for img_name in img_names:
                                            img_path = os.path.join(images_folder, img_name)
                                            if os.path.exists(img_path):
                                                section.pattern_image_widgets[pattern][obs].add_image(img_path)
                            
                            for img_name in obs_data.get('pattern_images', []):
                                img_path = os.path.join(images_folder, img_name)
                                if os.path.exists(img_path):
                                    pattern_widget = section.pattern_image_widgets.get(pattern)
                                    if pattern_widget and not isinstance(pattern_widget, dict):
                                        pattern_widget.add_image(img_path)
                
                for custom in section_info.get('custom_patterns', []):
                    section.add_custom_pattern()
                    if 'text' in custom and section.custom_patterns:
                        last_custom = section.custom_patterns[-1]
                        text_input = last_custom.get('text_input')
                        if text_input:
                            text_input.setText(custom['text'])
                
                if hasattr(section, 'notes_field'):
                    section.notes_field.setPlainText(section_info.get('notes', ''))
                
                for img_name in section_info.get('images', []):
                    img_path = os.path.join(images_folder, img_name)
                    if os.path.exists(img_path):
                        section.image_widget.add_image(img_path)
                        
    def add_footer_with_page_numbers(self, doc, rma_number):
        """
        Add a footer to the document with:
        - RMA number on the left
        - Page number right
        - Disclaimer centered
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches
        from docx.enum.table import WD_ROW_HEIGHT_RULE

        # Apply margins and footer distance to ALL sections
        for section in doc.sections:
            section.top_margin = Inches(0.45)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            # Set bottom margin to accommodate footer
            section.bottom_margin = Inches(0.45)
            section.footer_distance = Inches(0.05)
            
            footer = section.footer
            
            # REMOVE all default footer paragraphs
            for p in list(footer._element.findall('.//w:p', footer._element.nsmap)):
                footer._element.remove(p)

            # Create a table in the footer for layout (3 columns)
            footer_table = footer.add_table(rows=1, cols=3, width=Inches(7.15))
            footer_table.autofit = False
            
            # Set row height to exactly 0.45"
            footer_table.rows[0].height = Inches(0.45)
            footer_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            cells = footer_table.rows[0].cells
            
            # Set cell widths using XML (Word-compatible method)
            # Total available width: 7.15" (8.5" - 0.75" left - 0.6" right margins)
            # Widths in twips (1 inch = 1440 twips)
            cell_widths_twips = [
                2160,  # Column 0: 1.5" for RMA
                5904,  # Column 1: 4.1" for disclaimer
                2160   # Column 2: 1.5" for page number
            ]
            
            for i, cell in enumerate(cells):
                tcPr = cell._element.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(cell_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

            # Left cell - RMA
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p0 = cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Remove paragraph from XML to ensure clean slate
            pPr0 = p0._element.get_or_add_pPr()
            # Remove any spacing elements
            for spacing in pPr0.findall(qn('w:spacing'), pPr0.nsmap):
                pPr0.remove(spacing)
            # Add explicit zero spacing
            spacing0 = OxmlElement('w:spacing')
            spacing0.set(qn('w:before'), '0')
            spacing0.set(qn('w:after'), '0')
            spacing0.set(qn('w:line'), '240')
            spacing0.set(qn('w:lineRule'), 'auto')
            pPr0.append(spacing0)
            
            run0 = p0.add_run(f"{rma_number}")
            run0.font.size = Pt(10)

            # Center cell - Disclaimer
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p1 = cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Remove paragraph spacing
            pPr1 = p1._element.get_or_add_pPr()
            for spacing in pPr1.findall(qn('w:spacing'), pPr1.nsmap):
                pPr1.remove(spacing)
            spacing1 = OxmlElement('w:spacing')
            spacing1.set(qn('w:before'), '0')
            spacing1.set(qn('w:after'), '0')
            spacing1.set(qn('w:line'), '240')
            spacing1.set(qn('w:lineRule'), 'auto')
            pPr1.append(spacing1)
            
            run1 = p1.add_run(
                ""
            )
            run1.font.size = Pt(8)

            # Right cell - Page number
            cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p2 = cells[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Critical: Remove spacing BEFORE adding page field
            pPr2 = p2._element.get_or_add_pPr()
            for spacing in pPr2.findall(qn('w:spacing'), pPr2.nsmap):
                pPr2.remove(spacing)
            spacing2 = OxmlElement('w:spacing')
            spacing2.set(qn('w:before'), '0')
            spacing2.set(qn('w:after'), '0')
            spacing2.set(qn('w:line'), '240')
            spacing2.set(qn('w:lineRule'), 'auto')
            pPr2.append(spacing2)

            # Add page number field
            run_page = p2.add_run()
            add_field(run_page, "PAGE")

            p2.add_run(" of ")

            run_total = p2.add_run()
            add_field(run_total, "NUMPAGES")
            
            # Create field character elements
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
        

            # Remove table borders
            for cell in cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
                
                # Remove cell margins/padding
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)

    def generate_report(self):
        if not self.equipment_id.text():
            QMessageBox.warning(self, "Missing Information", "Please enter RMA     ")
            return
        
        # If save file exists for current RMA, save to it before generating report
        if self.has_matching_save_file():
            try:
                self.perform_autosave()
            except Exception as e:
                # If save fails, log but continue with report generation
                print(f"Warning: Could not save before report generation: {e}")
        
        try:

            doc = Document()

            sections = doc.sections
            
            # Helper function to add pictures with black borders and smart sizing
            def add_bordered_picture(doc, image_path_or_stream, width=None, height=None):
                """
                Wrapper for doc.add_picture() that adds a 1pt black border and smart sizing.
                Automatically sizes images based on orientation to fit 2 images + captions per page.
                Vertical images stay vertical, horizontal images stay horizontal.
                Author: Daniel Fraser
                """
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                from PIL import Image as PILImage
                from PIL import ImageOps
                from io import BytesIO
                
                # ALWAYS apply EXIF orientation first, regardless of width/height parameters
                transposed_image = None
                try:
                    # Open image
                    if isinstance(image_path_or_stream, str):
                        img = PILImage.open(image_path_or_stream)
                    else:
                        img = PILImage.open(image_path_or_stream)
                        image_path_or_stream.seek(0)  # Reset stream position
                    
                    # Apply EXIF orientation to get true dimensions
                    # This handles images that are stored rotated with EXIF metadata
                    try:
                        img = ImageOps.exif_transpose(img)
                        # Save the transposed image so we insert the correctly oriented version
                        transposed_image = img
                    except:
                        pass  # If EXIF handling fails, continue with original image
                except:
                    pass  # If image can't be opened, continue without transpose
                
                # Smart sizing based on image orientation
                # Always calculate sizing from transposed image (ignore provided width/height)
                if transposed_image is not None:
                    try:
                        img_width, img_height = transposed_image.size
                        
                        # Determine orientation: x=width, y=height
                        # if x > y (width > height) → landscape/horizontal
                        # if x < y (width < height) → portrait/vertical
                        
                        if img_width > img_height:
                            # Landscape (horizontal): x > y
                            # Constrain by WIDTH first (the long dimension)
                            # Max dimensions: 5" width, 3.5" height
                            width = Inches(5.0)
                            height = None  # Explicitly clear height to maintain aspect ratio
                            aspect_ratio = img_height / img_width
                            proposed_height_inches = (width / Inches(1)) * aspect_ratio
                            
                            # If too tall, constrain by height instead
                            max_height_inches = 3.5
                            if proposed_height_inches > max_height_inches:
                                height = Inches(max_height_inches)
                                width = None
                        else:
                            # Portrait (vertical): x < y (or equal)
                            # Constrain by HEIGHT first (the long dimension)
                            # Max dimensions: 3.5" width, 4" height
                            height = Inches(4.0)
                            width = None  # Explicitly clear width to maintain aspect ratio
                            aspect_ratio = img_width / img_height
                            proposed_width_inches = (height / Inches(1)) * aspect_ratio
                            
                            # If too wide, constrain by width instead
                            max_width_inches = 3.5
                            if proposed_width_inches > max_width_inches:
                                width = Inches(max_width_inches)
                                height = None
                    except:
                        # If we can't read image, use default horizontal sizing
                        width = Inches(4.0)
                        height = None
                elif width is None and height is None:
                    # No transposed image and no sizing provided - use default
                    width = Inches(4.0)
                    height = None
                
                # If we have a transposed image, save it to BytesIO and use that
                # This ensures the correctly oriented image is inserted into the document
                if transposed_image is not None:
                    img_stream = BytesIO()
                    # Always save as PNG - no EXIF complications, lossless
                    # Convert RGBA to RGB if needed for compatibility
                    if transposed_image.mode in ('RGBA', 'LA', 'P'):
                        # Create white background
                        background = PILImage.new('RGB', transposed_image.size, (255, 255, 255))
                        if transposed_image.mode == 'P':
                            transposed_image = transposed_image.convert('RGBA')
                        background.paste(transposed_image, mask=transposed_image.split()[-1] if transposed_image.mode in ('RGBA', 'LA') else None)
                        transposed_image = background
                    elif transposed_image.mode != 'RGB':
                        transposed_image = transposed_image.convert('RGB')
                    
                    transposed_image.save(img_stream, format='PNG')
                    img_stream.seek(0)
                    picture = doc.add_picture(img_stream, width=width, height=height)
                else:
                    # Use original image path/stream
                    picture = doc.add_picture(image_path_or_stream, width=width, height=height)
                
                # Add black border to the picture
                try:
                    # 1pt = 12,700 EMUs (English Metric Units)
                    width_emu = 12700
                    
                    # Access the picture XML element
                    pic = picture._inline.graphic.graphicData.pic
                    
                    # Get or create shape properties (spPr)
                    spPr = pic.spPr
                    if spPr is None:
                        spPr = OxmlElement('pic:spPr')
                        pic.insert(0, spPr)
                    
                    # Get or create line element (ln) for border
                    ln = spPr.find(qn('a:ln'))
                    if ln is None:
                        ln = OxmlElement('a:ln')
                        spPr.append(ln)
                    
                    # Set border width
                    ln.set('w', str(width_emu))
                    ln.set('cap', 'flat')  # Add cap property for Word compatibility
                    
                    # Get or create solid fill
                    solidFill = ln.find(qn('a:solidFill'))
                    if solidFill is None:
                        solidFill = OxmlElement('a:solidFill')
                        ln.append(solidFill)
                    
                    # Get or create RGB color element
                    srgbClr = solidFill.find(qn('a:srgbClr'))
                    if srgbClr is None:
                        srgbClr = OxmlElement('a:srgbClr')
                        solidFill.append(srgbClr)
                    
                    # Set color to black (000000)
                    srgbClr.set('val', '000000')
                    
                    # Add miter join for sharp corners
                    miter = ln.find(qn('a:miter'))
                    if miter is None:
                        miter = OxmlElement('a:miter')
                        ln.append(miter)
                except:
                    # If border fails, picture is still added without border
                    pass
                
                return picture
            
            def add_images_smart_layout(doc, image_paths, figure_counter, section_name=""):
                """
                Intelligently layout images - consecutive portrait photos side by side, landscape separate.
                Returns the updated figure counter.
                """
                from PIL import Image as PILImage
                from PIL import ImageOps
                from io import BytesIO
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                
                if not image_paths:
                    return figure_counter
                
                # Check orientation of all images
                image_info = []
                for img_path in image_paths:
                    try:
                        img = PILImage.open(img_path)
                        img = ImageOps.exif_transpose(img)
                        width, height = img.size
                        is_portrait = height > width
                        image_info.append({'path': img_path, 'is_portrait': is_portrait})
                    except:
                        image_info.append({'path': img_path, 'is_portrait': False})
                
                # Group consecutive portraits together, keep landscapes separate
                i = 0
                while i < len(image_info):
                    current_img = image_info[i]
                    
                    if current_img['is_portrait']:
                        # Look ahead to group consecutive portraits (max 2 per row)
                        portrait_group = [current_img]
                        j = i + 1
                        while j < len(image_info) and image_info[j]['is_portrait'] and len(portrait_group) < 2:
                            portrait_group.append(image_info[j])
                            j += 1
                        
                        if len(portrait_group) >= 2:
                            # Create side-by-side layout for portrait pair
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            for idx, img_info in enumerate(portrait_group):
                                try:
                                    img_path = img_info['path']
                                    img = PILImage.open(img_path)
                                    img = ImageOps.exif_transpose(img)
                                    
                                    # Portrait sizing: constrain height to 3.5" for side-by-side fit
                                    height = Inches(3.5)
                                    
                                    # Save transposed image to BytesIO
                                    img_stream = BytesIO()
                                    if img.mode in ('RGBA', 'LA', 'P'):
                                        background = PILImage.new('RGB', img.size, (255, 255, 255))
                                        if img.mode == 'P':
                                            img = img.convert('RGBA')
                                        background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                                        img = background
                                    elif img.mode != 'RGB':
                                        img = img.convert('RGB')
                                    
                                    img.save(img_stream, format='PNG')
                                    img_stream.seek(0)
                                    
                                    # Add picture to paragraph run
                                    run = para.add_run()
                                    picture = run.add_picture(img_stream, height=height)
                                    
                                    # Add 1pt black border
                                    try:
                                        pic = picture._inline.graphic.graphicData.pic
                                        spPr = pic.spPr
                                        if spPr is None:
                                            spPr = OxmlElement('pic:spPr')
                                            pic.insert(0, spPr)
                                        ln = spPr.find(qn('a:ln'))
                                        if ln is None:
                                            ln = OxmlElement('a:ln')
                                            spPr.append(ln)
                                        ln.set('w', '12700')
                                        ln.set('cap', 'flat')  # Add cap property for Word compatibility
                                        
                                        # Add solid fill
                                        solidFill = ln.find(qn('a:solidFill'))
                                        if solidFill is None:
                                            solidFill = OxmlElement('a:solidFill')
                                            ln.append(solidFill)
                                        srgbClr = solidFill.find(qn('a:srgbClr'))
                                        if srgbClr is None:
                                            srgbClr = OxmlElement('a:srgbClr')
                                            solidFill.append(srgbClr)
                                        srgbClr.set('val', '000000')
                                        
                                        # Add miter join for sharp corners
                                        miter = ln.find(qn('a:miter'))
                                        if miter is None:
                                            miter = OxmlElement('a:miter')
                                            ln.append(miter)
                                    except:
                                        pass
                                    
                                    # Add space between images
                                    if idx < len(portrait_group) - 1:
                                        run.add_text('    ')
                                        
                                except Exception as e:
                                    run = para.add_run()
                                    run.add_text(f"[Error: {str(e)}]")
                            
                            # Caption for side-by-side portraits
                            if section_name:
                                caption_text = f"Figure {figure_counter}: {section_name}"
                            else:
                                # Use filenames if no section name
                                filenames = [os.path.basename(img['path']) for img in portrait_group]
                                caption_text = f"Figure {figure_counter}: {', '.join(filenames)}"
                            
                            fig_para = doc.add_paragraph(caption_text)
                            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fig_para.paragraph_format.space_after = Pt(1)
                            figure_counter += 1
                            spacing_para = doc.add_paragraph()
                            spacing_para.paragraph_format.line_spacing = Pt(0)  # Small controlled gap
                            
                            # Move past all grouped portraits
                            i = j
                        else:
                            # Single portrait - use standard layout
                            img_path = current_img['path']
                            try:
                                if os.path.exists(img_path):
                                    add_bordered_picture(doc, img_path)
                                    last_paragraph = doc.paragraphs[-1]
                                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    # Use section_name if provided, otherwise filename
                                    caption = f"Figure {figure_counter}: {section_name}" if section_name else f"Figure {figure_counter}: {os.path.basename(img_path)}"
                                    fig_para = doc.add_paragraph(caption)
                                    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    fig_para.paragraph_format.space_after = Pt(1)
                                    figure_counter += 1
                                    spacing_para = doc.add_paragraph()
                                    spacing_para.paragraph_format.line_spacing = Pt(0)  # Small controlled gap
                            except Exception as e:
                                doc.add_paragraph(f"Error loading image: {str(e)}")
                            i += 1
                    else:
                        # Landscape image - always separate
                        img_path = current_img['path']
                        try:
                            if os.path.exists(img_path):
                                add_bordered_picture(doc, img_path)
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Use section_name if provided, otherwise filename
                                caption = f"Figure {figure_counter}: {section_name}" if section_name else f"Figure {figure_counter}: {os.path.basename(img_path)}"
                                fig_para = doc.add_paragraph(caption)
                                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                fig_para.paragraph_format.space_after = Pt(1)
                                figure_counter += 1
                                spacing_para = doc.add_paragraph()
                                spacing_para.paragraph_format.line_spacing = Pt(0)  # Small controlled gap
                        except Exception as e:
                            doc.add_paragraph(f"Error loading image: {str(e)}")
                        i += 1
                
                return figure_counter
            
            quality_lab = doc.add_heading('Quality Laboratory', 0)
            quality_lab.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title = doc.add_heading('Inspection Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            date_title = doc.add_heading(self.inspection_date.text(), 0)
            date_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            

            doc.add_heading('Inspection Details', level=1)
            
            header_data = [
                ('RMA:', self.equipment_id.text()),
                ('Unit Model:', self.unit_model.text()),
                ('Unit Serial:', self.unit_serial.text()),
                ('Part Number:', self.part_number.text()),
                ('Warranty Claim Number:', self.warranty_claim.text()),
                ('Customer and Email:', self.customer_email.text()),
                ('Total:', self.total_field.text()),
                ('Install Date:', self.install_date.text()),
                ('Failure Date:', self.failure_date.text()),
                ('Lab Technician:', self.lab_tech.text()),
                ('Manager:', self.manager_name.text()),
            ]
            
            header_table = doc.add_table(rows=0, cols=2)
            header_table.style = 'Light Grid Accent 1'
            header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set column widths to prevent long text from distorting layout
            # Left column (labels): 2.5 inches - fixed width (no wrapping)
            # Right column (values): 3.625 inches - allows wrapping
            header_table.columns[0].width = Inches(2.5)
            header_table.columns[1].width = Inches(4.625)

            for label, value in header_data:
                row = header_table.add_row().cells
                row[0].text = label
                row[1].text = value
                
                # Ensure text wrapping is enabled in the value cell
                # This prevents long strings from expanding the table
                row[1].paragraphs[0].paragraph_format.word_wrap = True
            
            doc.add_paragraph()
            
            # Add 4 blank lines before disclaimer
            for _ in range(4):
                doc.add_paragraph()
            
            # Add disclaimer (bold header, then regular text)
            disclaimer_heading_para = doc.add_paragraph()
            disclaimer_heading_run = disclaimer_heading_para.add_run("Disclaimer")
            disclaimer_heading_run.font.bold = True
            disclaimer_heading_run.font.size = Pt(11)
            
            doc.add_paragraph()  # blank line between heading and body
            
            disclaimer_text = ("This report has been prepared for internal review purposes only and reflects the findings of a limited technical "
                "evaluation conducted on the returned equipment. The inspection and testing performed were based solely on the condition "
                "of the equipment as received and under the specific test conditions applied at the time of examination. Nothing in this "
                "report creates a guarantee of performance, warranty determination, or final root cause analysis. Findings are subject to "
                "revision if additional information, testing, or analysis becomes available. If this report is provided to a customer, it "
                "is intended for informational purposes only and should not be relied upon as a comprehensive assessment of system-wide "
                "performance, installation conditions, or operational practices beyond the scope of this review.")
            
            disclaimer_body_para = doc.add_paragraph()
            disclaimer_body_run = disclaimer_body_para.add_run(disclaimer_text)
            disclaimer_body_run.font.bold = False
            disclaimer_body_run.font.size = Pt(9)
            
            doc.add_page_break()
            
            figure_counter = 1
            motor_counter = 1
            compressor_counter = 1
            coil_counter = 1
            receiver_counter = 1
            valve_counter = 1
            pcb_counter = 1
            general_counter = 1
            first_equipment = True
            
            for tab_idx in range(self.equipment_tabs_widget.count()):
                widget = self.equipment_tabs_widget.widget(tab_idx)
                
                if isinstance(widget, MotorTab):

                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'Motor {motor_counter}', level=1)
                    motor_counter += 1
                    

                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        

                        if section_type == 'test':
                            doc.add_heading('Test Results', level=2)
                            test_data = section.get_test_data()
                            

                            doc.add_heading('Audio Testing', level=3)
                            if test_data['audio']['numeric'] is not None:
                                if test_data['audio']['ok']:
                                    doc.add_paragraph(f"Audio testing measured {test_data['audio']['value']} which is within acceptable operating limits (≤72 dB).")
                                else:
                                    doc.add_paragraph(f"Audio testing measured {test_data['audio']['value']} which exceeds acceptable operating limits (>72 dB).")
                            elif test_data['audio']['value']:
                                doc.add_paragraph(f"Audio testing result: {test_data['audio']['value']}")
                            else:
                                doc.add_paragraph("Audio testing not performed.")
                            
                            if test_data['audio']['images']:
                                figure_counter = add_images_smart_layout(doc, test_data['audio']['images'], figure_counter, "Audio Testing")
                            

                            doc.add_heading('Vibration Testing', level=3)
                            if test_data['vibration']['rpm_numeric'] and test_data['vibration']['numeric']:
                                classification = test_data['vibration']['classification']
                                rpm_val = test_data['vibration']['rpm']
                                vib_val = test_data['vibration']['value']
                                
                                if test_data['vibration']['ok']:
                                    doc.add_paragraph(f"Vibration testing at {rpm_val} RPM measured {vib_val} mm/s, classified as '{classification}', which is within acceptable operating limits.")
                                else:
                                    doc.add_paragraph(f"Vibration testing at {rpm_val} RPM measured {vib_val} mm/s, classified as '{classification}', which exceeds acceptable operating limits.")
                            elif test_data['vibration']['value']:
                                parts = []
                                if test_data['vibration']['rpm']:
                                    parts.append(f"RPM: {test_data['vibration']['rpm']}")
                                if test_data['vibration']['value']:
                                    parts.append(f"Vibration: {test_data['vibration']['value']}")
                                if parts:
                                    doc.add_paragraph(f"Vibration testing result: {', '.join(parts)}")
                            else:
                                doc.add_paragraph("Vibration testing not performed.")
                            
                            if test_data['vibration']['images']:
                                figure_counter = add_images_smart_layout(doc, test_data['vibration']['images'], figure_counter, "Vibration Testing")
                            

                            doc.add_heading('Temperature Testing', level=3)
                            if test_data['temperature']['numeric'] is not None:
                                if test_data['temperature']['ok']:
                                    doc.add_paragraph(f"Temperature testing measured {test_data['temperature']['value']} which is within acceptable operating limits (≤70°F).")
                                else:
                                    doc.add_paragraph(f"Temperature testing measured {test_data['temperature']['value']} which exceeds acceptable operating limits (>70°F).")
                            elif test_data['temperature']['value']:
                                doc.add_paragraph(f"Temperature testing result: {test_data['temperature']['value']}")
                            else:
                                doc.add_paragraph("Temperature testing not performed.")
                            
                            if test_data['temperature']['images']:
                                figure_counter = add_images_smart_layout(doc, test_data['temperature']['images'], figure_counter, "Temperature Testing")
                            
                            doc.add_heading('Electrical Resistance Testing', level=3)
                            status = "OK" if test_data['resistance']['ok'] else "NOT OK"
                            
                            resistance_parts = []
                            black_red = test_data['resistance'].get('black_red')
                            black_white = test_data['resistance'].get('black_white')
                            black_yellow = test_data['resistance'].get('black_yellow', '')
                            black_blue = test_data['resistance'].get('black_blue', '')
                            is_five_wire = test_data['resistance'].get('is_five_wire', False)
                            
                            # Build list of measurements
                            measurements = []
                            if black_red:
                                measurements.append(f"Black/Red at {black_red}Ω")
                            if black_white:
                                measurements.append(f"Black/White at {black_white}Ω")
                            if is_five_wire and black_yellow:
                                measurements.append(f"Black/Yellow at {black_yellow}Ω")
                            if is_five_wire and black_blue:
                                measurements.append(f"Black/Blue at {black_blue}Ω")
                            
                            if measurements:
                                measurements_text = ", ".join(measurements)
                                if test_data['resistance']['ok']:
                                    doc.add_paragraph(f"Electrical resistance testing measured {measurements_text}, which are within acceptable operating limits.")
                                else:
                                    doc.add_paragraph(f"Electrical resistance testing measured {measurements_text}, which are outside acceptable operating limits.")
                            else:
                                if test_data['resistance']['ok']:
                                    doc.add_paragraph("Electrical resistance testing found connections to be within acceptable operating limits.")
                                else:
                                    doc.add_paragraph("Electrical resistance testing found connections to be outside acceptable operating limits.")
                            
                            if test_data['resistance']['images']:
                                figure_counter = add_images_smart_layout(doc, test_data['resistance']['images'], figure_counter, "Electrical Resistance Testing")
                            
                            continue
                    
                        
                        doc.add_heading(section.title, level=2)
                        

                        selected_patterns = section.get_selected_patterns()
                        
                        if selected_patterns:
                            is_special_section = section.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B", "Electrical and Mechanical Inspection"]
                        

                            if is_special_section and selected_patterns == ["Status: OK"]:
                                # Just add the OK statement, no "Observed Wear Patterns" header
                                selected_obs = section.get_selected_observations("Status: OK")
                                if selected_obs and len(selected_obs) > 0:
                                    doc.add_paragraph(selected_obs[0])
                                
                                    pattern_images = section.get_pattern_images("Status: OK")
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                                continue
                        

                            doc.add_heading('Observed Wear Patterns:', level=3)
                        
                            for pattern in selected_patterns:
                                if is_special_section and pattern == "Status: OK":
                                    continue
                            
                                # Skip heading for "No issues detected" pattern
                                if pattern == "No issues detected":
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        para = doc.add_paragraph(selected_obs[0])
                                        para.paragraph_format.space_after = Pt(12)
                                    
                                    # Add images for "No issues detected" pattern
                                    pattern_images = section.get_pattern_images(pattern)
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                                    
                                    continue

                                if not (is_special_section and pattern == "Status: NOT OK"):
                                    doc.add_heading(pattern, level=4)
                                
                                    # Add example image FIRST (at the top, right after heading) for bearing sections
                                    is_bearing = section.title in ["Drive End Bearing", "Non-Drive End Bearing"]
                                    if is_bearing and pattern in BEARING_REFERENCE_IMAGES:
                                        ref_image_name = BEARING_REFERENCE_IMAGES[pattern]
                                        if getattr(sys, 'frozen', False):
                                            ref_image_path = os.path.join(sys._MEIPASS, '_assets', ref_image_name)
                                        else:
                                            ref_image_path = os.path.join(os.path.dirname(__file__), '_assets', ref_image_name)
                                        
                                        if os.path.exists(ref_image_path):
                                            try:
                                                # Add example image
                                                add_bordered_picture(doc, ref_image_path, width=Inches(4.5))
                                                last_paragraph = doc.paragraphs[-1]
                                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                
                                                # Add caption with bold red formatting - EXAMPLE in all caps
                                                fig_para = doc.add_paragraph()
                                                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                fig_para.paragraph_format.space_after = Pt(0)
                                                run = fig_para.add_run(f"Figure {figure_counter}: EXAMPLE - {pattern}")
                                                run.bold = True
                                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                                figure_counter += 1
                                                
                                                # Add source with bold red formatting
                                                para = doc.add_paragraph()
                                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                source_run = para.add_run("(Source: SKF, Bearing Damage and Failure Analysis.)")
                                                source_run.bold = True
                                                source_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                            except Exception as e:
                                                pass
                            

                                selected_obs = section.get_selected_observations(pattern)
                            
                                if selected_obs:
                                    if is_special_section and pattern == "Status: NOT OK":

                                        grouped_obs = []
                                        current_group = []
                                        
                                        for obs in selected_obs:
                                            obs_images = section.get_observation_images(pattern, obs)
                                            if obs_images:

                                                if current_group:
                                                    grouped_obs.append(("text", current_group))
                                                    current_group = []
                                                grouped_obs.append(("image", obs, obs_images))
                                            else:
                                                current_group.append(obs)
                                        
                                        # Add custom issues to the grouping
                                        custom_patterns_temp = section.get_custom_patterns()
                                        for custom in custom_patterns_temp:
                                            if custom.get('text'):
                                                if custom.get('images'):
                                                    # Custom issue with images - flush current group first
                                                    if current_group:
                                                        grouped_obs.append(("text", current_group))
                                                        current_group = []
                                                    grouped_obs.append(("image", custom['text'], custom['images']))
                                                else:
                                                    # Custom issue without images - add to current group
                                                    current_group.append(custom['text'])
                                        
                                        if current_group:
                                            grouped_obs.append(("text", current_group))
                                        
                                        for item in grouped_obs:
                                            if item[0] == "text":
                                                observations_text = " ".join(item[1])
                                                para = doc.add_paragraph(observations_text)
                                                para.paragraph_format.space_after = Pt(0)
                                            else:
                                                obs = item[1]
                                                obs_images = item[2]
                                                para = doc.add_paragraph(obs)
                                                para.paragraph_format.space_after = Pt(12)
                                                figure_counter = add_images_smart_layout(doc, obs_images, figure_counter, obs)
                                    else:

                                        observations_text = " ".join(selected_obs)
                                        para = doc.add_paragraph(observations_text)
                                        para.paragraph_format.space_after = Pt(12)
                                    
                                        pattern_images = section.get_pattern_images(pattern)
                                        if pattern_images:
                                            figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                                else:
                                    custom_patterns_preview = section.get_custom_patterns()
                                    if not custom_patterns_preview:
                                        para = doc.add_paragraph("Wear pattern observed (no specific details selected).")
                                        para.paragraph_format.space_after = Pt(12)
                                    
                                    if not is_special_section:
                                        pattern_images = section.get_pattern_images(pattern)
                                        if pattern_images:
                                            figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                                
                                if is_special_section and pattern == "Status: NOT OK":
                                    pattern_images = section.get_pattern_images(pattern)
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                        
                        # Process custom patterns for bearing sections
                        # (Motor sections with selected patterns already handled in integrated grouping above)
                        is_motor_special = section.title in ["Motor Housing", "Motor Shaft", "Electrical Connection"]
                        
                        if not (is_motor_special and selected_patterns):
                            custom_patterns = section.get_custom_patterns()
                            if custom_patterns:
                                if not selected_patterns:
                                    doc.add_heading('Observed Wear Patterns:', level=3)
                                
                                # Group custom patterns together like regular observations
                                custom_texts = []
                                for custom in custom_patterns:
                                    if 'text' in custom:
                                        if custom['text']:
                                            custom_texts.append(custom['text'])
                                    elif 'name' in custom:
                                        # Legacy format support
                                        if 'observations' in custom and custom['observations']:
                                            observations_text = " ".join(custom['observations'])
                                            custom_texts.append(observations_text)
                                
                                # Add all custom texts as a single paragraph
                                if custom_texts:
                                    combined_text = " ".join(custom_texts)
                                    para = doc.add_paragraph(combined_text)
                                    para.paragraph_format.space_after = Pt(12)
                                
                                # Handle images for custom patterns (if any)
                                for custom in custom_patterns:
                                    if custom.get('images'):
                                        figure_counter = add_images_smart_layout(doc, custom['images'], figure_counter, "Custom Observation")
                        
                        if not selected_patterns:
                            has_custom = bool(section.get_custom_patterns())
                            if not has_custom:
                                # For bearing sections, say "not tested"
                                if section.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                                    doc.add_paragraph(f"{section.title}: Not tested.")
                                else:
                                    doc.add_paragraph("No significant wear patterns observed.")
                        

                        notes = section.get_notes()
                        if notes:
                            doc.add_heading('Additional Notes:', level=3)
                            doc.add_paragraph(notes)
                        

                        images = section.get_images()
                        if images:
                            doc.add_heading('General Section Images:', level=3)
                            figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                    
                    doc.add_heading('Motor Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=3)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set table width to full page width
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    # Set table width to 100% of page width (5000 = 100%)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)
                    for col in summary_table.columns:
                        col.width = Inches(2.375)
                    

                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Visual Inspection'
                    header_cells[1].text = 'Electrical and Mechanical Inspection'
                    header_cells[2].text = 'Internal Inspection'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    
                    # Visual Inspection (Motor Housing, Motor Shaft, Electrical Connection) - separate by section
                    visual_sections = {}
                    visual_status = {}  # Track status patterns separately
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title in ["Motor Housing", "Motor Shaft", "Electrical Connection"]:
                            section_items = []
                            status_items = []
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        status_items.extend(selected_obs)
                                    else:
                                        status_items.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        section_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    section_items.append(custom['text'])
                            
                            if section_items:
                                visual_sections[section.title] = section_items
                            if status_items:
                                visual_status[section.title] = status_items
                    
                    # Build visual inspection text with bold labels and issue counts
                    body_cells[0].text = ''
                    first_section = True
                    
                    for section_name in ["Motor Housing", "Motor Shaft", "Electrical Connection"]:
                        # Show sections with actual issues
                        if section_name in visual_sections:
                            if not first_section:
                                para = body_cells[0].add_paragraph()
                            else:
                                para = body_cells[0].paragraphs[0]
                            first_section = False
                            
                            items = visual_sections[section_name]
                            count = len(items)
                            
                            # Bold label with count in parentheses and colon
                            issue_word = "Issue" if count == 1 else "Issues"
                            run = para.add_run(f"{section_name} ({count} {issue_word}): ")
                            run.font.bold = True
                            
                            # Normal text for issues
                            run = para.add_run(' '.join(items))
                            run.font.bold = False
                        
                        # Show sections with only status (no issues)
                        elif section_name in visual_status:
                            if not first_section:
                                para = body_cells[0].add_paragraph()
                            else:
                                para = body_cells[0].paragraphs[0]
                            first_section = False
                            
                            status = visual_status[section_name]
                            
                            # Bold label without count
                            run = para.add_run(f"{section_name}: ")
                            run.font.bold = True
                            
                            # Normal text for status
                            run = para.add_run(' '.join(status))
                            run.font.bold = False
                        
                        # Show sections with nothing selected (fallback - not counted as issue)
                        else:
                            if not first_section:
                                para = body_cells[0].add_paragraph()
                            else:
                                para = body_cells[0].paragraphs[0]
                            first_section = False
                            
                            # Bold label without count
                            run = para.add_run(f"{section_name}: ")
                            run.font.bold = True
                            
                            # Normal text for fallback
                            run = para.add_run("No data recorded")
                            run.font.bold = False
                    
                    # Electrical and Mechanical Inspection (Test Results)
                    body_cells[1].text = ''
                    first_item = True
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'test':
                            test_data = section.get_test_data()

                            audio_data = test_data.get('audio', {})
                            if audio_data.get('numeric') is not None:
                                if not first_item:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_item = False
                                
                                run = para.add_run("Audio: ")
                                run.font.bold = True
                                if audio_data.get('ok'):
                                    run = para.add_run(f"{audio_data.get('value')} (within limits)")
                                    run.font.bold = False
                                else:
                                    run = para.add_run(f"{audio_data.get('value')} (exceeds limits)")
                                    run.font.bold = False
                            elif audio_data.get('value'):
                                if not first_item:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_item = False
                                
                                run = para.add_run("Audio: ")
                                run.font.bold = True
                                run = para.add_run(f"{audio_data.get('value')}")
                                run.font.bold = False
                            
                            vib_data = test_data.get('vibration', {})
                            if vib_data.get('rpm_numeric') and vib_data.get('numeric'):
                                if not first_item:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_item = False
                                
                                run = para.add_run("Vibration: ")
                                run.font.bold = True
                                if vib_data.get('ok'):
                                    run = para.add_run(f"{vib_data.get('classification')} (within acceptable limits)")
                                    run.font.bold = False
                                else:
                                    run = para.add_run(f"{vib_data.get('classification')} (exceeds limits)")
                                    run.font.bold = False

                            temp_data = test_data.get('temperature', {})
                            if temp_data.get('value'):
                                if not first_item:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_item = False
                                
                                run = para.add_run("Temperature: ")
                                run.font.bold = True
                                if temp_data.get('ok'):
                                    run = para.add_run(f"{temp_data.get('value')} (within limits)")
                                    run.font.bold = False
                                else:
                                    run = para.add_run(f"{temp_data.get('value')} (exceeds limits)")
                                    run.font.bold = False

                            res_data = test_data.get('resistance', {})
                            black_red = res_data.get('black_red')
                            black_white = res_data.get('black_white')
                            if black_red or black_white:
                                if not first_item:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_item = False
                                
                                run = para.add_run("Resistance: ")
                                run.font.bold = True
                                res_text = f"Black-Red: {black_red}, Black-White: {black_white}"
                                if not res_data.get('ok'):
                                    res_text += " (NOT OK)"
                                run = para.add_run(res_text)
                                run.font.bold = False
                    
                    if first_item:
                        body_cells[1].text = 'No test results recorded.'
                    
                    # Internal Inspection (Bearings)
                    body_cells[2].text = ''
                    first_bearing = True
                    bearings_shown = []
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                            selected_patterns = section.get_selected_patterns()
                            section_items = []
                            status_items = []
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        status_items.extend(selected_obs)
                                    else:
                                        status_items.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        section_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    section_items.append(custom['text'])
                            
                            bearing_label = "Drive End" if section.title == "Drive End Bearing" else "Non-Drive End"
                            bearings_shown.append(bearing_label)
                            
                            # Show section with actual issues
                            if section_items:
                                if not first_bearing:
                                    para = body_cells[2].add_paragraph()
                                else:
                                    para = body_cells[2].paragraphs[0]
                                first_bearing = False
                                
                                count = len(section_items)
                                
                                # Bold label with count in parentheses and colon
                                issue_word = "Issue" if count == 1 else "Issues"
                                run = para.add_run(f"{bearing_label} ({count} {issue_word}): ")
                                run.font.bold = True
                                
                                # Normal text for issues
                                run = para.add_run(' '.join(section_items))
                                run.font.bold = False
                            
                            # Show section with only status (no issues)
                            elif status_items:
                                if not first_bearing:
                                    para = body_cells[2].add_paragraph()
                                else:
                                    para = body_cells[2].paragraphs[0]
                                first_bearing = False
                                
                                # Bold label without count
                                run = para.add_run(f"{bearing_label}: ")
                                run.font.bold = True
                                
                                # Normal text for status
                                run = para.add_run(' '.join(status_items))
                                run.font.bold = False
                            
                            # Show section with nothing selected (fallback - not counted as issue)
                            else:
                                if not first_bearing:
                                    para = body_cells[2].add_paragraph()
                                else:
                                    para = body_cells[2].paragraphs[0]
                                first_bearing = False
                                
                                # Bold label without count
                                run = para.add_run(f"{bearing_label}: ")
                                run.font.bold = True
                                
                                # Normal text for fallback
                                run = para.add_run("No data recorded")
                                run.font.bold = False
                    
                    doc.add_paragraph()

                elif isinstance(widget, CompressorTab):

                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'Compressor {compressor_counter}', level=1)
                    compressor_counter += 1
                    

                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        

                        if section_type == 'compressor_electrical':
                            doc.add_heading('Electrical Testing', level=2)
                            test_data = section.get_test_data()
                            
                            if test_data['ok']:
                                doc.add_paragraph("Electrical testing completed with no issues detected. No ground detected on housing, windings are intact. Resistance values are within operating range.")
                                
                                # Status: OK images
                                if test_data.get('ok_images'):
                                    figure_counter = add_images_smart_layout(doc, test_data['ok_images'], figure_counter, f"Electrical Testing")
                            elif test_data['not_ok']:
                                if test_data['issues'] or test_data['custom_patterns']:
                                    doc.add_heading('Observed Issues:', level=3)
                                    

                                    all_issues = []
                                    if test_data['issues']:
                                        all_issues.extend(test_data['issues'])
                                    if test_data['custom_patterns']:
                                        for custom in test_data['custom_patterns']:
                                            if custom['text']:
                                                all_issues.append(custom['text'])
                                    
                                    if all_issues:
                                        issues_text = " ".join(all_issues)
                                        doc.add_paragraph(issues_text)
                                else:
                                    doc.add_paragraph("Electrical testing revealed issues but specifics were not provided.")
                                
                                # NOT OK images (for all issues) - moved outside the issues check
                                if test_data['not_ok_images']:
                                    figure_counter = add_images_smart_layout(doc, test_data['not_ok_images'], figure_counter, f"Electrical Testing")
                            else:
                                # Neither OK nor NOT OK selected
                                doc.add_paragraph("Electrical testing not performed.")
                            
                            # Resistance Testing section
                            if test_data['resistance'] or test_data.get('resistance_2'):
                                doc.add_heading('Resistance Testing:', level=3)
                                
                                # Resistance values
                                if test_data['resistance']:
                                    doc.add_paragraph(f"Resistance 1 measured: {test_data['resistance']}Ω")
                                
                                if test_data.get('resistance_2'):
                                    doc.add_paragraph(f"Resistance 2 measured: {test_data['resistance_2']}Ω")
                                
                                # Resistance images
                                if test_data.get('resistance_images'):
                                    figure_counter = add_images_smart_layout(doc, test_data['resistance_images'], figure_counter, "Resistance Testing")
                            

                            if test_data.get('general_images'):
                                figure_counter = add_images_smart_layout(doc, test_data['general_images'], figure_counter, "Images")
                            
                            if test_data['notes']:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(test_data['notes'])
                            
                            continue
                        

                        doc.add_heading(section.title, level=2)
                        
                        selected_patterns = section.get_selected_patterns()
                        
                        if selected_patterns:
                            is_special_section = section.title in ["Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection"]
                            

                            if is_special_section and selected_patterns == ["Status: OK"]:
                                selected_obs = section.get_selected_observations("Status: OK")
                                if selected_obs and len(selected_obs) > 0:
                                    doc.add_paragraph(selected_obs[0])
                                    
                                    pattern_images = section.get_pattern_images("Status: OK")
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                                continue
                            

                            doc.add_heading('Observed Issues:', level=3)
                            
                            for pattern in selected_patterns:
                                if is_special_section and pattern == "Status: OK":
                                    continue
                                
                                # Don't add pattern name as heading for Status: NOT OK
                                if not (is_special_section and pattern == "Status: NOT OK"):
                                    doc.add_heading(pattern, level=4)
                                
                                selected_obs = section.get_selected_observations(pattern)
                                
                                if selected_obs:
                                    if is_special_section and pattern == "Status: NOT OK":
                                        # Group observations intelligently for special sections
                                        grouped_obs = []
                                        current_group = []
                                        
                                        for obs in selected_obs:
                                            obs_images = section.get_observation_images(pattern, obs)
                                            if obs_images:
                                                if current_group:
                                                    grouped_obs.append(("text", current_group))
                                                    current_group = []
                                                grouped_obs.append(("image", obs, obs_images))
                                            else:
                                                current_group.append(obs)
                                        
                                        if current_group:
                                            grouped_obs.append(("text", current_group))
                                        
                                        for item in grouped_obs:
                                            if item[0] == "text":
                                                observations_text = " ".join(item[1])
                                                para = doc.add_paragraph(observations_text)
                                                para.paragraph_format.space_after = Pt(0)
                                            else:
                                                obs = item[1]
                                                obs_images = item[2]
                                                para = doc.add_paragraph(obs)
                                                para.paragraph_format.space_after = Pt(12)
                                                figure_counter = add_images_smart_layout(doc, obs_images, figure_counter, obs)
                                    else:

                                        observations_text = " ".join(selected_obs)
                                        para = doc.add_paragraph(observations_text)
                                        para.paragraph_format.space_after = Pt(12)
                                else:
                                    # No observations selected for Status: NOT OK
                                    if is_special_section and pattern == "Status: NOT OK":
                                        is_compressor = section.title in ["Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection"]
                                        is_valve = section.title in ["External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]
                                        custom_patterns = section.get_custom_patterns()
                                        has_custom = any(custom.get('text') for custom in custom_patterns)
                                        
                                        if not has_custom:
                                            if is_compressor:
                                                doc.add_paragraph(f"{section.title} revealed issues but specifics were not provided.")
                                            elif is_valve:
                                                doc.add_paragraph(f"{section.title} revealed issues but specifics were not provided.")
                                            else:
                                                # For Visual Inspection (coil)
                                                doc.add_paragraph("Visual inspection revealed issues but specifics were not provided.")
                                
                                # Add pattern-level images for Status: NOT OK (whether or not observations were selected)
                                if is_special_section and pattern == "Status: NOT OK":
                                    pattern_images = section.get_pattern_images(pattern)
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                        


                        is_compressor_section = section.title in ["Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "Visual Inspection"]
                        is_motor_section = section.title in ["Motor Housing", "Motor Shaft", "Electrical Connection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B", "Electrical and Mechanical Inspection"]
                        
                        # Only process custom patterns here if NOT a motor section with selected patterns
                        # (motor sections with selected patterns already handled in integrated grouping above)
                        if not (is_motor_section and selected_patterns):
                            custom_patterns = section.get_custom_patterns()
                            
                            if custom_patterns:
                                if is_compressor_section and selected_patterns:

                                    custom_texts = [custom['text'] for custom in custom_patterns if custom.get('text')]
                                    if custom_texts:
                                        last_para = None
                                        for para in reversed(doc.paragraphs):
                                            if para.text and para.style.name.startswith('Heading'):
                                                continue
                                            last_para = para
                                            break
                                        
                                        if last_para:

                                            last_para.add_run(" " + " ".join(custom_texts))
                                        else:
                                            para = doc.add_paragraph(" ".join(custom_texts))
                                            para.paragraph_format.space_after = Pt(12)
                                elif not selected_patterns:
                                    doc.add_heading('Observed Issues:', level=3)
                                    
                                    for custom in custom_patterns:
                                        if custom['text']:
                                            para = doc.add_paragraph(custom['text'])
                                            para.paragraph_format.space_after = Pt(0)
                                            
                                            if 'images' in custom and custom['images']:
                                                figure_counter = add_images_smart_layout(doc, custom['images'], figure_counter, "Images")
                        
                        # For motor sections with selected patterns, custom patterns already integrated above
                        if not selected_patterns:
                            has_custom = bool(section.get_custom_patterns())
                            if not has_custom:
                                # For compressor and valve sections, say "testing not performed"
                                if section.title in ["Oil Evaluation", "Sleeve Bearing Inspection", "Scroll Plate Inspection", "External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]:
                                    doc.add_paragraph(f"{section.title} not performed.")
                                else:
                                    doc.add_paragraph("No significant issues observed.")
                        
                        notes = section.get_notes()
                        if notes:
                            doc.add_heading('Additional Notes:', level=3)
                            doc.add_paragraph(notes)
                        

                        images = section.get_images()
                        if images:
                            doc.add_heading('General Section Images:', level=3)
                            figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                    
                    # Add Compressor Summary Table at end of compressor section
                    doc.add_heading('Compressor Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=4)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set table width to full page width
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    # Set table width to 100% of page width (5000 = 100%)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)

                    for col in summary_table.columns:
                        col.width = Inches(1.775)
                    

                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Electrical Testing'
                    header_cells[1].text = 'Oil Evaluation'
                    header_cells[2].text = 'Scroll Plate Inspection'
                    header_cells[3].text = 'Sleeve Bearing Inspection'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    

                    # Compressor Electrical Testing
                    body_cells[0].text = ''
                    first_item = True
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'compressor_electrical':
                            test_data = section.get_test_data()
                            if test_data.get('ok'):
                                if not first_item:
                                    para0 = body_cells[0].add_paragraph()
                                else:
                                    para0 = body_cells[0].paragraphs[0]
                                first_item = False
                                
                                run = para0.add_run("No issues detected.")
                                run.font.bold = False
                            else:
                                issues_list = []
                                if test_data.get('issues'):
                                    issues_list.extend(test_data['issues'])
                                if test_data.get('custom_patterns'):
                                    for custom in test_data['custom_patterns']:
                                        if custom.get('text'):
                                            issues_list.append(custom['text'])
                                
                                if issues_list:
                                    if not first_item:
                                        para0 = body_cells[0].add_paragraph()
                                    else:
                                        para0 = body_cells[0].paragraphs[0]
                                    first_item = False
                                    
                                    count = len(issues_list)
                                    run = para0.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                                    run.font.bold = True
                                    run = para0.add_run(' '.join(issues_list))
                                    run.font.bold = False
                            
                            if test_data.get('resistance'):
                                para0 = body_cells[0].add_paragraph()
                                run = para0.add_run("Resistance 1: ")
                                run.font.bold = True
                                run = para0.add_run(f"{test_data['resistance']}Ω")
                                run.font.bold = False
                            
                            if test_data.get('resistance_2'):
                                para0 = body_cells[0].add_paragraph()
                                run = para0.add_run("Resistance 2: ")
                                run.font.bold = True
                                run = para0.add_run(f"{test_data['resistance_2']}Ω")
                                run.font.bold = False
                            
                            if test_data.get('notes'):
                                para0 = body_cells[0].add_paragraph()
                                run = para0.add_run(test_data['notes'])
                                run.font.bold = False
                    
                    if first_item:
                        body_cells[0].text = 'No data recorded.'
                    

                    # Oil Evaluation
                    body_cells[1].text = ''
                    para1 = body_cells[1].paragraphs[0]
                    oil_items = []
                    oil_status = []
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Oil Evaluation":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        oil_status.extend(selected_obs)
                                    else:
                                        oil_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        oil_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    oil_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                oil_items.append(notes)
                    
                    if oil_items:
                        count = len(oil_items)
                        run = para1.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para1.add_run(' '.join(oil_items))
                        run.font.bold = False
                    elif oil_status:
                        run = para1.add_run(' '.join(oil_status))
                        run.font.bold = False
                    else:
                        body_cells[1].text = 'No issues detected.'
                    
                    # Scroll Plate Inspection
                    body_cells[2].text = ''
                    para_scroll = body_cells[2].paragraphs[0]
                    scroll_items = []
                    scroll_status = []
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Scroll Plate Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        scroll_status.extend(selected_obs)
                                    else:
                                        scroll_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        scroll_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    scroll_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                scroll_items.append(notes)
                    
                    if scroll_items:
                        count = len(scroll_items)
                        run = para_scroll.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para_scroll.add_run(' '.join(scroll_items))
                        run.font.bold = False
                    elif scroll_status:
                        run = para_scroll.add_run(' '.join(scroll_status))
                        run.font.bold = False
                    else:
                        body_cells[2].text = 'No issues detected.'
                    
                    # Sleeve Bearing Inspection
                    body_cells[3].text = ''
                    para_sleeve = body_cells[3].paragraphs[0]
                    sleeve_items = []
                    sleeve_status = []
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Sleeve Bearing Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        sleeve_status.extend(selected_obs)
                                    else:
                                        sleeve_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        sleeve_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    sleeve_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                sleeve_items.append(notes)
                    
                    if sleeve_items:
                        count = len(sleeve_items)
                        run = para_sleeve.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para_sleeve.add_run(' '.join(sleeve_items))
                        run.font.bold = False
                    elif sleeve_status:
                        run = para_sleeve.add_run(' '.join(sleeve_status))
                        run.font.bold = False
                    else:
                        body_cells[3].text = 'No issues detected.'
                    
                    doc.add_paragraph()
                
                elif isinstance(widget, CoilTab):

                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'Coil {coil_counter}', level=1)
                    coil_counter += 1
                    
                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type in ('coil_pressure', 'receiver_pressure'):
                            # Pressure Testing Section
                            doc.add_heading('Pressure Testing', level=2)
                            pressure_data = section.get_pressure_data()
                            
                            # 250 PSI Test
                            psi_250_tested = pressure_data['250psi_leak'] or pressure_data['250psi_no_leak']
                            
                            if psi_250_tested:
                                if pressure_data['250psi_leak']:
                                    leak_text = "250 PSI pressure test detected a leak"
                                    # Add location if specified
                                    locations = []
                                    if pressure_data.get('250psi_leak_face'):
                                        locations.append("face")
                                    if pressure_data.get('250psi_leak_ubends'):
                                        locations.append("U-bends")
                                    
                                    if locations:
                                        leak_text += f" on {' and '.join(locations)}"
                                    else:
                                        leak_text += " (location not specified)"
                                    
                                    doc.add_paragraph(leak_text + ".")
                                elif pressure_data['250psi_no_leak']:
                                    doc.add_paragraph("250 PSI pressure test showed no leak. Testing will be done with 500 PSI to verify.")
                            else:
                                doc.add_paragraph("250 PSI pressure testing was not performed.")
                            
                            # 250 PSI Images
                            if pressure_data.get('250psi_images'):
                                figure_counter = add_images_smart_layout(doc, pressure_data['250psi_images'], figure_counter, "250 PSI Pressure Testing")
                            
                            # 500 PSI Test
                            if pressure_data['500psi_leak']:
                                leak_text = "500 PSI pressure test detected a leak"
                                # Add location if specified
                                locations = []
                                if pressure_data.get('500psi_leak_face'):
                                    locations.append("face")
                                if pressure_data.get('500psi_leak_ubends'):
                                    locations.append("U-bends")
                                
                                if locations:
                                    leak_text += f" on {' and '.join(locations)}"
                                else:
                                    leak_text += " (location not specified)"
                                
                                doc.add_paragraph(leak_text + ".")
                            elif pressure_data['500psi_no_leak']:
                                doc.add_paragraph("500 PSI pressure test showed no leak.")
                            else:
                                # If "Not Performed" checked OR nothing selected, use the same logic
                                # Check if it was unnecessary (leak in 250 PSI) or just not performed
                                if psi_250_tested and pressure_data['250psi_leak']:
                                    doc.add_paragraph("500 PSI pressure testing was not necessary.")
                                else:
                                    doc.add_paragraph("500 PSI pressure testing was not performed.")
                            
                            # 500 PSI Images
                            if pressure_data.get('500psi_images'):
                                figure_counter = add_images_smart_layout(doc, pressure_data['500psi_images'], figure_counter, "500 PSI Pressure Testing")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                        
                        elif section_type == 'component':
                            # Visual Inspection Section
                            doc.add_heading(section.title, level=2)
                            
                            selected_patterns = section.get_selected_patterns()
                            
                            if "Status: OK" in selected_patterns:
                                observations = section.get_selected_observations("Status: OK")
                                for obs in observations:
                                    doc.add_paragraph(obs)
                            
                            if "Status: NOT OK" in selected_patterns:
                                selected_obs = section.get_selected_observations("Status: NOT OK")
                                
                                if selected_obs:
                                    grouped_obs = []
                                    for obs in selected_obs:
                                        grouped_obs.append(obs)
                                    
                                    if grouped_obs:
                                        para_text = " ".join(grouped_obs)
                                        para = doc.add_paragraph(para_text)
                                        para.paragraph_format.space_after = Pt(12)
                                
                                custom_patterns = section.get_custom_patterns()
                                if custom_patterns and selected_obs:
                                    custom_texts = [custom['text'] for custom in custom_patterns if custom.get('text')]
                                    if custom_texts:
                                        last_para = None
                                        for para in reversed(doc.paragraphs):
                                            if para.text and para.style.name.startswith('Heading'):
                                                continue
                                            last_para = para
                                            break
                                        
                                        if last_para:
                                            last_para.add_run(" " + " ".join(custom_texts))
                                        else:
                                            para = doc.add_paragraph(" ".join(custom_texts))
                                            para.paragraph_format.space_after = Pt(12)
                                elif custom_patterns and not selected_obs:
                                    doc.add_heading('Observed Issues:', level=3)
                                    for custom in custom_patterns:
                                        if custom.get('text'):
                                            doc.add_paragraph(custom['text'], style='List Bullet')
                                
                                # Pattern-level images for NOT OK
                                pattern_images = section.get_pattern_images("Status: NOT OK")
                                if pattern_images:
                                    figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                            
                            # Fallback message if no patterns selected at all
                            if not selected_patterns:
                                doc.add_paragraph("Visual inspection not performed.")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                    
                    # Add Coil Summary Table
                    doc.add_heading('Coil Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=2)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set table width to full page width
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    # Set table width to 100% of page width (5000 = 100%)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)

                    for col in summary_table.columns:
                        col.width = Inches(3.575)
                    
                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Visual Inspection'
                    header_cells[1].text = 'Pressure Testing'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    
                    visual_items = []
                    visual_status = []
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Visual Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_status.extend(selected_obs)
                                    else:
                                        visual_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    visual_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                visual_items.append(notes)
                    
                    # Build visual inspection text with issue count
                    body_cells[0].text = ''
                    if visual_items:
                        para = body_cells[0].paragraphs[0]
                        count = len(visual_items)
                        run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para.add_run(' '.join(visual_items))
                        run.font.bold = False
                    elif visual_status:
                        para = body_cells[0].paragraphs[0]
                        run = para.add_run(' '.join(visual_status))
                        run.font.bold = False
                    else:
                        body_cells[0].text = 'No issues detected.'
                    
                    # Build pressure testing text with bold labels
                    body_cells[1].text = ''
                    first_test = True
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type in ('coil_pressure', 'receiver_pressure'):
                            pressure_data = section.get_pressure_data()
                            psi_250_tested = pressure_data['250psi_leak'] or pressure_data['250psi_no_leak']
                            psi_500_tested = pressure_data['500psi_leak'] or pressure_data['500psi_no_leak']
                            
                            if psi_250_tested:
                                if not first_test:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_test = False
                                
                                run = para.add_run("250 PSI: ")
                                run.font.bold = True
                                
                                if pressure_data['250psi_leak']:
                                    leak_locations = []
                                    if pressure_data.get('250psi_leak_face'):
                                        leak_locations.append("face")
                                    if pressure_data.get('250psi_leak_ubends'):
                                        leak_locations.append("U-bends")
                                    
                                    if leak_locations:
                                        location_text = " and ".join(leak_locations)
                                        run = para.add_run(f"Leak detected on {location_text}")
                                        run.font.bold = False
                                    else:
                                        run = para.add_run("Leak detected")
                                        run.font.bold = False
                                elif pressure_data['250psi_no_leak']:
                                    run = para.add_run("No leak")
                                    run.font.bold = False
                            
                            if psi_500_tested:
                                if not first_test:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_test = False
                                
                                run = para.add_run("500 PSI: ")
                                run.font.bold = True
                                
                                if pressure_data['500psi_leak']:
                                    leak_locations = []
                                    if pressure_data.get('500psi_leak_face'):
                                        leak_locations.append("face")
                                    if pressure_data.get('500psi_leak_ubends'):
                                        leak_locations.append("U-bends")
                                    
                                    if leak_locations:
                                        location_text = " and ".join(leak_locations)
                                        run = para.add_run(f"Leak detected on {location_text}")
                                        run.font.bold = False
                                    else:
                                        run = para.add_run("Leak detected")
                                        run.font.bold = False
                                elif pressure_data['500psi_no_leak']:
                                    run = para.add_run("No leak")
                                    run.font.bold = False
                            
                            notes = section.get_notes()
                            if notes:
                                para = body_cells[1].add_paragraph()
                                run = para.add_run(notes)
                                run.font.bold = False
                    
                    if first_test:
                        body_cells[1].text = 'No data recorded.'
                    
                    doc.add_paragraph()
                
                
                elif isinstance(widget, ReceiverTab):
                    # Receiver Report Generation
                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'Receiver {receiver_counter}', level=1)
                    receiver_counter += 1
                    
                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type == 'receiver_pressure':
                            # Pressure Testing Section
                            doc.add_heading('Pressure Testing', level=2)
                            pressure_data = section.get_pressure_data()
                            
                            # 250 PSI Test
                            psi_250_tested = pressure_data['250psi_leak'] or pressure_data['250psi_no_leak']
                            
                            if psi_250_tested:
                                if pressure_data['250psi_leak']:
                                    doc.add_paragraph("250 PSI pressure test detected a leak.")
                                elif pressure_data['250psi_no_leak']:
                                    doc.add_paragraph("250 PSI pressure test showed no leak. Testing will be done with 500 PSI to verify.")
                            else:
                                doc.add_paragraph("250 PSI pressure testing was not performed.")
                            
                            # 250 PSI Images
                            if pressure_data.get('250psi_images'):
                                figure_counter = add_images_smart_layout(doc, pressure_data['250psi_images'], figure_counter, "250 PSI Pressure Testing")
                            
                            # 500 PSI Test
                            if pressure_data['500psi_leak']:
                                doc.add_paragraph("500 PSI pressure test detected a leak.")
                            elif pressure_data['500psi_no_leak']:
                                doc.add_paragraph("500 PSI pressure test showed no leak.")
                            else:
                                # If "Not Performed" checked OR nothing selected, use the same logic
                                # Check if it was unnecessary (leak in 250 PSI) or just not performed
                                if psi_250_tested and pressure_data['250psi_leak']:
                                    doc.add_paragraph("500 PSI pressure testing was not necessary.")
                                else:
                                    doc.add_paragraph("500 PSI pressure testing was not performed.")
                            
                            # 500 PSI Images
                            if pressure_data.get('500psi_images'):
                                figure_counter = add_images_smart_layout(doc, pressure_data['500psi_images'], figure_counter, "500 PSI Pressure Testing")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                        
                        elif section_type == 'component':
                            # Visual Inspection Section
                            doc.add_heading(section.title, level=2)
                            
                            selected_patterns = section.get_selected_patterns()
                            
                            if "Status: OK" in selected_patterns:
                                observations = section.get_selected_observations("Status: OK")
                                for obs in observations:
                                    doc.add_paragraph(obs)
                                # Add Status: OK images
                                ok_images = section.get_pattern_images("Status: OK")
                                if ok_images:
                                    figure_counter = add_images_smart_layout(doc, ok_images, figure_counter, section.title)
                            
                            if "Status: NOT OK" in selected_patterns:
                                selected_obs = section.get_selected_observations("Status: NOT OK")
                                
                                if selected_obs:
                                    grouped_obs = []
                                    for obs in selected_obs:
                                        grouped_obs.append(obs)
                                    
                                    if grouped_obs:
                                        para_text = " ".join(grouped_obs)
                                        para = doc.add_paragraph(para_text)
                                        para.paragraph_format.space_after = Pt(12)
                                
                                custom_patterns = section.get_custom_patterns()
                                if custom_patterns and selected_obs:
                                    custom_texts = [custom['text'] for custom in custom_patterns if custom.get('text')]
                                    if custom_texts:
                                        last_para = None
                                        for para in reversed(doc.paragraphs):
                                            if para.text and para.style.name.startswith('Heading'):
                                                continue
                                            last_para = para
                                            break
                                        
                                        if last_para:
                                            last_para.add_run(" " + " ".join(custom_texts))
                                        else:
                                            para = doc.add_paragraph(" ".join(custom_texts))
                                            para.paragraph_format.space_after = Pt(12)
                                elif custom_patterns and not selected_obs:
                                    doc.add_heading('Observed Issues:', level=3)
                                    for custom in custom_patterns:
                                        if custom.get('text'):
                                            doc.add_paragraph(custom['text'], style='List Bullet')
                                
                                pattern_images = section.get_pattern_images("Status: NOT OK")
                                if pattern_images:
                                    figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                            
                            # Fallback message if no patterns selected at all
                            if not selected_patterns:
                                doc.add_paragraph("Visual inspection not performed.")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                    
                    # Add Receiver Summary Table
                    doc.add_heading('Receiver Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=2)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set table width to full page width
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    # Set table width to 100% of page width (5000 = 100%)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)

                    for col in summary_table.columns:
                        col.width = Inches(3.575)
                    
                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Visual Inspection'
                    header_cells[1].text = 'Pressure Testing'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    
                    visual_items = []
                    visual_status = []
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Visual Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_status.extend(selected_obs)
                                    else:
                                        visual_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    visual_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                visual_items.append(notes)
                    
                    # Build visual inspection text with issue count
                    body_cells[0].text = ''
                    if visual_items:
                        para = body_cells[0].paragraphs[0]
                        count = len(visual_items)
                        run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para.add_run(' '.join(visual_items))
                        run.font.bold = False
                    elif visual_status:
                        para = body_cells[0].paragraphs[0]
                        run = para.add_run(' '.join(visual_status))
                        run.font.bold = False
                    else:
                        body_cells[0].text = 'No issues detected.'
                    
                    # Build pressure testing text with bold labels
                    body_cells[1].text = ''
                    first_test = True
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'receiver_pressure':
                            pressure_data = section.get_pressure_data()
                            psi_250_tested = pressure_data['250psi_leak'] or pressure_data['250psi_no_leak']
                            psi_500_tested = pressure_data['500psi_leak'] or pressure_data['500psi_no_leak']
                            
                            if psi_250_tested:
                                if not first_test:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_test = False
                                
                                run = para.add_run("250 PSI: ")
                                run.font.bold = True
                                
                                if pressure_data['250psi_leak']:
                                    run = para.add_run("Leak detected")
                                    run.font.bold = False
                                elif pressure_data['250psi_no_leak']:
                                    run = para.add_run("No leak")
                                    run.font.bold = False
                            
                            if psi_500_tested:
                                if not first_test:
                                    para = body_cells[1].add_paragraph()
                                else:
                                    para = body_cells[1].paragraphs[0]
                                first_test = False
                                
                                run = para.add_run("500 PSI: ")
                                run.font.bold = True
                                
                                if pressure_data['500psi_leak']:
                                    run = para.add_run("Leak detected")
                                    run.font.bold = False
                                elif pressure_data['500psi_no_leak']:
                                    run = para.add_run("No leak")
                                    run.font.bold = False
                            
                            notes = section.get_notes()
                            if notes:
                                para = body_cells[1].add_paragraph()
                                run = para.add_run(notes)
                                run.font.bold = False
                    
                    if first_test:
                        body_cells[1].text = 'No data recorded.'
                    
                    doc.add_paragraph()
                
                elif isinstance(widget, ValveTab):
                    # Valve Report Generation
                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'Four-Way Valve {valve_counter}', level=1)
                    valve_counter += 1
                    
                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type == 'valve_electrical':
                            doc.add_heading('Electrical Inspection', level=2)
                            test_data = section.get_test_data()
                            
                            if test_data['ok']:
                                doc.add_paragraph("Solenoid coil electrical inspection completed with no issues detected.")
                                
                                figure_counter = add_images_smart_layout(doc, test_data['ok_images'], figure_counter, f"Electrical Inspection")
                            
                            elif test_data['not_ok']:
                                issues = test_data.get('issues', [])
                                issue_images = test_data.get('issue_images', {})
                                custom_patterns = test_data.get('custom_patterns', [])
                                
                                if issues or custom_patterns:
                                    # Group issues intelligently
                                    grouped_items = []
                                    current_group = []
                                    
                                    # Process regular issues
                                    for issue in issues:
                                        if issue in issue_images and issue_images[issue]:
                                            # Issue with images - flush current group first
                                            if current_group:
                                                grouped_items.append(("text", current_group))
                                                current_group = []
                                            grouped_items.append(("image", issue, issue_images[issue]))
                                        else:
                                            # Issue without images - add to current group
                                            current_group.append(issue)
                                    
                                    # Process custom patterns
                                    for custom in custom_patterns:
                                        if custom.get('text'):
                                            if custom.get('images'):
                                                # Custom with images - flush current group first
                                                if current_group:
                                                    grouped_items.append(("text", current_group))
                                                    current_group = []
                                                grouped_items.append(("image", custom['text'], custom['images']))
                                            else:
                                                # Custom without images - add to current group
                                                current_group.append(custom['text'])
                                    
                                    # Add any remaining grouped text
                                    if current_group:
                                        grouped_items.append(("text", current_group))
                                    
                                    # Output grouped items
                                    for item in grouped_items:
                                        if item[0] == "text":
                                            observations_text = " ".join(item[1])
                                            para = doc.add_paragraph(observations_text)
                                            para.paragraph_format.space_after = Pt(0)
                                        else:
                                            issue_text = item[1]
                                            issue_imgs = item[2]
                                            para = doc.add_paragraph(issue_text)
                                            para.paragraph_format.space_after = Pt(12)
                                            figure_counter = add_images_smart_layout(doc, issue_imgs, figure_counter, issue_text)
                                
                                # Pattern-level NOT OK images
                                figure_counter = add_images_smart_layout(doc, test_data.get('not_ok_images', []), figure_counter, "Electrical Inspection")
                                
                                if not issues and not custom_patterns:
                                    doc.add_paragraph("Electrical inspection revealed issues but specifics were not provided.")
                            else:
                                doc.add_paragraph("Electrical inspection not performed.")
                            
                            if test_data.get('resistance'):
                                doc.add_paragraph(f"Solenoid Coil Resistance: {test_data['resistance']} Ω")
                            
                            if test_data.get('notes'):
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(test_data['notes'])
                            
                            figure_counter = add_images_smart_layout(doc, test_data.get('general_images', []), figure_counter, "Images")
                        
                        elif section_type == 'valve_mechanical':
                            doc.add_heading('Mechanical Inspection', level=2)
                            test_data = section.get_test_data()
                            
                            has_bore_data = test_data.get('bore_top_bottom') or test_data.get('bore_left_right')
                            has_movement_data = (test_data.get('left_smooth') or test_data.get('left_not_smooth') or 
                                                 test_data.get('right_smooth') or test_data.get('right_not_smooth'))
                            has_status = test_data['ok'] or test_data['not_ok']
                            has_notes = bool(test_data.get('notes'))
                            has_images = bool(test_data.get('general_images', []))
                            
                            if not (has_bore_data or has_movement_data or has_status or has_notes or has_images):
                                doc.add_paragraph("Mechanical inspection not performed.")
                            else:
                                bore_top_bottom = test_data.get('bore_top_bottom')
                                bore_left_right = test_data.get('bore_left_right')
                                
                                if bore_top_bottom or bore_left_right:
                                    doc.add_heading('Cylinder Bore Measurements:', level=3)
                                    
                                    if bore_top_bottom and bore_left_right:
                                        doc.add_paragraph(f"Top to Bottom: {bore_top_bottom} mm, Left to Right: {bore_left_right} mm")
                                        
                                        if test_data.get('bore_consistent'):
                                            doc.add_paragraph("Bore measurements are consistent across all measurements.")
                                    elif bore_top_bottom:
                                        doc.add_paragraph(f"Top to Bottom: {bore_top_bottom} mm (Left to Right measurement not taken)")
                                    else:
                                        doc.add_paragraph(f"Left to Right: {bore_left_right} mm (Top to Bottom measurement not taken)")
                                    
                                    # Add cylinder measurements images if present
                                    cylinder_meas_images = test_data.get('cylinder_measurements_images', [])
                                    if cylinder_meas_images:
                                        figure_counter = add_images_smart_layout(doc, cylinder_meas_images, figure_counter, "Cylinder Measurements")
                                
                                # New valve movement logic
                                left_smooth = test_data.get('left_smooth')
                                left_not_smooth = test_data.get('left_not_smooth')
                                right_smooth = test_data.get('right_smooth')
                                right_not_smooth = test_data.get('right_not_smooth')
                                
                                has_left_data = left_smooth or left_not_smooth
                                has_right_data = right_smooth or right_not_smooth
                                
                                if has_left_data or has_right_data:
                                    doc.add_heading('Valve Movement:', level=3)
                                    
                                    if has_left_data:
                                        if left_smooth:
                                            doc.add_paragraph("Left valve movement: Valve moves smoothly to the left.")
                                        else:  # left_not_smooth
                                            doc.add_paragraph("Left valve movement: Valve does not move smoothly to the left.")
                                    
                                    if has_right_data:
                                        if right_smooth:
                                            doc.add_paragraph("Right valve movement: Valve moves smoothly to the right.")
                                        else:  # right_not_smooth
                                            doc.add_paragraph("Right valve movement: Valve does not move smoothly to the right.")
                                    
                                    if not has_left_data:
                                        doc.add_paragraph("Left valve movement: Not tested.")
                                    if not has_right_data:
                                        doc.add_paragraph("Right valve movement: Not tested.")
                                elif has_bore_data or has_status:
                                    doc.add_heading('Valve Movement:', level=3)
                                    doc.add_paragraph("Left valve movement: Not tested.")
                                    doc.add_paragraph("Right valve movement: Not tested.")
                                
                                if test_data['ok']:
                                    doc.add_paragraph("Valve mechanical inspection completed with smooth operation confirmed.")
                                    
                                    figure_counter = add_images_smart_layout(doc, test_data['ok_images'], figure_counter, "Mechanical Inspection")
                                
                                elif test_data['not_ok']:
                                    issues = test_data.get('issues', [])
                                    issue_images = test_data.get('issue_images', {})
                                    custom_patterns = test_data.get('custom_patterns', [])
                                    
                                    if issues or custom_patterns:
                                        # Group issues intelligently
                                        grouped_items = []
                                        current_group = []
                                        
                                        # Process regular issues
                                        for issue in issues:
                                            if issue in issue_images and issue_images[issue]:
                                                # Issue with images - flush current group first
                                                if current_group:
                                                    grouped_items.append(("text", current_group))
                                                    current_group = []
                                                grouped_items.append(("image", issue, issue_images[issue]))
                                            else:
                                                # Issue without images - add to current group
                                                current_group.append(issue)
                                        
                                        # Process custom patterns
                                        for custom in custom_patterns:
                                            if custom.get('text'):
                                                if custom.get('images'):
                                                    # Custom with images - flush current group first
                                                    if current_group:
                                                        grouped_items.append(("text", current_group))
                                                        current_group = []
                                                    grouped_items.append(("image", custom['text'], custom['images']))
                                                else:
                                                    # Custom without images - add to current group
                                                    current_group.append(custom['text'])
                                        
                                        # Add any remaining grouped text
                                        if current_group:
                                            grouped_items.append(("text", current_group))
                                        
                                        # Output grouped items
                                        for item in grouped_items:
                                            if item[0] == "text":
                                                observations_text = " ".join(item[1])
                                                para = doc.add_paragraph(observations_text)
                                                para.paragraph_format.space_after = Pt(0)
                                            else:
                                                issue_text = item[1]
                                                issue_imgs = item[2]
                                                para = doc.add_paragraph(issue_text)
                                                para.paragraph_format.space_after = Pt(12)
                                                figure_counter = add_images_smart_layout(doc, issue_imgs, figure_counter, issue_text)
                                    
                                    # Pattern-level NOT OK images
                                    figure_counter = add_images_smart_layout(doc, test_data.get('not_ok_images', []), figure_counter, "Mechanical Inspection")
                                    
                                    if not issues and not custom_patterns:
                                        doc.add_paragraph("Mechanical inspection revealed issues but specifics were not provided.")
                                
                                else:
                                    if has_bore_data or has_movement_data:
                                        doc.add_paragraph("No status indicated.")
                                
                                if test_data.get('notes'):
                                    doc.add_heading('Additional Notes:', level=3)
                                    doc.add_paragraph(test_data['notes'])
                                
                                figure_counter = add_images_smart_layout(doc, test_data.get('general_images', []), figure_counter, "Images")
                        
                        else:
                            doc.add_heading(section.title, level=2)
                            
                            selected_patterns = section.get_selected_patterns()
                            
                            if selected_patterns:
                                is_valve_section = section.title in ["External Inspection", "Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]
                                
                                if is_valve_section and selected_patterns == ["Status: OK"]:
                                    # Just add the OK statement
                                    selected_obs = section.get_selected_observations("Status: OK")
                                    if selected_obs and len(selected_obs) > 0:
                                        doc.add_paragraph(selected_obs[0])
                                    
                                    # Status: OK images
                                    pattern_images = section.get_pattern_images("Status: OK")
                                    if pattern_images:
                                        figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                                
                                else:
                                    for pattern in selected_patterns:
                                        selected_obs = section.get_selected_observations(pattern)
                                        
                                        if selected_obs:
                                            if pattern == "Status: NOT OK":
                                                doc.add_heading('Observed Issues:', level=3)
                                            
                                            grouped_obs = []
                                            for obs in selected_obs:
                                                obs_images = section.get_observation_images(pattern, obs)
                                                if not obs_images:
                                                    grouped_obs.append(obs)
                                                else:
                                                    if grouped_obs:
                                                        para_text = " ".join(grouped_obs)
                                                        para = doc.add_paragraph(para_text)
                                                        para.paragraph_format.space_after = Pt(12)
                                                        grouped_obs = []
                                                    
                                                    para = doc.add_paragraph(obs)
                                                    para.paragraph_format.space_after = Pt(12)
                                                    
                                                    figure_counter = add_images_smart_layout(doc, obs_images, figure_counter, "Images")
                                            
                                            if grouped_obs:
                                                para_text = " ".join(grouped_obs)
                                                para = doc.add_paragraph(para_text)
                                                para.paragraph_format.space_after = Pt(12)
                                            
                                            custom_patterns = section.get_custom_patterns()
                                            if custom_patterns and selected_obs:
                                                custom_texts = [custom['text'] for custom in custom_patterns if custom.get('text')]
                                                if custom_texts:
                                                    last_para = None
                                                    for para in reversed(doc.paragraphs):
                                                        if para.text and para.style.name.startswith('Heading'):
                                                            continue
                                                        last_para = para
                                                        break
                                                    
                                                    if last_para:
                                                        last_para.add_run(" " + " ".join(custom_texts))
                                                    else:
                                                        para = doc.add_paragraph(" ".join(custom_texts))
                                                        para.paragraph_format.space_after = Pt(12)
                                            elif custom_patterns and not selected_obs:
                                                doc.add_heading('Observed Issues:', level=3)
                                                for custom in custom_patterns:
                                                    if custom.get('text'):
                                                        doc.add_paragraph(custom['text'], style='List Bullet')
                                        else:
                                            # No observations selected for Status: NOT OK
                                            if is_valve_section and pattern == "Status: NOT OK":
                                                custom_patterns = section.get_custom_patterns()
                                                has_custom = any(custom.get('text') for custom in custom_patterns)
                                                
                                                if not has_custom:
                                                    doc.add_paragraph(f"{section.title} revealed issues but specifics were not provided.")
                                        
                                        if pattern == "Status: NOT OK":
                                            pattern_images = section.get_pattern_images(pattern)
                                            if pattern_images:
                                                figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title if pattern in ["Status: OK", "Status: NOT OK"] else pattern)
                            else:
                                doc.add_paragraph(f"{section.title} not performed.")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                    
                    doc.add_heading('Valve Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=3)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set table width to full page width
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    # Set table width to 100% of page width (5000 = 100%)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)

                    for col in summary_table.columns:
                        col.width = Inches(2.375)
                    
                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'External Inspection'
                    header_cells[1].text = 'Electrical and Mechanical Inspection'
                    header_cells[2].text = 'Internal Inspection'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    
                    # External Inspection
                    external_items = []
                    external_status = []
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type == 'component' and section.title == "External Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        external_status.extend(selected_obs)
                                    else:
                                        external_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        external_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    external_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                external_items.append(notes)
                    
                    # Build external inspection text with issue count
                    body_cells[0].text = ''
                    if external_items:
                        para = body_cells[0].paragraphs[0]
                        count = len(external_items)
                        run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para.add_run(' '.join(external_items))
                        run.font.bold = False
                    elif external_status:
                        para = body_cells[0].paragraphs[0]
                        run = para.add_run(' '.join(external_status))
                        run.font.bold = False
                    else:
                        body_cells[0].text = 'No issues detected.'
                    
                    # Electrical and Mechanical Inspection
                    body_cells[1].text = ''
                    first_section = True
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type == 'valve_electrical':
                            test_data = section.get_test_data()
                            if test_data['ok']:
                                if not first_section:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_section = False
                                
                                run = para1.add_run("Electrical: ")
                                run.font.bold = True
                                run = para1.add_run("OK")
                                run.font.bold = False
                            elif test_data['not_ok']:
                                if not first_section:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_section = False
                                
                                issues_list = []
                                if test_data.get('issues'):
                                    issues_list.extend(test_data['issues'])
                                custom_patterns = test_data.get('custom_patterns', [])
                                for custom in custom_patterns:
                                    if custom.get('text'):
                                        issues_list.append(custom['text'])
                                
                                if issues_list:
                                    count = len(issues_list)
                                    issue_word = "Issue" if count == 1 else "Issues"
                                    run = para1.add_run(f"Electrical ({count} {issue_word}): ")
                                    run.font.bold = True
                                    run = para1.add_run(' '.join(issues_list))
                                    run.font.bold = False
                                else:
                                    run = para1.add_run("Electrical: ")
                                    run.font.bold = True
                                    run = para1.add_run("Issues detected")
                                    run.font.bold = False
                            
                            notes = test_data.get('notes', '')
                            if notes:
                                para1 = body_cells[1].add_paragraph()
                                run = para1.add_run(notes)
                                run.font.bold = False
                            
                            if test_data.get('resistance'):
                                para1 = body_cells[1].add_paragraph()
                                run = para1.add_run("Resistance: ")
                                run.font.bold = True
                                run = para1.add_run(f"{test_data['resistance']} Ω")
                                run.font.bold = False
                        
                        elif section_type == 'valve_mechanical':
                            test_data = section.get_test_data()
                            if test_data['ok']:
                                if not first_section:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_section = False
                                
                                run = para1.add_run("Mechanical: ")
                                run.font.bold = True
                                run = para1.add_run("OK")
                                run.font.bold = False
                            elif test_data['not_ok']:
                                if not first_section:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_section = False
                                
                                issues_list = []
                                if test_data.get('issues'):
                                    issues_list.extend(test_data['issues'])
                                custom_patterns = test_data.get('custom_patterns', [])
                                for custom in custom_patterns:
                                    if custom.get('text'):
                                        issues_list.append(custom['text'])
                                
                                if issues_list:
                                    count = len(issues_list)
                                    issue_word = "Issue" if count == 1 else "Issues"
                                    run = para1.add_run(f"Mechanical ({count} {issue_word}): ")
                                    run.font.bold = True
                                    run = para1.add_run(' '.join(issues_list))
                                    run.font.bold = False
                                else:
                                    run = para1.add_run("Mechanical: ")
                                    run.font.bold = True
                                    run = para1.add_run("Issues detected")
                                    run.font.bold = False
                            
                            notes = test_data.get('notes', '')
                            if notes:
                                para1 = body_cells[1].add_paragraph()
                                run = para1.add_run(notes)
                                run.font.bold = False
                            
                            if test_data.get('bore_top_bottom') or test_data.get('bore_left_right'):
                                para1 = body_cells[1].add_paragraph()
                                if test_data.get('bore_top_bottom'):
                                    run = para1.add_run("TB: ")
                                    run.font.bold = True
                                    run = para1.add_run(f"{test_data['bore_top_bottom']}mm")
                                    run.font.bold = False
                                    if test_data.get('bore_left_right'):
                                        run = para1.add_run(", ")
                                        run.font.bold = False
                                if test_data.get('bore_left_right'):
                                    run = para1.add_run("LR: ")
                                    run.font.bold = True
                                    run = para1.add_run(f"{test_data['bore_left_right']}mm")
                                    run.font.bold = False
                    
                    if first_section:
                        body_cells[1].text = 'No data recorded.'
                    
                    # Internal Inspection
                    internal_items = []
                    internal_status = []
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title in ["Internal Cylinder Inspection A", "Internal Cylinder Inspection B"]:
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                # Separate status patterns from actual issues
                                if pattern in ["No issues detected", "Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        internal_status.extend(selected_obs)
                                    else:
                                        internal_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        internal_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    internal_items.append(custom['text'])
                    
                    # Build internal inspection text with issue count
                    body_cells[2].text = ''
                    if internal_items:
                        para = body_cells[2].paragraphs[0]
                        count = len(internal_items)
                        run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para.add_run(' '.join(internal_items))
                        run.font.bold = False
                    elif internal_status:
                        para = body_cells[2].paragraphs[0]
                        run = para.add_run(' '.join(internal_status))
                        run.font.bold = False
                    else:
                        body_cells[2].text = 'No issues detected.'
                    
                    doc.add_paragraph()
                
                elif isinstance(widget, PCBTab):
                    # PCB/Inverter Board Report Generation
                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    doc.add_heading(f'PCB/Inverter Board {pcb_counter}', level=1)
                    pcb_counter += 1
                    
                    sections = widget.get_all_sections()
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        
                        if section_type == 'component':
                            # Visual Inspection Section
                            doc.add_heading(section.title, level=2)
                            
                            selected_patterns = section.get_selected_patterns()
                            
                            if "Status: OK" in selected_patterns:
                                observations = section.get_selected_observations("Status: OK")
                                for obs in observations:
                                    doc.add_paragraph(obs)
                                pattern_images = section.get_pattern_images("Status: OK")
                                if pattern_images:
                                    figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                            
                            if "Status: NOT OK" in selected_patterns:
                                selected_obs = section.get_selected_observations("Status: NOT OK")
                                
                                if selected_obs:
                                    para_text = " ".join(selected_obs)
                                    para = doc.add_paragraph(para_text)
                                    para.paragraph_format.space_after = Pt(12)
                                
                                custom_patterns = section.get_custom_patterns()
                                if custom_patterns and selected_obs:
                                    custom_texts = [custom['text'] for custom in custom_patterns if custom.get('text')]
                                    if custom_texts:
                                        last_para = None
                                        for para in reversed(doc.paragraphs):
                                            if para.text and para.style.name.startswith('Heading'):
                                                continue
                                            last_para = para
                                            break
                                        if last_para:
                                            last_para.add_run(" " + " ".join(custom_texts))
                                        else:
                                            doc.add_paragraph(" ".join(custom_texts))
                                elif custom_patterns and not selected_obs:
                                    for custom in custom_patterns:
                                        if custom.get('text'):
                                            doc.add_paragraph(custom['text'], style='List Bullet')
                                
                                pattern_images = section.get_pattern_images("Status: NOT OK")
                                if pattern_images:
                                    figure_counter = add_images_smart_layout(doc, pattern_images, figure_counter, section.title)
                            
                            if not selected_patterns:
                                doc.add_paragraph("Visual inspection not performed.")
                            
                            notes = section.get_notes()
                            if notes:
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(notes)
                            
                            images = section.get_images()
                            if images:
                                doc.add_heading('General Section Images:', level=3)
                                figure_counter = add_images_smart_layout(doc, images, figure_counter, f"{section.title} - General Section Images")
                        
                        elif section_type == 'pcb_electrical':
                            doc.add_heading('Electrical Testing', level=2)
                            test_data = section.get_test_data()
                            korea_testing = test_data.get('korea_testing', False)
                            
                            if test_data['ok']:
                                doc.add_paragraph("Diode testing completed with all readings within specification. No failures detected.")
                                figure_counter = add_images_smart_layout(doc, test_data.get('ok_images', []), figure_counter, "Electrical Testing")
                            
                            elif test_data['not_ok']:
                                issues = test_data.get('issues', [])
                                issue_images = test_data.get('issue_images', {})
                                custom_patterns = test_data.get('custom_patterns', [])
                                
                                if issues or custom_patterns:
                                    grouped_items = []
                                    current_group = []
                                    
                                    for issue in issues:
                                        if issue in issue_images and issue_images[issue]:
                                            if current_group:
                                                grouped_items.append(("text", current_group))
                                                current_group = []
                                            grouped_items.append(("image", issue, issue_images[issue]))
                                        else:
                                            current_group.append(issue)
                                    
                                    for custom in custom_patterns:
                                        if custom.get('text'):
                                            if custom.get('images'):
                                                if current_group:
                                                    grouped_items.append(("text", current_group))
                                                    current_group = []
                                                grouped_items.append(("image", custom['text'], custom['images']))
                                            else:
                                                current_group.append(custom['text'])
                                    
                                    if current_group:
                                        grouped_items.append(("text", current_group))
                                    
                                    for item in grouped_items:
                                        if item[0] == "text":
                                            observations_text = " ".join(item[1])
                                            para = doc.add_paragraph(observations_text)
                                            para.paragraph_format.space_after = Pt(0)
                                        else:
                                            issue_text = item[1]
                                            issue_imgs = item[2]
                                            para = doc.add_paragraph(issue_text)
                                            para.paragraph_format.space_after = Pt(12)
                                            figure_counter = add_images_smart_layout(doc, issue_imgs, figure_counter, issue_text)
                                
                                figure_counter = add_images_smart_layout(doc, test_data.get('not_ok_images', []), figure_counter, "Electrical Testing")
                                
                                if not issues and not custom_patterns:
                                    doc.add_paragraph("Diode testing revealed failures but specifics were not provided.")
                            else:
                                doc.add_paragraph("Electrical testing not performed.")
                            
                            if test_data.get('notes'):
                                doc.add_heading('Additional Notes:', level=3)
                                doc.add_paragraph(test_data['notes'])
                            
                            figure_counter = add_images_smart_layout(doc, test_data.get('general_images', []), figure_counter, "Electrical Testing - General Images")
                            
                            if korea_testing:
                                doc.add_paragraph("Sent for live testing in Korea.")
                    
                    # PCB Summary Table
                    doc.add_heading('PCB/Inverter Board Summary', level=2)
                    
                    summary_table = doc.add_table(rows=2, cols=2)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)

                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)

                    for col in summary_table.columns:
                        col.width = Inches(3.575)
                    
                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Visual Inspection'
                    header_cells[1].text = 'Electrical Testing'
                    
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    body_cells = summary_table.rows[1].cells
                    
                    # Visual Inspection summary
                    visual_items = []
                    visual_status = []
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'component' and section.title == "Visual Inspection":
                            selected_patterns = section.get_selected_patterns()
                            for pattern in selected_patterns:
                                if pattern in ["Status: OK"]:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_status.extend(selected_obs)
                                    else:
                                        visual_status.append(pattern)
                                else:
                                    selected_obs = section.get_selected_observations(pattern)
                                    if selected_obs:
                                        visual_items.extend(selected_obs)
                            custom_patterns = section.get_custom_patterns()
                            for custom in custom_patterns:
                                if custom.get('text'):
                                    visual_items.append(custom['text'])
                            notes = section.get_notes()
                            if notes:
                                visual_items.append(notes)
                    
                    body_cells[0].text = ''
                    if visual_items:
                        para = body_cells[0].paragraphs[0]
                        count = len(visual_items)
                        run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                        run.font.bold = True
                        run = para.add_run(' '.join(visual_items))
                        run.font.bold = False
                    elif visual_status:
                        para = body_cells[0].paragraphs[0]
                        run = para.add_run(' '.join(visual_status))
                        run.font.bold = False
                    else:
                        body_cells[0].text = 'No issues detected.'
                    
                    # Electrical Testing summary
                    body_cells[1].text = ''
                    first_elec = True
                    elec_korea = False
                    
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        if section_type == 'pcb_electrical':
                            test_data = section.get_test_data()
                            elec_korea = test_data.get('korea_testing', False)
                            
                            if test_data['ok']:
                                if not first_elec:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_elec = False
                                run = para1.add_run("Status: ")
                                run.font.bold = True
                                run = para1.add_run("OK - No diode failures detected")
                                run.font.bold = False
                            elif test_data['not_ok']:
                                if not first_elec:
                                    para1 = body_cells[1].add_paragraph()
                                else:
                                    para1 = body_cells[1].paragraphs[0]
                                first_elec = False
                                
                                issues_list = []
                                if test_data.get('issues'):
                                    issues_list.extend(test_data['issues'])
                                for custom in test_data.get('custom_patterns', []):
                                    if custom.get('text'):
                                        issues_list.append(custom['text'])
                                
                                if issues_list:
                                    count = len(issues_list)
                                    issue_word = "Issue" if count == 1 else "Issues"
                                    run = para1.add_run(f"Diode Failures ({count} {issue_word}): ")
                                    run.font.bold = True
                                    run = para1.add_run(' '.join(issues_list))
                                    run.font.bold = False
                                else:
                                    run = para1.add_run("Status: ")
                                    run.font.bold = True
                                    run = para1.add_run("NOT OK - Diode failures detected")
                                    run.font.bold = False
                            
                            notes = test_data.get('notes', '')
                            if notes:
                                para1 = body_cells[1].add_paragraph()
                                run = para1.add_run(notes)
                                run.font.bold = False
                    
                    if first_elec:
                        body_cells[1].text = 'No data recorded.'
                    
                    # Korea testing always last in summary
                    if elec_korea:
                        para1 = body_cells[1].add_paragraph()
                        run = para1.add_run("Sent for live testing in Korea.")
                        run.font.bold = False
                    
                    doc.add_paragraph()

                elif isinstance(widget, GeneralTab):
                    # General Component Report Generation
                    if first_equipment:
                        first_equipment = False
                    else:
                        doc.add_page_break()
                    component_label = f"{widget.component_name} {general_counter}" if widget.component_name else f"General {general_counter}"
                    doc.add_heading(component_label, level=1)
                    general_counter += 1

                    sections = widget.get_all_sections()

                    for section, _, section_type in sections:
                        doc.add_heading(section.title, level=2)
                        sec_data = section.get_section_data()

                        if sec_data['ok']:
                            doc.add_paragraph(f"{section.title} completed with no significant issues observed.")
                            figure_counter = add_images_smart_layout(doc, sec_data.get('ok_images', []), figure_counter, section.title)

                        elif sec_data['not_ok']:
                            custom_patterns = sec_data.get('custom_patterns', [])
                            issue_texts = [c['text'] for c in custom_patterns if c.get('text')]
                            if issue_texts:
                                para = doc.add_paragraph(' '.join(issue_texts))
                                para.paragraph_format.space_after = Pt(12)
                            else:
                                doc.add_paragraph(f"{section.title} issues identified but specifics were not provided.")
                            figure_counter = add_images_smart_layout(doc, sec_data.get('not_ok_images', []), figure_counter, f"{section.title} - Issues")

                        else:
                            doc.add_paragraph(f"{section.title} not performed.")

                        if sec_data.get('notes'):
                            doc.add_heading('Additional Notes:', level=3)
                            doc.add_paragraph(sec_data['notes'])

                        figure_counter = add_images_smart_layout(doc, sec_data.get('general_images', []), figure_counter, f"{section.title} - General Images")

                    # General Summary Table
                    summary_label = f"{widget.component_name} Summary" if widget.component_name else "General Summary"
                    doc.add_heading(summary_label, level=2)

                    summary_table = doc.add_table(rows=2, cols=2)
                    summary_table.style = 'Light Grid Accent 1'
                    summary_table.autofit = False
                    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    tbl = summary_table._element
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')
                    tblPr.append(tblW)
                    for col in summary_table.columns:
                        col.width = Inches(3.575)

                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = 'Visual Inspection'
                    header_cells[1].text = 'Testing'
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    body_cells = summary_table.rows[1].cells
                    for col_idx, section in enumerate(sections):
                        sec, _, _ = section
                        sec_data = sec.get_section_data()
                        body_cells[col_idx].text = ''
                        para = body_cells[col_idx].paragraphs[0]
                        if sec_data['ok']:
                            run = para.add_run("Status: ")
                            run.font.bold = True
                            run = para.add_run("OK")
                            run.font.bold = False
                        elif sec_data['not_ok']:
                            issues_list = [c['text'] for c in sec_data.get('custom_patterns', []) if c.get('text')]
                            if issues_list:
                                count = len(issues_list)
                                run = para.add_run(f"{count} issue{'s' if count != 1 else ''}: ")
                                run.font.bold = True
                                run = para.add_run(' '.join(issues_list))
                                run.font.bold = False
                            else:
                                run = para.add_run("Status: ")
                                run.font.bold = True
                                run = para.add_run("NOT OK")
                                run.font.bold = False
                        else:
                            body_cells[col_idx].text = 'Not performed.'
                        if sec_data.get('notes'):
                            note_para = body_cells[col_idx].add_paragraph()
                            note_para.add_run(sec_data['notes'])
                            bold = False

                    doc.add_paragraph()

            # Add disclaimer on last page (1 blank line, then disclaimer)
            doc.add_paragraph()
            
            last_disclaimer_heading_para = doc.add_paragraph()
            last_disclaimer_heading_run = last_disclaimer_heading_para.add_run("Disclaimer")
            last_disclaimer_heading_run.font.bold = True
            last_disclaimer_heading_run.font.size = Pt(11)
            
            doc.add_paragraph()  # blank line
            
            last_disclaimer_text = ("This report has been prepared for internal review purposes only and reflects the findings of a limited technical "
                "evaluation conducted on the returned equipment. The inspection and testing performed were based solely on the condition "
                "of the equipment as received and under the specific test conditions applied at the time of examination. Nothing in this "
                "report creates a guarantee of performance, warranty determination, or final root cause analysis. Findings are subject to "
                "revision if additional information, testing, or analysis becomes available. If this report is provided to a customer, it "
                "is intended for informational purposes only and should not be relied upon as a comprehensive assessment of system-wide "
                "performance, installation conditions, or operational practices beyond the scope of this review.")
            
            last_disclaimer_body_para = doc.add_paragraph()
            last_disclaimer_body_run = last_disclaimer_body_para.add_run(last_disclaimer_text)
            last_disclaimer_body_run.font.bold = False
            last_disclaimer_body_run.font.size = Pt(9)
            
            has_electrical_fluting = False
            for tab_idx in range(self.equipment_tabs_widget.count()):
                widget = self.equipment_tabs_widget.widget(tab_idx)
                if isinstance(widget, MotorTab):
                    sections = widget.get_all_sections()
                    for section_data in sections:
                        section, patterns_dict, section_type = section_data
                        # Check only bearing sections (Drive End and Non-Drive End)
                        if section_type == 'component' and section.title in ["Drive End Bearing", "Non-Drive End Bearing"]:
                            selected_patterns = section.get_selected_patterns()
                            if "Electrical Fluting (Bearing Flare)" in selected_patterns:
                                has_electrical_fluting = True
                                break
                if has_electrical_fluting:
                    break
            

            base_filename = self.equipment_id.text()
            if has_electrical_fluting:
                base_filename += " (needs further review)"
            

            filename, _ = QFileDialog.getSaveFileName(
                self, "Save Report", 
                base_filename + ".docx",
                "Word Documents (*.docx)"
            )
            
            if filename:

                if not filename.endswith('.docx'):
                    filename += '.docx'
                
                self.add_footer_with_page_numbers(doc, self.equipment_id.text())

                doc.save(filename)
                
                pdf_filename = filename.replace('.docx', '.pdf')
                pdf_success = False
                pdf_error = None
                
                # PDF Conversion using LibreOffice (fast and reliable)
                try:
                    import subprocess
                    import platform
                    
                    if platform.system() == 'Windows':
                        libreoffice_paths = [
                            r'C:\Program Files\LibreOffice\program\soffice.exe',
                            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                        ]
                    elif platform.system() == 'Darwin':
                        libreoffice_paths = ['/Applications/LibreOffice.app/Contents/MacOS/soffice']
                    else:
                        libreoffice_paths = ['libreoffice', 'soffice']
                    
                    libreoffice_cmd = None
                    for path in libreoffice_paths:
                        if os.path.exists(path) if '/' in path or '\\' in path else True:
                            libreoffice_cmd = path
                            break
                    
                    if not libreoffice_cmd:
                        raise FileNotFoundError("LibreOffice not found. Please install LibreOffice to enable PDF conversion.")
                    
                    output_dir = os.path.dirname(filename)
                    
                    # Run LibreOffice conversion with 30 second timeout
                    result = subprocess.run([
                        libreoffice_cmd,
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', output_dir,
                        filename
                    ], check=True, capture_output=True, timeout=30)
                    
                    # Verify PDF was created
                    if not os.path.exists(pdf_filename):
                        raise FileNotFoundError("PDF was not created by LibreOffice")
                    
                    pdf_success = True
                except subprocess.TimeoutExpired:
                    pdf_error = "PDF conversion timed out (>30s). The document may be too large or complex."
                except FileNotFoundError as e:
                    pdf_error = str(e)
                except subprocess.CalledProcessError as e:
                    pdf_error = f"LibreOffice conversion failed: {e.stderr.decode() if e.stderr else 'Unknown error'}"
                except Exception as e:
                    pdf_error = f"PDF conversion error: {str(e)}"
                

                if pdf_success:
                    QMessageBox.information(
                        self, "Success", 
                        f"Report generated successfully!\n\nSaved as:\n{filename}\n{pdf_filename}"
                    )
                else:
                    error_details = pdf_error if pdf_error else "Unknown error"
                    
                    QMessageBox.warning(
                        self, "Partial Success", 
                        f"DOCX saved successfully:\n{filename}\n\n"
                        f"PDF conversion failed:\n{error_details}\n\n"
                        f"To enable PDF generation:\n"
                        f"1. Install LibreOffice (free):\n"
                        f"   - Download from libreoffice.org\n"
                        f"   - Restart this application after installation\n\n"
                        f"2. Manual conversion:\n"
                        f"   - Open the DOCX file in Word or LibreOffice\n"
                        f"   - File > Export As > PDF"
                    )
                
                # Auto-save to project file after successful report generation
                if self.has_matching_save_file():
                    try:
                        save_path = self.last_save_path
                        
                        base_name = os.path.splitext(save_path)[0]
                        parent_dir = os.path.dirname(base_name)
                        folder_name = os.path.basename(base_name)
                        
                        if platform.system() == 'Windows':
                            project_folder = base_name + "_project"
                        else:
                            project_folder = os.path.join(parent_dir, "." + folder_name + "_project")
                        
                        images_folder = os.path.join(project_folder, "images")
                        os.makedirs(images_folder, exist_ok=True)
                        
                        data = {
                            'version': '1.0',
                            'header': {
                                'rma': self.equipment_id.text(),
                                'inspection_date': self.inspection_date.text(),
                                'unit_model': self.unit_model.text(),
                                'unit_serial': self.unit_serial.text(),
                                'customer_email': self.customer_email.text(),
                                'total_field': self.total_field.text(),
                                'lab_tech': self.lab_tech.text(),
                                'manager_name': self.manager_name.text(),
                                'install_date': self.install_date.text(),
                                'failure_date': self.failure_date.text(),
                                'warranty_claim': self.warranty_claim.text(),
                                'part_number': self.part_number.text()
                            },
                            'equipment': []
                        }
                        
                        for i in range(self.equipment_tabs_widget.count()):
                            widget = self.equipment_tabs_widget.widget(i)
                            if isinstance(widget, MotorTab):
                                equipment_type = 'motor'
                            elif isinstance(widget, CoilTab):
                                equipment_type = 'coil'
                            elif isinstance(widget, ReceiverTab):
                                equipment_type = 'receiver'
                            elif isinstance(widget, ValveTab):
                                equipment_type = 'valve'
                            elif isinstance(widget, PCBTab):
                                equipment_type = 'pcb'
                            elif isinstance(widget, GeneralTab):
                                equipment_type = 'general'
                            else:
                                equipment_type = 'compressor'
                            
                            equipment_data = {
                                'type': equipment_type,
                                'name': self.equipment_tabs_widget.tabText(i)
                            }
                            
                            if isinstance(widget, MotorTab):
                                equipment_data['sections'] = self._save_motor_tab(widget, images_folder)
                            elif isinstance(widget, CoilTab):
                                equipment_data['sections'] = self._save_coil_tab(widget, images_folder)
                            elif isinstance(widget, ReceiverTab):
                                equipment_data['sections'] = self._save_receiver_tab(widget, images_folder)
                            elif isinstance(widget, ValveTab):
                                equipment_data['sections'] = self._save_valve_tab(widget, images_folder)
                            elif isinstance(widget, PCBTab):
                                equipment_data['sections'] = self._save_pcb_tab(widget, images_folder)
                            elif isinstance(widget, GeneralTab):
                                equipment_data['sections'] = self._save_general_tab(widget, images_folder)
                            else:
                                equipment_data['sections'] = self._save_compressor_tab(widget, images_folder)
                            
                            data['equipment'].append(equipment_data)
                        
                        json_path = os.path.join(project_folder, "report_state.json")
                        with open(json_path, 'w', encoding='utf-8') as f:
                            json.dump(data, f, indent=2, ensure_ascii=False)
                        
                        shutil.copy(json_path, save_path)
                        
                    except Exception as e:
                        # Silent failure - don't interrupt report generation success
                        print(f"Auto-save after report generation failed: {str(e)}")
        
        except Exception as e:
            QMessageBox.critical(
                self, "Error", 
                f"Error generating report:\n{str(e)}\n\nPlease ensure all required packages are installed."
            )

def main():
    try:
        app = QApplication(sys.argv)
        app.setStyle('Fusion')
        

        splash_widget = QWidget()
        splash_widget.setWindowFlags(Qt.SplashScreen | Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        splash_widget.setStyleSheet("background-color: #A6D3C0; border: 3px solid black;")
        
        splash_layout = QVBoxLayout()
        splash_layout.setContentsMargins(20, 20, 20, 20)
        splash_layout.setSpacing(10)
        
        text_label = QLabel("Software by Daniel Fraser")
        text_label.setAlignment(Qt.AlignCenter)
        text_label.setStyleSheet("""
            QLabel {
                color: #2c3e50;
                font-size: 18px;
                font-weight: bold;
                background-color: transparent;
                padding: 10px;
                border: 2px solid black;
                border-radius: 5px;
            }
        """)
        splash_layout.addWidget(text_label)
        
        icon_label = QLabel()
        icon_label.setAlignment(Qt.AlignCenter)
        icon_label.setStyleSheet("background-color: transparent; border: none;")
        
        splash_path = None
        icon_path = None
        if getattr(sys, 'frozen', False):
            splash_path = os.path.join(sys._MEIPASS, '_assets', 'splash.png')
            icon_path = os.path.join(sys._MEIPASS, '_assets', 'icon.ico')
        else:
            splash_path = os.path.join(os.path.dirname(__file__), '_assets', 'splash.png')
            icon_path = os.path.join(os.path.dirname(__file__), '_assets', 'icon.ico')
        
        if os.path.exists(splash_path):
            pixmap = QPixmap(splash_path)
            if not pixmap.isNull():
                scaled_pixmap = pixmap.scaled(300, 300, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                icon_label.setPixmap(scaled_pixmap)
            else:
                icon_label.setText("🔧")
                icon_label.setStyleSheet("font-size: 120px; background-color: transparent; border: none;")
        else:

            icon_label.setText("🔧")
            icon_label.setStyleSheet("font-size: 120px; background-color: transparent; border: none;")
        
        splash_layout.addWidget(icon_label)
        
        splash_widget.setLayout(splash_layout)
        splash_widget.adjustSize()
        
        screen = app.primaryScreen().geometry()
        splash_x = (screen.width() - splash_widget.width()) // 2
        splash_y = (screen.height() - splash_widget.height()) // 2
        splash_widget.move(splash_x, splash_y)
        
        splash_widget.show()
        app.processEvents()
        

        from time import sleep
        sleep(2)
        
        splash_widget.close()
        

        if os.path.exists(icon_path):
            app.setWindowIcon(QIcon(icon_path))
        

        window = ReportGeneratorApp()
        
        # Check if a file path was passed as a command-line argument
        if len(sys.argv) > 1:
            file_path = sys.argv[1]
            # Load the file after the window is shown
            QTimer.singleShot(100, lambda: window.load_progress_from_path(file_path))
        
        if os.path.exists(icon_path):
            window.setWindowIcon(QIcon(icon_path))
        
        window.show()
        
        sys.exit(app.exec_())
    except Exception as e:
        print(f"Error starting application: {e}")
        print("\nPlease ensure you have installed all required packages:")
        print("pip install PyQt5 python-docx pillow")
        sys.exit(1)

if __name__ == '__main__':
    main()